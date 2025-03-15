import streamlit as st
import requests
import json
import pandas as pd
from datetime import datetime
import plotly.express as px
import os
import re
import docx
from io import BytesIO
import time
import functools

# ================= CONFIGURATION =================

# Configuração da página - otimizado para mobile e desktop
st.set_page_config(
    page_title="NEXUS",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Definindo paleta de cores - simplificada e mais Apple-like
NEXUS_COLORS = {
    "primary": "#007AFF",       # Apple blue
    "secondary": "#5E5CE6",     # purple
    "accent": "#FF9500",        # orange
    "background": "#F5F5F7",    # light gray
    "card": "#FFFFFF",          # white
    "text_primary": "#1D1D1F",  # nearly black
    "text_secondary": "#86868B", # medium gray
    "success": "#34C759",       # green
    "warning": "#FF3B30",       # red
    "border": "#E5E5EA"         # light border
}

# Esconder elementos padrão do Streamlit e fazer o app ocupar toda largura
st.markdown("""
<style>
    /* Ocultar elementos Streamlit desnecessários */
    #MainMenu, footer, header {visibility: hidden;}
    [data-testid="collapsedControl"] {display: none;}
    section[data-testid="stSidebar"] {display: none;}
    
    /* Otimização para mobile */
    .main .block-container {
        padding-left: 1rem;
        padding-right: 1rem;
        max-width: 960px;
        margin: 0 auto;
    }
    
    /* Tipografia mais clara - estilo Apple */
    body {
        font-family: 'SF Pro Text', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        -webkit-font-smoothing: antialiased;
        color: #1D1D1F;
    }
    
    h1, h2, h3, h4, h5 {
        font-weight: 600;
    }
    
    /* Espaçamento e tamanho responsivos */
    @media (max-width: 640px) {
        .main .block-container {
            padding: 0.5rem;
        }
        
        .header-container h1 {
            font-size: 24px;
        }
        
        .feature-card {
            padding: 0.8rem;
        }
    }
</style>
""", unsafe_allow_html=True)

# CSS Aprimorado - Mais limpo, moderno e inspirado em design Apple
st.markdown("""
<style>
    /* Variáveis de cores */
    :root {
        --primary: #007AFF;
        --secondary: #5E5CE6;
        --accent: #FF9500;
        --background: #F5F5F7;
        --card: #FFFFFF;
        --text-primary: #1D1D1F;
        --text-secondary: #86868B;
        --success: #34C759;
        --warning: #FF3B30;
        --border: #E5E5EA;
    }
    
    /* Reset e estilos gerais */
    * {
        margin: 0;
        padding: 0;
        box-sizing: border-box;
    }
    
    /* Corpo principal */
    .main {
        background-color: var(--background);
        padding: 1rem;
    }
    
    .stApp {
        background-color: var(--background);
    }
    
    /* Header redesenhado - mais Apple-like */
    .header-container {
        background-color: var(--card);
        border-radius: 12px;
        padding: 1.2rem 1.5rem;
        margin-bottom: 1.5rem;
        display: flex;
        justify-content: space-between;
        align-items: center;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
        border: 1px solid var(--border);
    }
    
    .header-container h1 {
        margin: 0;
        color: var(--text-primary);
        font-size: 28px;
        font-weight: 600;
    }
    
    .header-status {
        display: flex;
        align-items: center;
        gap: 8px;
    }
    
    .status-indicator {
        width: 8px;
        height: 8px;
        border-radius: 50%;
    }
    
    .status-active {
        background-color: var(--success);
    }
    
    /* Cards de funcionalidades - redesenhados */
    .feature-grid {
        display: grid;
        grid-template-columns: repeat(auto-fill, minmax(280px, 1fr));
        gap: 1rem;
        margin-bottom: 1.5rem;
    }
    
    .feature-card {
        background-color: var(--card);
        border-radius: 12px;
        padding: 1.2rem;
        transition: all 0.2s ease;
        box-shadow: 0 1px 2px rgba(0,0,0,0.05);
        border: 1px solid var(--border);
        height: 100%;
        display: flex;
        flex-direction: column;
    }
    
    .feature-card:hover {
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    }
    
    .feature-card h3 {
        margin: 0;
        color: var(--text-primary);
        font-size: 18px;
        font-weight: 600;
        margin-bottom: 8px;
    }
    
    .feature-card p {
        color: var(--text-secondary);
        font-size: 14px;
        line-height: 1.4;
        margin-bottom: 1rem;
        flex-grow: 1;
    }
    
    .feature-card .feature-icon {
        font-size: 24px;
        margin-bottom: 12px;
    }
    
    /* Formulários - redesenhados para maior clareza */
    .form-container {
        background-color: var(--card);
        border-radius: 12px;
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        border: 1px solid var(--border);
    }
    
    .form-title {
        font-size: 18px;
        color: var(--text-primary);
        margin-bottom: 1.2rem;
        font-weight: 600;
    }
    
    /* Campos de entrada - estilo Apple */
    .stTextInput > div > div > input {
        border-radius: 8px;
        border: 1px solid var(--border);
        padding: 12px 16px;
        font-size: 15px;
        background-color: rgba(245, 245, 247, 0.5);
        transition: all 0.2s;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: var(--primary);
        box-shadow: 0 0 0 2px rgba(0, 122, 255, 0.15);
        background-color: white;
    }
    
    .stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 1px solid var(--border);
        padding: 12px 16px;
        font-size: 15px;
        background-color: rgba(245, 245, 247, 0.5);
        transition: all 0.2s;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: var(--primary);
        box-shadow: 0 0 0 2px rgba(0, 122, 255, 0.15);
        background-color: white;
    }
    
    /* Botões - estilo Apple */
    .button-primary {
        background-color: var(--primary);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: 500;
        font-size: 15px;
        cursor: pointer;
        transition: all 0.2s;
        width: 100%;
        text-align: center;
        display: inline-block;
    }
    
    .button-secondary {
        background-color: rgba(0, 122, 255, 0.1);
        color: var(--primary);
        border: none;
        border-radius: 8px;
        padding: 12px 24px;
        font-weight: 500;
        font-size: 15px;
        cursor: pointer;
        transition: all 0.2s;
        width: 100%;
        text-align: center;
        display: inline-block;
    }
    
    .button-primary:hover {
        background-color: #0062CC;
    }
    
    .button-secondary:hover {
        background-color: rgba(0, 122, 255, 0.15);
    }
    
    /* Resultado */
    .result-area {
        background-color: var(--card);
        border-radius: 12px;
        padding: 1.5rem;
        margin: 1.5rem 0;
        border: 1px solid var(--border);
    }
    
    /* Histórico - redesenhado */
    .history-item {
        padding: 1rem;
        border-radius: 8px;
        margin-bottom: 0.8rem;
        background-color: var(--card);
        border: 1px solid var(--border);
        transition: all 0.2s;
    }
    
    .history-item:hover {
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    
    /* Responsividade aprimorada */
    @media (max-width: 768px) {
        .feature-grid {
            grid-template-columns: 1fr;
        }
        
        .form-container, .result-area {
            padding: 1rem;
        }
        
        .button-primary, .button-secondary {
            padding: 10px 16px;
            font-size: 14px;
        }
    }
</style>
""", unsafe_allow_html=True)

# Base de conhecimento sobre PMBOK 7 - mantida como um módulo separado com lazy loading
@functools.lru_cache(maxsize=1)
def get_pmbok_knowledge():
    return {
        "princípios": [
            "1. Ser um administrador diligente, respeitoso e cuidadoso",
            "2. Criar um ambiente colaborativo da equipe do projeto",
            "3. Envolver efetivamente as partes interessadas",
            "4. Focar no valor",
            "5. Reconhecer, avaliar e responder às interações do sistema",
            "6. Demonstrar comportamentos de liderança",
            "7. Adaptar com base no contexto",
            "8. Incorporar qualidade nos processos e resultados",
            "9. Navegar na complexidade",
            "10. Otimizar respostas a riscos",
            "11. Abraçar adaptabilidade e resiliência",
            "12. Permitir mudança para alcançar o estado futuro previsto"
        ],
        
        "domínios": [
            "1. Stakeholders (Partes Interessadas)",
            "2. Team (Equipe)",
            "3. Development Approach and Life Cycle (Abordagem de Desenvolvimento e Ciclo de Vida)",
            "4. Planning (Planejamento)",
            "5. Project Work (Trabalho do Projeto)",
            "6. Delivery (Entrega)",
            "7. Measurement (Mensuração)",
            "8. Uncertainty (Incerteza)"
        ],
        
        "metodologias": {
            "preditiva": "Abordagem tradicional (cascata) com fases sequenciais",
            "adaptativa": "Abordagens ágeis e iterativas (Scrum, Kanban, etc.)",
            "híbrida": "Combinação de elementos preditivos e adaptativos"
        },
        
        "mudancas_principais": [
            "1. Transição de processos para princípios e domínios de performance",
            "2. Foco em entrega de valor em vez de apenas escopo, tempo e custo",
            "3. Maior ênfase em adaptabilidade e contexto",
            "4. Abordagem de sistemas em vez de processos isolados",
            "5. Reconhecimento de múltiplas abordagens (adaptativa, preditiva, híbrida)",
            "6. Maior ênfase na liderança e soft skills",
            "7. Visão holística do gerenciamento de projetos"
        ]
    }

# Definição de funcionalidades disponíveis - refatorado para mais eficiência
FEATURE_OPTIONS = {
    "Gerador de Comunicações": {
        "description": "Crie e-mails, relatórios e comunicados profissionais",
        "icon": "📧",
        "color": "primary",
        "subtypes": ["E-mail Profissional", "Relatório de Status", "Comunicado Formal"]
    },
    "Assistente de Reuniões": {
        "description": "Gere agendas, atas e follow-ups estruturados",
        "icon": "📅",
        "color": "secondary",
        "subtypes": ["Agenda de Reunião", "Ata de Reunião", "Follow-up de Reunião"]
    },
    "Tradutor de Jargão": {
        "description": "Simplifique linguagem técnica para diferentes públicos",
        "icon": "🔄",
        "color": "accent",
        "subtypes": ["Simplificação Técnica", "Adaptação para Executivos", "Adaptação para Clientes", "Adaptação para Equipe Técnica"]
    },
    "Facilitador de Feedback": {
        "description": "Estruture feedbacks construtivos e conversas difíceis",
        "icon": "💬",
        "color": "primary",
        "subtypes": ["Feedback de Desempenho", "Feedback sobre Entregáveis", "Roteiro para Conversa"]
    },
    "Detector de Riscos": {
        "description": "Identifique ambiguidades e problemas potenciais",
        "icon": "🔍",
        "color": "secondary",
        "subtypes": ["Análise de E-mail", "Análise de Comunicado", "Análise de Proposta", "Análise de Requisitos"]
    },
    "Consultor PMBOK 7": {
        "description": "Obtenha orientações sobre gerenciamento de projetos",
        "icon": "📚",
        "color": "accent",
        "subtypes": ["Princípios", "Domínios", "Metodologias", "Ferramentas", "Melhores Práticas"]
    }
}

# Limites do sistema - mais realistas para evitar problemas de desempenho
TOKEN_LIMIT = 50000     # Reduzido para evitar sobrecarga
REQUEST_LIMIT = 30      # Valor mais realista baseado em uso típico

# ================= SESSION STATE INITIALIZATION =================

# Inicialização da sessão otimizada - usando funções auxiliares
def init_session_state():
    """Inicializa todas as variáveis do session_state de uma vez"""
    
    # Variáveis básicas
    defaults = {
        'api_key_configured': True,
        'usage_data': [],
        'generated_content': "",
        'history': [],
        'token_count': 0,
        'request_count': 0,
        'last_request_time': 0,
        'session_id': datetime.now().strftime("%Y%m%d%H%M%S") + str(hash(datetime.now()))[:5],
        'current_feature': "",
        'optimized_content': "",
        'previous_screen': None
    }
    
    # Aplicar apenas valores ausentes
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value
    
    # Configurar API key automaticamente a partir dos secrets
    if 'api_key' not in st.session_state:
        if st.secrets.get("OPENAI_API_KEY"):
            st.session_state.api_key = st.secrets.get("OPENAI_API_KEY")
        else:
            st.session_state.api_key = None
            st.session_state.api_key_configured = False

# Inicializar o session state ao carregar o app
init_session_state()

# ================= HELPER FUNCTIONS =================

def header():
    """Renderiza o cabeçalho com design minimalista estilo Apple"""
    
    st.markdown("""
    <div class="header-container">
        <h1>NEXUS</h1>
        <div class="header-status">
            <div class="status-indicator status-active"></div>
            <span style="font-size:14px; color:var(--text-secondary);">Conectado</span>
        </div>
    </div>
    """, unsafe_allow_html=True)

def enrich_pmbok_prompt(prompt, pmbok_topic):
    """Enriquece o prompt com informações relevantes do PMBOK 7 baseado no tópico selecionado"""
    
    pmbok_data = get_pmbok_knowledge()  # Lazy loading com cache
    additional_info = ""
    
    if "Princípios" in pmbok_topic:
        additional_info += "\n\nPrincípios de Gerenciamento do PMBOK 7:\n"
        additional_info += "\n".join(pmbok_data["princípios"])
        
    elif "Domínios" in pmbok_topic:
        additional_info += "\n\nDomínios de Performance do PMBOK 7:\n"
        additional_info += "\n".join(pmbok_data["domínios"])
        
    elif "Metodologias" in pmbok_topic:
        additional_info += "\n\nAbordagens de Desenvolvimento no PMBOK 7:\n"
        for k, v in pmbok_data["metodologias"].items():
            additional_info += f"- {k.capitalize()}: {v}\n"
            
    elif "Melhores Práticas" in pmbok_topic:
        additional_info += "\n\nMudanças Principais no PMBOK 7:\n"
        additional_info += "\n".join([f"- {item}" for item in pmbok_data["mudancas_principais"]])
    
    # Adicionar a informação relevante ao prompt
    return prompt + additional_info

# Função para exportar conteúdo como DOCX - otimizada
@st.cache_data  # Usar cache para documentos gerados
def export_as_docx(content, filename="documento"):
    """Exporta o conteúdo para um documento DOCX"""
    
    doc = docx.Document()
    
    # Adicionar título
    doc.add_heading(f"{filename}", 0)
    
    # Dividir por linhas e adicionar parágrafos
    paragraphs = content.split('\n')
    for para in paragraphs:
        if not para.strip():
            continue
        
        # Verificar se é um cabeçalho
        header_match = re.match(r'^(#{1,6})\s+(.*)', para)
        if header_match:
            level = min(len(header_match.group(1)), 9)  # Limitar a 9 para evitar erro
            text = header_match.group(2)
            doc.add_heading(text, level)
        else:
            doc.add_paragraph(para)
    
    # Salvar para um buffer em memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Função para gerar conteúdo via API OpenAI - refatorada para melhor gerenciamento de erros e performance
def generate_content(prompt, model="gpt-3.5-turbo", temperature=0.7):
    """Gera conteúdo via API OpenAI com tratamento de erros aprimorado"""
    
    # Verificações básicas
    if not st.session_state.api_key_configured or not st.session_state.api_key:
        return "API não configurada. Por favor, contate o administrador."
    
    # Verificar limites
    if st.session_state.token_count >= TOKEN_LIMIT:
        return "Você atingiu o limite de tokens para esta sessão. Por favor, tente novamente mais tarde."
    
    if st.session_state.request_count >= REQUEST_LIMIT:
        return "Você atingiu o limite de requisições para esta sessão. Por favor, tente novamente mais tarde."
    
    try:
        with st.spinner("Gerando conteúdo..."):
            # Atualizar o timestamp da última requisição
            st.session_state.last_request_time = time.time()
            
            # Incrementar contador de requisições
            st.session_state.request_count += 1
            
            # Configurar requisição direta à API
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {st.session_state.api_key}"
            }
            
            # System prompt mais eficaz e direcionado
            system_prompt = """
            Você é o NEXUS, um assistente especializado em comunicação estratégica e gerenciamento de projetos.
            Forneça respostas profissionais, estruturadas e objetivas, focando em clareza e valor prático.
            """
            
            # Payload otimizado
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                "temperature": temperature,
                "max_tokens": 3000  # Reduzido para melhorar performance 
            }
            
            # Fazer a requisição à API com timeout apropriado
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                data=json.dumps(payload),
                timeout=30  # Reduzido para melhorar experiência do usuário
            )
            
            # Processar a resposta
            if response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                # Atualizar contadores de tokens
                prompt_tokens = result['usage']['prompt_tokens']
                completion_tokens = result['usage']['completion_tokens']
                total_tokens = result['usage']['total_tokens']
                st.session_state.token_count += total_tokens
                
                # Registrar uso - apenas dados essenciais
                st.session_state.usage_data.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'feature': st.session_state.current_feature,
                    'tokens': total_tokens,
                    'model': model
                })
                
                # Adicionar ao histórico - versão otimizada
                st.session_state.history.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'feature': st.session_state.current_feature,
                    'input': prompt[:100] + "..." if len(prompt) > 100 else prompt,
                    'output': content[:100] + "..." if len(content) > 100 else content,
                    'model': model
                })
                
                return content
            else:
                # Mensagem de erro mais detalhada e útil
                error_msg = response.json().get('error', {}).get('message', f"Erro {response.status_code}")
                return f"Erro na requisição: {error_msg}. Por favor, tente novamente."
        
    except requests.exceptions.Timeout:
        return "Tempo limite excedido ao tentar conectar com a API. Verifique sua conexão e tente novamente."
    except requests.exceptions.ConnectionError:
        return "Erro de conexão. Verifique sua internet e tente novamente."
    except Exception as e:
        return f"Erro inesperado: {str(e)}"

# Função para criar cards de funcionalidades - redesenhada com grid mais responsivo
def create_feature_cards():
    """Cria os cards de seleção de funcionalidades com design inspirado na Apple"""
    
    st.markdown('<div class="feature-grid">', unsafe_allow_html=True)
    
    for name, details in FEATURE_OPTIONS.items():
        icon = details["icon"]
        description = details["description"]
        color = details["color"]
        
        # HTML/CSS para criar cards mais modernos e responsivos
        card_html = f"""
        <div class="feature-card" id="card_{name.lower().replace(' ', '_')}">
            <div class="feature-icon" style="color: var(--{color});">{icon}</div>
            <h3>{name}</h3>
            <p>{description}</p>
            <div style="text-align: right;">
                <button 
                    onclick="document.getElementById('button_{name.lower().replace(' ', '_')}').click()"
                    style="background-color: var(--{color}); color: white; border: none; 
                           border-radius: 6px; padding: 8px 16px; font-size: 14px;
                           font-weight: 500; cursor: pointer; transition: all 0.2s;">
                    Selecionar
                </button>
            </div>
        </div>
        """
        
        st.markdown(card_html, unsafe_allow_html=True)
        
        # Botão oculto para capturar o clique
        if st.button("Selecionar", key=f"button_{name.lower().replace(' ', '_')}", style="display: none;"):
            st.session_state.current_feature = name
            st.session_state.previous_screen = "home"
            st.rerun()

# Interface principal do aplicativo - completamente redesenhada
def main():
    """Função principal do aplicativo - redesenhada para maior eficiência e UX aprimorada"""
    
    # Renderizar o cabeçalho minimalista
    header()
    
    # Botão VOLTAR quando estiver em uma funcionalidade - redesenhado
    if st.session_state.current_feature:
        if st.button("◀️ Voltar", key="back_to_home", 
                    help="Voltar à tela inicial", type="secondary"):
            st.session_state.current_feature = ""
            st.session_state.previous_screen = None
            st.rerun()
    
    # Exibir mensagem sobre versão e uso do app
    with st.expander("ℹ️ Sobre o NEXUS", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"""
            **NEXUS - Assistente de Comunicação**
            
            Facilite suas comunicações de projeto com IA:
            * E-mails e relatórios profissionais
            * Agendas e atas de reuniões
            * Simplificação de linguagem técnica
            * Estruturação de feedback
            """)
        with col2:
            st.markdown(f"""
            **Uso atual:**
            * Requisições: {st.session_state.request_count}/{REQUEST_LIMIT}
            * Tokens: {st.session_state.token_count}/{TOKEN_LIMIT}
            
            Projeto desenvolvido por Ricardo Brigante
            """)
    
    # Histórico simplificado - apenas quando relevante
    if st.session_state.history and len(st.session_state.history) > 0:
        with st.expander("Histórico recente", expanded=False):
            # Mostrar apenas os últimos 3 itens para melhor performance
            for i, item in enumerate(reversed(st.session_state.history[-3:])):
                feature = item['feature']
                timestamp = item['timestamp']
                
                st.markdown(f"""
                <div class="history-item">
                    <p style="margin:0; font-weight:500;">{feature}</p>
                    <p style="margin:0; font-size:12px; color:var(--text-secondary);">{timestamp}</p>
                </div>
                """, unsafe_allow_html=True)
                
                col1, col2 = st.columns([3, 1])
                with col2:
                    if st.button(f"Carregar", key=f"load_{i}", type="secondary"):
                        st.session_state.current_feature = item['feature']
                        st.session_state.generated_content = item.get('output', "")
                        st.rerun()
    
    # Se nenhuma funcionalidade selecionada, mostrar a tela principal
    if not st.session_state.current_feature:
        st.markdown('<h2 style="font-size:24px; margin-bottom:16px;">Selecione uma funcionalidade</h2>', 
                   unsafe_allow_html=True)
        
        # Cards de funcionalidades redesenhados
        create_feature_cards()
        
    else:
        # Interface da funcionalidade selecionada
        current_feature = st.session_state.current_feature
        feature_details = FEATURE_OPTIONS[current_feature]
        
        # Título da funcionalidade
        st.markdown(f"""
        <h2 style="font-size:24px; margin-bottom:16px; display:flex; align-items:center; gap:8px;">
            <span style="color:var(--{feature_details['color']})">{feature_details['icon']}</span>
            {current_feature}
        </h2>
        """, unsafe_allow_html=True)
        
        # Verificar limites antes de mostrar a interface
        if st.session_state.token_count >= TOKEN_LIMIT:
            st.error(f"Limite de tokens atingido. Por favor, tente novamente mais tarde.")
        elif st.session_state.request_count >= REQUEST_LIMIT:
            st.error(f"Limite de requisições atingido. Por favor, tente novamente mais tarde.")
        else:
            # Interface da funcionalidade com formulário otimizado
            with st.form(key=f"{current_feature}_form", clear_on_submit=False):
                st.markdown('<div class="form-container">', unsafe_allow_html=True)
                
                # Campo de subtipo com descrições mais claras
                subtype_options = feature_details['subtypes']
                subtype = st.selectbox("Tipo de comunicação", 
                                      subtype_options,
                                      help="Selecione o tipo específico de conteúdo que você precisa")
                
                # Layout adaptativo - comportamento responsivo para mobile
                is_mobile = st.checkbox("Otimizar para visualização em dispositivo móvel", 
                                      value=False, 
                                      help="Ative para melhor visualização em telas pequenas",
                                      key="mobile_toggle")
                
                # Campos comuns a todas as funcionalidades (exceto PMBOK)
                context = ""
                if current_feature != "Consultor PMBOK 7":
                    context_height = 80 if is_mobile else 120
                    context = st.text_area("Contexto do projeto", 
                                        help="Descreva brevemente o projeto e sua situação atual",
                                        height=context_height,
                                        placeholder="Ex: Projeto de app mobile em fase de testes")
                
                # Prompt padrão que será personalizado por funcionalidade
                prompt = ""
                
                # Campos específicos por funcionalidade - redesenhados para UX melhorada
                if current_feature == "Gerador de Comunicações":
                    # Layout adaptativo para campos
                    field_height = 80 if is_mobile else 120
                    
                    audience = st.text_input("Destinatário", 
                                        help="Para quem esta comunicação será enviada",
                                        placeholder="Ex: Cliente, diretor de marketing")
                    
                    key_points = st.text_area("Pontos principais", 
                                            help="Liste os pontos essenciais a incluir",
                                            height=field_height,
                                            placeholder="Ex: Atraso de 3 dias, plano de recuperação...")
                    
                    tone = st.select_slider("Tom da comunicação", 
                                        options=["Muito Formal", "Formal", "Neutro", "Amigável", "Casual"],
                                        value="Neutro")
                    
                    prompt = f"""
                    Gere um {subtype} com base nas seguintes informações:
                    
                    Contexto do Projeto: {context}
                    Destinatário: {audience}
                    Pontos principais: {key_points}
                    Tom desejado: {tone}
                    
                    Formate a saída adequadamente para um {subtype}, incluindo assunto/título e estrutura apropriada.
                    """
                
                elif current_feature == "Assistente de Reuniões":
                    # Layout adaptativo para campos
                    field_height = 80 if is_mobile else 120
                    
                    participants = st.text_area("Participantes", 
                                            help="Liste os participantes e suas funções",
                                            height=field_height,
                                            placeholder="Ex: João (Gerente), Maria (Dev)")
                    
                    topics = st.text_area("Tópicos", 
                                        help="Liste os tópicos a serem discutidos",
                                        height=field_height,
                                        placeholder="Ex: Cronograma, bugs pendentes, feedback")
                    
                    # Valor padrão mais realista
                    duration = st.number_input("Duração (minutos)", 
                                            min_value=15, 
                                            max_value=180, 
                                            value=30, 
                                            step=15)
                    
                    if subtype == "Agenda de Reunião":
                        prompt = f"""
                        Crie uma agenda para reunião de {duration} minutos com:
                        
                        Contexto: {context}
                        Participantes: {participants}
                        Tópicos: {topics}
                        
                        Inclua alocação de tempo e objetivos claros para cada item.
                        """
                    elif subtype == "Ata de Reunião":
                        # Campos adicionais apenas quando relevantes
                        decisions = st.text_area("Decisões tomadas", 
                                            height=field_height,
                                            placeholder="Ex: Aprovação do design, extensão de prazo")
                        
                        actions = st.text_area("Ações acordadas", 
                                            height=field_height,
                                            placeholder="Ex: João: corrigir bug #123 até amanhã")
                        
                        prompt = f"""
                        Crie uma ata detalhada para reunião de {duration} minutos:
                        
                        Contexto: {context}
                        Participantes: {participants}
                        Tópicos: {topics}
                        Decisões: {decisions}
                        Ações: {actions}
                        
                        Organize por tópicos, destacando decisões e próximos passos.
                        """
                    else:  # Follow-up
                        meeting_outcome = st.text_area("Resultado da reunião", 
                                                height=field_height,
                                                placeholder="Ex: Definidas prioridades para o sprint")
                        
                        action_items = st.text_area("Itens de ação", 
                                                height=field_height,
                                                placeholder="Ex: João: revisão até 25/03")
                        
                        prompt = f"""
                        Crie uma mensagem de follow-up para reunião:
                        
                        Contexto: {context}
                        Participantes: {participants}
                        Tópicos: {topics}
                        Resultado: {meeting_outcome}
                        Ações: {action_items}
                        
                        A mensagem deve agradecer a participação, resumir pontos principais e detalhar próximos passos.
                        """
                
                elif current_feature == "Tradutor de Jargão":
                    # Área de texto maior para conteúdo técnico
                    content_height = 150 if is_mobile else 200
                    
                    technical_content = st.text_area("Conteúdo técnico", 
                                                help="Cole o texto técnico a ser traduzido",
                                                height=content_height,
                                                placeholder="Cole aqui o texto com jargão técnico...")
                    
                    audience = st.selectbox("Público-alvo", 
                                        ["Executivos", "Clientes não-técnicos", "Equipe de Negócios", "Equipe Técnica Junior"])
                    
                    key_concepts = st.text_input("Conceitos-chave a preservar", 
                                            placeholder="Ex: API, front-end, estado")
                    
                    prompt = f"""
                    Adapte o seguinte conteúdo técnico para {audience}:
                    
                    Contexto: {context}
                    Conteúdo Original: {technical_content}
                    Conceitos a preservar: {key_concepts}
                    
                    Para {audience}, foque em: 
                    {'Impacto nos negócios e resultados' if audience == 'Executivos' else
                    'Benefícios e funcionalidades em linguagem acessível' if audience == 'Clientes não-técnicos' else
                    'Conexão com objetivos de negócios' if audience == 'Equipe de Negócios' else
                    'Explicações técnicas mais detalhadas com conceitos explicados'}
                    
                    Mantenha a precisão conceitual mesmo simplificando a linguagem.
                    """
                
                elif current_feature == "Facilitador de Feedback":
                    # Layout adaptativo
                    field_height = 80 if is_mobile else 120
                    
                    situation = st.text_area("Situação", 
                                        help="Descreva a situação específica para feedback",
                                        height=field_height,
                                        placeholder="Ex: Atraso na entrega de componentes...")
                    
                    strengths = st.text_area("Pontos fortes", 
                                        height=field_height,
                                        placeholder="Ex: Qualidade do código, comunicação")
                    
                    areas_for_improvement = st.text_area("Áreas para melhoria", 
                                                    height=field_height,
                                                    placeholder="Ex: Estimativas irrealistas, bloqueios")
                    
                    relationship = st.selectbox("Relação com o receptor", 
                                            ["Membro da equipe", "Colega", "Superior", "Cliente", "Fornecedor"])
                    
                    prompt = f"""
                    Estruture um {subtype} construtivo com base em:
                    
                    Contexto: {context}
                    Situação: {situation}
                    Pontos fortes: {strengths}
                    Áreas de melhoria: {areas_for_improvement}
                    Relação: {relationship}
                    
                    O feedback deve ser específico, equilibrado, incluir exemplos e oferecer sugestões acionáveis.
                    Formate como um roteiro que o usuário pode seguir na conversa.
                    """
                
                elif current_feature == "Detector de Riscos":
                    # Área maior para o conteúdo a analisar
                    content_height = 150 if is_mobile else 200
                    
                    content_to_analyze = st.text_area("Conteúdo para análise", 
                                                    help="Cole o texto a ser analisado",
                                                    height=content_height,
                                                    placeholder="Cole o texto para identificar riscos...")
                    
                    audience = st.text_input("Público-alvo", 
                                        placeholder="Ex: Cliente executivo com conhecimento técnico limitado")
                    
                    stakes = st.select_slider("Importância da comunicação", 
                                            options=["Baixa", "Média", "Alta", "Crítica"],
                                            value="Média")
                    
                    prompt = f"""
                    Analise o seguinte {subtype} quanto a riscos de comunicação:
                    
                    Contexto: {context}
                    Público: {audience}
                    Importância: {stakes}
                    
                    Conteúdo:
                    ---
                    {content_to_analyze}
                    ---
                    
                    Identifique:
                    1. Ambiguidades e informações confusas
                    2. Possíveis mal-entendidos baseados no público
                    3. Problemas de tom ou linguagem inapropriada
                    4. Informações sensíveis ou potencialmente problemáticas
                    5. Sugestões específicas de melhoria
                    
                    Forneça uma avaliação geral e uma versão revisada completa do texto.
                    """
                
                elif current_feature == "Consultor PMBOK 7":
                    # Layout simplificado para consultas PMBOK
                    pmbok_topic = subtype
                    field_height = 100 if is_mobile else 150
                    
                    project_context = st.text_area("Contexto do projeto", 
                                            height=field_height,
                                            placeholder="Ex: Projeto de software com metodologia híbrida...")
                    
                    specific_question = st.text_area("Sua dúvida específica", 
                                                height=field_height,
                                                placeholder="Ex: Como aplicar os princípios de valor do PMBOK 7...")
                    
                    experience_level = st.select_slider("Seu nível de experiência", 
                                                    options=["Iniciante", "Intermediário", "Avançado"],
                                                    value="Intermediário")
                    
                    # Campo opcional com visualização condicional
                    show_org_context = st.checkbox("Adicionar contexto organizacional", value=False)
                    organization_context = ""
                    if show_org_context:
                        organization_context = st.text_input("Contexto organizacional",
                                                        placeholder="Ex: Empresa de médio porte do setor financeiro")
                    
                    base_prompt = f"""
                    Forneça orientação sobre "{pmbok_topic}" do PMBOK 7:
                    
                    Contexto: {project_context}
                    Dúvida: {specific_question}
                    Nível de experiência: {experience_level}
                    {"Contexto organizacional: " + organization_context if organization_context else ""}
                    
                    Sua resposta deve:
                    1. Explicar conceitos relevantes do PMBOK 7
                    2. Fornecer orientações práticas para o contexto específico
                    3. Incluir exemplos concretos de aplicação
                    4. Destacar boas práticas considerando o nível de experiência ({experience_level})
                    
                    Formate a resposta de maneira estruturada e objetiva.
                    """
                    
                    # Enriquecemos o prompt com informações relevantes do PMBOK 7
                    prompt = enrich_pmbok_prompt(base_prompt, pmbok_topic)
                
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Botão de geração mais atraente
                st.markdown("""
                <style>
                    /* Estilo Apple para o botão de submit */
                    button[type="submit"] {
                        background-color: var(--primary);
                        color: white;
                        border: none;
                        border-radius: 8px;
                        padding: 12px 24px;
                        font-weight: 500;
                        font-size: 15px;
                        width: 100%;
                        cursor: pointer;
                        transition: all 0.2s;
                        margin-top: 8px;
                    }
                    
                    button[type="submit"]:hover {
                        background-color: #0062CC;
                        box-shadow: 0 4px 10px rgba(0,0,0,0.1);
                    }
                </style>
                """, unsafe_allow_html=True)
                
                submit_col1, submit_col2, submit_col3 = st.columns([1, 2, 1])
                with submit_col2:
                    submit_button = st.form_submit_button("GERAR CONTEÚDO")

            # Processamento após o envio do formulário
            if submit_button:
                # Verificar API e limites novamente
                if not st.session_state.api_key_configured:
                    st.error("API não configurada. Por favor, contate o administrador.")
                elif st.session_state.token_count >= TOKEN_LIMIT:
                    st.error(f"Limite de tokens atingido.")
                elif st.session_state.request_count >= REQUEST_LIMIT:
                    st.error(f"Limite de requisições atingido.")
                else:
                    # Feedback visual de carregamento melhorado
                    with st.status("Gerando conteúdo personalizado...", expanded=True) as status:
                        st.write("Processando sua solicitação...")
                        
                        # Gerar conteúdo com modelo apropriado
                        model = "gpt-3.5-turbo"  # Modelo padrão
                        temperature = 0.7        # Criatividade padrão
                        
                        # Ajustar modelo e temperatura baseado na funcionalidade
                        if current_feature in ["Detector de Riscos", "Consultor PMBOK 7"]:
                            temperature = 0.3  # Menos criatividade para análises
                        
                        generated_content = generate_content(prompt, model=model, temperature=temperature)
                        st.session_state.generated_content = generated_content
                        
                        # Atualizar status
                        status.update(label="Conteúdo gerado com sucesso!", state="complete", expanded=False)
                    
                    # Exibir resultado em container estilizado
                    st.markdown('<div class="result-container">', unsafe_allow_html=True)
                    st.markdown("### Resultado")
                    
                    # Container com estilo aprimorado para o resultado
                    st.markdown(f"""
                    <div style="background-color: white; border-radius: 12px; padding: 20px; 
                                border: 1px solid var(--border); margin-bottom: 20px;">
                        <div style="white-space: pre-wrap; font-size: 15px; line-height: 1.5;">
                            {generated_content.replace('\n', '<br>')}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Botões de ação para o resultado
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        # Voltar para editar
                        if st.button("✏️ Editar solicitação", key="edit_request", type="secondary"):
                            # Manter o current_feature, apenas limpar o resultado
                            st.session_state.generated_content = ""
                            st.rerun()
                    
                    with col2:
                        # Download como texto
                        st.download_button(
                            label="📄 Baixar como TXT",
                            data=generated_content,
                            file_name=f"nexus_{subtype.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            key="download_txt"
                        )
                    
                    with col3:
                        # Download como DOCX
                        docx_buffer = export_as_docx(generated_content, filename=f"NEXUS - {subtype}")
                        st.download_button(
                            label="📝 Baixar como DOCX",
                            data=docx_buffer,
                            file_name=f"nexus_{subtype.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="download_docx"
                        )

# Iniciar a aplicação
if __name__ == "__main__":
    main()
