import streamlit as st
import requests
import json
import pandas as pd
from datetime import datetime
import plotly.express as px
import os
import re
import docx
from io import BytesIO
import time

# ================= CONFIGURATION =================

# Configuração da página
st.set_page_config(
    page_title="NEXUS - Assistente de Comunicação de Projetos",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Definindo paleta de cores
nexus_colors = {
    "purple": "#6247AA",
    "orange": "#FF6D2A",
    "teal": "#00C1D5",
    "dark_purple": "#231A41",
    "background": "#F5F5F7",
    "text_primary": "#1D1D1F",
    "text_secondary": "#86868B"
}

st.markdown("""
<style>
    /* Ocultar completamente a sidebar */
    [data-testid="collapsedControl"] {
        display: none;
    }
    section[data-testid="stSidebar"] {
        display: none;
    }
    /* Ajustar a área principal para ocupar toda a largura */
    .main .block-container {
        padding-left: 2rem;
        padding-right: 2rem;
        max-width: 1200px;
    }
</style>
""", unsafe_allow_html=True)

# Base de conhecimento sobre PMBOK 7
PMBOK7_KNOWLEDGE_BASE = {
    "princípios": {
        "1. Ser um administrador diligente, respeitoso e cuidadoso": {
            "definição": "Este princípio fundamenta o gerenciamento de projetos na ética e integridade como alicerces inegociáveis. A administração vai além da simples gestão de recursos - representa uma responsabilidade fiduciária e moral com todas as partes envolvidas e impactadas pelo projeto.",
            "dimensões": [
                "Stewardship Organizacional: Uso responsável de recursos da organização, garantindo retorno de investimento adequado e alinhamento com objetivos estratégicos.",
                "Stewardship Humano: Cuidado com o bem-estar, desenvolvimento profissional e equilíbrio trabalho-vida da equipe.",
                "Stewardship Social: Consideração do impacto social do projeto, incluindo efeitos em comunidades locais e stakeholders indiretos.",
                "Stewardship Ambiental: Minimização de impactos ecológicos negativos e consideração da sustentabilidade a longo prazo.",
                "Stewardship Ético: Manutenção dos mais altos padrões de honestidade, transparência e integridade em todas as decisões."
            ],
            "implementação": [
                "Implementação de sistemas de governança transparentes e responsáveis",
                "Criação de mecanismos para receber e incorporar feedback de todas as partes interessadas",
                "Desenvolvimento de estruturas de prestação de contas em múltiplos níveis",
                "Estabelecimento de processos éticos de tomada de decisão que consideram impactos multidimensionais",
                "Documentação clara de compromissos e responsabilidades"
            ],
            "indicadores": [
                "Satisfação equilibrada entre múltiplos stakeholders",
                "Ausência de violações éticas ou de conformidade",
                "Sustentabilidade dos resultados do projeto após sua conclusão",
                "Feedback positivo sobre a conduta da equipe de gerenciamento do projeto",
                "Uso eficiente e responsável de recursos"
            ],
            "desafios": [
                {"desafio": "Pressões para cortar caminho em questões éticas devido a restrições de tempo ou orçamento", 
                 "superação": "Estabelecer 'linhas vermelhas' não-negociáveis e mecanismos de escalação para questões éticas"},
                {"desafio": "Conflitos entre diferentes responsabilidades (ex: lucro vs. impacto ambiental)", 
                 "superação": "Desenvolver frameworks de decisão que equilibrem múltiplos objetivos e criar transparência sobre trade-offs"},
                {"desafio": "Determinar o nível apropriado de envolvimento com stakeholders secundários", 
                 "superação": "Realizar mapeamento abrangente de impacto para identificar todos os grupos afetados e estabelecer protocolos de engajamento proporcionais"}
            ]
        },
        
        "2. Criar um ambiente colaborativo da equipe do projeto": {
            "definição": "Este princípio reconhece que o sucesso do projeto está intrinsecamente ligado à qualidade das interações da equipe e ao ambiente psicológico em que operam. Vai além da simples colaboração, focando na criação de condições ideais para alto desempenho coletivo, inovação e resolução de problemas.",
            "dimensões": [
                "Segurança Psicológica: Criação de um ambiente onde os membros da equipe sentem-se seguros para assumir riscos interpessoais, expressar preocupações e compartilhar ideias sem medo de retaliação ou humilhação.",
                "Diversidade Cognitiva: Valorização e aproveitamento de diferentes perspectivas, experiências e estilos de pensamento para enriquecer a tomada de decisão.",
                "Coesão Estruturada: Estabelecimento de processos que permitem autonomia individual enquanto mantêm alinhamento com objetivos comuns.",
                "Comunicação Multidirecional: Promoção de fluxos de informação abertos em todas as direções (vertical e horizontal).",
                "Inteligência Coletiva: Desenvolvimento de sistemas que maximizem a inteligência combinada da equipe acima da soma das capacidades individuais."
            ],
            "implementação": [
                "Estabelecimento de normas explícitas que promovam vulnerabilidade produtiva e abertura",
                "Implementação de estruturas de reunião que garantem que todas as vozes sejam ouvidas",
                "Criação de mecanismos de feedback contínuo e aprendizado em equipe",
                "Desenvolvimento de protocolos para resolução construtiva de conflitos",
                "Desenho de espaços físicos e virtuais que facilitam colaboração formal e informal",
                "Estabelecimento de rituais de equipe que fortalecem identidade e propósito compartilhado"
            ],
            "indicadores": [
                "Participação equilibrada entre membros da equipe em discussões e tomada de decisão",
                "Alta qualidade de debate construtivo sem conflitos disfuncionais",
                "Rápida resolução de obstáculos através de colaboração",
                "Inovação incremental e contínua nos processos e soluções",
                "Baixa rotatividade e alto engajamento dos membros da equipe",
                "Resiliência da equipe frente a desafios e mudanças"
            ],
            "desafios": [
                {"desafio": "Equipes distribuídas globalmente ou trabalhando remotamente", 
                 "superação": "Implementar práticas específicas para colaboração virtual, incluindo ferramentas síncronas e assíncronas, check-ins estruturados e tempo dedicado para construção de relacionamentos"},
                {"desafio": "Membros de equipe com diferentes culturas organizacionais ou nacionais", 
                 "superação": "Criar acordos de equipe explícitos que acomodem diferenças culturais e estabelecer glossários compartilhados para alinhar entendimento de termos-chave"},
                {"desafio": "Desequilíbrios de poder que inibem contribuições genuínas", 
                 "superação": "Implementar técnicas como 'round robin' para coleta de ideias, sessões anônimas de brainstorming, e alternância de papéis de facilitação"}
            ]
        },
        
        "3. Envolver efetivamente as partes interessadas": {
            "definição": "Este princípio transcende a simples comunicação com stakeholders, representando uma abordagem sofisticada e estratégica para cultivar relacionamentos produtivos com todos os indivíduos e grupos que podem influenciar ou ser influenciados pelo projeto. O engajamento efetivo reconhece que stakeholders são parceiros ativos no sucesso do projeto, não apenas receptores de informação.",
            "dimensões": [
                "Mapeamento Multidimensional: Identificação não apenas da posição formal dos stakeholders, mas de suas redes de influência, motivações subjacentes, preocupações não expressas e potencial para impacto no projeto.",
                "Engajamento Personalizado: Adaptação da abordagem de comunicação e envolvimento às necessidades específicas, preferências e contexto cultural de cada stakeholder ou grupo.",
                "Engajamento Evolutivo: Reconhecimento de que as posições, interesses e influência dos stakeholders mudam ao longo do tempo, exigindo reavaliação contínua.",
                "Construção de Coalizões: Desenvolvimento estratégico de alianças e redes de apoio entre grupos de stakeholders.",
                "Gestão de Expectativas: Alinhamento proativo de percepções e expectativas com realidades do projeto."
            ],
            "implementação": [
                "Criação de matrizes de stakeholders multidimensionais que vão além do poder/interesse para incluir fatores como atitude, influência, conhecimento e necessidades",
                "Desenvolvimento de planos de engajamento personalizados para stakeholders-chave",
                "Estabelecimento de múltiplos canais de comunicação e feedback",
                "Criação de mecanismos para detecção precoce de mudanças na posição ou percepção dos stakeholders",
                "Implementação de sistemas de rastreamento de compromissos com stakeholders",
                "Documentação e gerenciamento ativo de expectativas dos stakeholders"
            ],
            "indicadores": [
                "Ausência de 'surpresas' negativas de stakeholders durante o projeto",
                "Feedback positivo sobre transparência e comunicação",
                "Participação ativa dos stakeholders em momentos críticos de decisão",
                "Stakeholders atuando como defensores e facilitadores do projeto",
                "Resolução eficiente de conflitos entre stakeholders com interesses divergentes",
                "Continuidade de relações positivas mesmo após conclusão do projeto"
            ],
            "desafios": [
                {"desafio": "Stakeholders com expectativas irrealistas ou conflitantes", 
                 "superação": "Implementar processos de alinhamento de expectativas baseados em dados, incluindo benchmarks externos e análises de viabilidade transparentes"},
                {"desafio": "Stakeholders ocultos ou que emergem tardiamente", 
                 "superação": "Realizar análises de redes organizacionais expandidas periodicamente e manter canais abertos para identificação de novos stakeholders"},
                {"desafio": "Fadiga de engajamento em projetos longos", 
                 "superação": "Variar formatos e frequência de comunicação, focar em valor incremental e desenvolver narrativas envolventes sobre o progresso"}
            ]
        },
        
        "4. Focar no valor": {
            "definição": "Este princípio reorienta o gerenciamento de projetos de um foco em entrega de outputs (produtos, serviços) para um foco em outcomes e benefícios. O valor transcende métricas financeiras simples para incluir todas as formas de impacto positivo para a organização e stakeholders, reconhecendo que o valor é multidimensional, contextual e evolutivo.",
            "dimensões": [
                "Valor Financeiro: ROI, fluxo de caixa, aumento de receita, redução de custos, valor presente líquido e outras métricas financeiras.",
                "Valor Estratégico: Contribuição para objetivos organizacionais de longo prazo, posicionamento competitivo e capacidades organizacionais.",
                "Valor Operacional: Melhorias em eficiência, qualidade, velocidade e flexibilidade nos processos organizacionais.",
                "Valor para o Cliente: Aumento de satisfação, lealdade, experiência e percepção da marca.",
                "Valor Organizacional: Desenvolvimento de conhecimento, habilidades, cultura e capacidade de inovação.",
                "Valor Social: Impacto positivo em comunidades, meio ambiente e sociedade em geral."
            ],
            "implementação": [
                "Desenvolvimento de frameworks personalizados para definir valor em múltiplas dimensões",
                "Criação de mecanismos para rastrear a entrega incremental de valor ao longo do projeto",
                "Implementação de processos de tomada de decisão explicitamente baseados em valor",
                "Estabelecimento de critérios de priorização que maximizam valor total",
                "Criação de visualizações de 'cadeia de valor' que conectam atividades específicas a outcomes",
                "Implementação de revisões periódicas de realização de benefícios"
            ],
            "indicadores": [
                "Clareza compartilhada sobre a definição de valor para o projeto específico",
                "Capacidade de articular como cada componente do projeto contribui para valor",
                "Decisões de escopo consistentemente baseadas em análises de valor",
                "Mecanismos funcionais para medir e rastrear a realização de benefícios",
                "Percepção dos stakeholders de que o projeto entrega valor substancial",
                "Realização mensurável de benefícios planejados após a conclusão do projeto"
            ],
            "desafios": [
                {"desafio": "Dificuldade em quantificar formas intangíveis de valor", 
                 "superação": "Desenvolver proxies mensuráveis e indicadores indiretos; utilizar técnicas de valoração baseadas em preferência e disposição a pagar"},
                {"desafio": "Pressão para focar em métricas de curto prazo vs. valor de longo prazo", 
                 "superação": "Criar scorecards balanceados que incluem indicadores de lead e lag; estabelecer métricas intermediárias que preveem valor futuro"},
                {"desafio": "Divergências entre stakeholders sobre o que constitui valor", 
                 "superação": "Facilitar workshops estruturados de alinhamento de valor; desenvolver modelos multi-critério que acomodam diferentes perspectivas"}
            ]
        },

        "5. Reconhecer, avaliar e responder às interações do sistema": {
            "definição": "Este princípio fundamenta-se no entendimento de que projetos são sistemas complexos adaptativos, não processos lineares. Reconhece que componentes, pessoas, e organizações interagem de formas não-lineares, criando comportamentos emergentes, loops de feedback, e efeitos secundários que devem ser compreendidos e gerenciados holisticamente.",
            "dimensões": [
                "Mapeamento de Interdependências: Identificação das relações e dependências entre componentes do projeto, tanto técnicos quanto organizacionais.",
                "Análise de Feedback: Compreensão dos loops de feedback (reforçadores e equilibradores) que influenciam o comportamento do sistema.",
                "Identificação de Alavancas: Reconhecimento de pontos de alta influência no sistema onde pequenas intervenções podem produzir grandes resultados.",
                "Previsão de Efeitos Emergentes: Antecipação de propriedades e comportamentos que emergem das interações do sistema, não visíveis no nível individual.",
                "Gestão de Fronteiras: Definição consciente dos limites do sistema do projeto e suas interfaces com outros sistemas."
            ],
            "implementação": [
                "Criação de mapas de sistemas visuais que mostram relações, interdependências e loops de feedback",
                "Implementação de análises de impacto sistêmico antes de mudanças significativas",
                "Estabelecimento de processos para identificar e monitorar efeitos secundários e consequências não intencionais",
                "Desenvolvimento de simulações e cenários que exploram comportamentos emergentes",
                "Implementação de reuniões periódicas de reflexão sistêmica com múltiplas perspectivas",
                "Criação de indicadores que capturam comportamentos no nível do sistema, não apenas componentes"
            ],
            "indicadores": [
                "Capacidade de prever impactos secundários de mudanças antes que ocorram",
                "Ausência de falhas causadas por interdependências não reconhecidas",
                "Maior estabilidade e previsibilidade do projeto apesar da complexidade",
                "Intervenções eficientes que produzem resultados amplificados",
                "Integração harmoniosa do projeto com sistemas organizacionais existentes",
                "Adaptações bem-sucedidas a mudanças nas condições do sistema"
            ],
            "desafios": [
                {"desafio": "Tendência a simplificar excessivamente sistemas complexos", 
                 "superação": "Utilizar ferramentas de visualização e modelagem que revelam complexidade; incorporar múltiplas perspectivas através de equipes multidisciplinares"},
                {"desafio": "Dificuldade em comunicar pensamento sistêmico a stakeholders", 
                 "superação": "Desenvolver metáforas e visualizações acessíveis; demonstrar através de casos reais como interações sistêmicas afetaram resultados"},
                {"desafio": "Paralisia analítica devido à percepção de complexidade excessiva", 
                 "superação": "Focar em 'complexidade gerenciável' através de decomposição de sistemas; aplicar princípio de Pareto para identificar as interações mais críticas"}
            ]
        },
        
        "6. Demonstrar comportamentos de liderança": {
            "definição": "Este princípio transcende a visão tradicional de liderança como função hierárquica, reconhecendo-a como um conjunto de comportamentos que podem e devem ser demonstrados por todos os envolvidos no projeto, independentemente de cargo ou autoridade formal. Enfatiza liderança como influência, visão, facilitação e habilitação, não controle.",
            "dimensões": [
                "Liderança Adaptativa: Capacidade de ajustar estilo e abordagem baseado no contexto, necessidades da equipe e fase do projeto.",
                "Liderança Distribuída: Cultivo de capacidades de liderança em toda a equipe, criando um sistema onde a liderança emerge naturalmente de quem tem conhecimento ou insight relevante.",
                "Liderança Servil: Foco em remover obstáculos, fornecer recursos e criar condições para que outros tenham sucesso.",
                "Liderança Transformacional: Inspiração de níveis mais altos de motivação e desempenho através de visão compartilhada e propósito claro.",
                "Liderança Autêntica: Demonstração de autoconsciência, transparência e integridade que constrói confiança e credibilidade."
            ],
            "implementação": [
                "Estabelecimento de uma visão clara e envolvente do sucesso do projeto",
                "Desenvolvimento intencional de líderes em múltiplos níveis da equipe do projeto",
                "Criação de oportunidades para membros da equipe exercitarem liderança em áreas específicas",
                "Implementação de processos de decisão que balanceiam autonomia e alinhamento",
                "Modelagem consistente de comportamentos desejados (liderar pelo exemplo)",
                "Estabelecimento de mecanismos para reflexão e desenvolvimento de capacidades de liderança"
            ],
            "indicadores": [
                "Emergência natural de liderança de diferentes membros da equipe em momentos apropriados",
                "Decisões eficazes tomadas no nível apropriado sem escalações desnecessárias",
                "Alta confiança e engajamento entre membros da equipe",
                "Capacidade da equipe de navegar desafios e incertezas sem direção constante",
                "Melhoria contínua nas capacidades individuais e coletivas de liderança",
                "Percepção dos stakeholders de liderança coesa e eficaz no projeto"
            ],
            "desafios": [
                {"desafio": "Tensão entre liderança compartilhada e accountability clara", 
                 "superação": "Implementar frameworks de 'responsabilidade sem culpa' com papéis explícitos mas flexíveis; distinguir claramente entre accountability por decisões e autoridade para tomar decisões"},
                {"desafio": "Resistência a liderança distribuída em culturas hierárquicas", 
                 "superação": "Introduzir gradualmente através de áreas designadas de autonomia; demonstrar sucessos iniciais; obter apoio explícito de líderes formais"},
                {"desafio": "Desenvolvimento insuficiente de capacidades de liderança na equipe", 
                 "superação": "Criar programa intencional de desenvolvimento com coaching, oportunidades estruturadas, feedback e reflexão"}
            ]
        },
        
        "7. Adaptar com base no contexto": {
            "definição": "Este princípio reconhece que não existe abordagem universal para gerenciamento de projetos. Em vez de aplicar fórmulas padronizadas, exige uma avaliação sofisticada das características específicas do ambiente, projeto, organização e stakeholders para criar uma abordagem personalizada e evolutiva que maximiza a probabilidade de sucesso.",
            "dimensões": [
                "Adaptação Metodológica: Seleção e customização de abordagens, métodos e frameworks baseados nas características específicas do projeto.",
                "Sensibilidade Cultural: Ajuste de práticas para alinhar com culturas organizacionais, regionais e profissionais relevantes.",
                "Calibração de Formalidade: Determinação do nível apropriado de formalidade, documentação e controle baseado nas necessidades reais.",
                "Evolução Temporal: Modificação da abordagem conforme o projeto avança e as condições mudam.",
                "Personalização de Interação: Adaptação de estilos de comunicação e engajamento às preferências e necessidades dos stakeholders."
            ],
            "implementação": [
                "Condução de análises de contexto detalhadas no início do projeto e periodicamente",
                "Desenvolvimento de frameworks de decisão para selecionar abordagens apropriadas",
                "Criação de 'playbooks' customizados que documentam a abordagem adaptada",
                "Implementação de revisões periódicas da eficácia da abordagem atual",
                "Construção de flexibilidade intencional em processos e artefatos",
                "Estabelecimento de mecanismos para detectar mudanças contextuais que podem exigir adaptação"
            ],
            "indicadores": [
                "Abordagem claramente alinhada com o contexto específico do projeto",
                "Capacidade de articular o raciocínio por trás das escolhas metodológicas",
                "Ajuste fluido de práticas quando necessário sem disrupção",
                "Ausência de imposição de processos desnecessariamente pesados ou inadequados",
                "Feedback positivo dos stakeholders sobre a adequação da abordagem",
                "Balanço apropriado entre consistência e flexibilidade"
            ],
            "desafios": [
                {"desafio": "Pressão organizacional para seguir metodologias padronizadas", 
                 "superação": "Desenvolver frameworks de 'conformidade com variação' que demonstram alinhamento com princípios organizacionais enquanto permitem adaptação; documentar e comunicar claramente o raciocínio para adaptações específicas"},
                {"desafio": "Dificuldade em determinar o nível apropriado de adaptação", 
                 "superação": "Implementar abordagem experimental com feedback rápido; utilizar decisões baseadas em dados para avaliar eficácia de adaptações"},
                {"desafio": "Equilíbrio entre estabilidade necessária e adaptação contínua", 
                 "superação": "Definir explicitamente 'zonas de estabilidade' vs. 'zonas de experimentação'; estabelecer cadência para revisões de abordagem para evitar mudanças constantes"}
            ]
        },
        
        "8. Incorporar qualidade nos processos e resultados": {
            "definição": "Este princípio transcende a visão tradicional de qualidade como conformidade com especificações, adotando uma abordagem holística que integra qualidade no DNA do projeto. Reconhece que qualidade é multidimensional, abrangendo não apenas o produto final, mas também processos, experiências dos stakeholders e impactos de longo prazo.",
            "dimensões": [
                "Qualidade Incorporada: Integração de controles de qualidade dentro dos processos de trabalho, não como atividade separada.",
                "Qualidade Expandida: Consideração de múltiplas dimensões de qualidade, incluindo técnica, experiencial, estética e durabilidade.",
                "Qualidade Preventiva: Foco em prevenir defeitos e problemas na origem, não em detectá-los posteriormente.",
                "Qualidade Cultural: Desenvolvimento de uma mentalidade coletiva onde qualidade é responsabilidade de todos.",
                "Qualidade Evolutiva: Refinamento contínuo da definição e padrões de qualidade baseado em feedback e aprendizado."
            ],
            "implementação": [
                "Definição colaborativa e explícita de qualidade para cada entrega significativa",
                "Implementação de práticas como 'shift left testing' e revisão por pares integrada ao fluxo de trabalho",
                "Estabelecimento de loops de feedback curtos para detecção precoce de problemas",
                "Criação de mecanismos para capturar e aplicar aprendizados relacionados à qualidade",
                "Desenvolvimento de 'padrões de excelência' específicos para o projeto",
                "Implementação de revisões periódicas de processos para melhoria contínua"
            ],
            "indicadores": [
                "Redução significativa em retrabalho e correções tardias",
                "Definições claras e compartilhadas de qualidade para todas as entregas",
                "Integração visível de práticas de qualidade no trabalho diário",
                "Orgulho demonstrável da equipe na qualidade de seu trabalho",
                "Feedback positivo dos stakeholders sobre a qualidade das entregas",
                "Equilíbrio apropriado entre qualidade, tempo e custo"
            ],
            "desafios": [
                {"desafio": "Pressão de cronograma levando a compromissos de qualidade", 
                 "superação": "Implementar 'definições de pronto' não-negociáveis; tornar visíveis os custos reais de qualidade insuficiente; criar mecanismos para decisões explícitas sobre trade-offs"},
                {"desafio": "Definições subjetivas ou ambíguas de qualidade", 
                 "superação": "Facilitar workshops dedicados à definição operacional de qualidade; desenvolver critérios mensuráveis e exemplos concretos para cada aspecto de qualidade"},
                {"desafio": "Silos de responsabilidade de qualidade", 
                 "superação": "Implementar práticas de propriedade coletiva e responsabilidade compartilhada; integrar especialistas de qualidade nas equipes de desenvolvimento em vez de mantê-los separados"}
            ]
        },

       "9. Navegar na complexidade": {
            "definição": "Este princípio reconhece que projetos modernos frequentemente operam em ambientes caracterizados por alta complexidade - múltiplas variáveis interdependentes, causalidade não-linear, ambiguidade e emergência. Em vez de tentar eliminar a complexidade (frequentemente impossível) ou ignorá-la (perigoso), este princípio enfoca o desenvolvimento de capacidades para funcionar efetivamente dentro dela.",
            "dimensões": [
                "Complexidade Técnica: Gerenciamento de sistemas com múltiplos componentes interdependentes e interfaces.",
                "Complexidade Social: Navegação em ambientes com múltiplos stakeholders, interesses divergentes e relações em evolução.",
                "Complexidade Temporal: Adaptação a condições que mudam ao longo do tempo, frequentemente de formas imprevisíveis.",
                "Complexidade Cognitiva: Gerenciamento da capacidade da equipe e stakeholders para compreender e processar informações complexas.",
                "Complexidade Contextual: Funcionamento em ecossistemas com múltiplos fatores externos influenciando o projeto."
            ],
            "implementação": [
                "Aplicação de técnicas de visualização para tornar a complexidade visível e compreensível",
                "Implementação de abordagens exploratórias e experimentais para aprendizado em ambientes complexos",
                "Desenvolvimento de modelos mentais compartilhados que capturam elementos essenciais da complexidade",
                "Criação de sistemas de detecção precoce para identificar padrões emergentes",
                "Estabelecimento de práticas regulares de sentido coletivo (collective sensemaking)",
                "Implementação de estratégias de simplificação onde apropriado sem redução excessiva"
            ],
            "indicadores": [
                "Capacidade da equipe para articular e visualizar a complexidade relevante",
                "Abordagens apropriadamente matizadas em vez de soluções simplistas",
                "Detecção precoce de padrões emergentes e consequências não intencionais",
                "Capacidade demonstrada de pivotar quando padrões complexos mudam",
                "Equilíbrio efetivo entre ação decisiva e exploração cautelosa",
                "Ausência de surpresas catastróficas devido a dinâmicas complexas não reconhecidas"
            ],
            "desafios": [
                {"desafio": "Tendência organizacional para soluções simples e previsibilidade linear", 
                 "superação": "Educar stakeholders sobre a natureza da complexidade usando exemplos concretos; desenvolver métricas e narrativas que reconhecem a natureza não-linear do progresso"},
                {"desafio": "Sobrecarga de informação e paralisia analítica", 
                 "superação": "Implementar técnicas de simplificação estratégica; focar na 'complexidade essencial' vs. 'complexidade acidental'; utilizar visualização e abstrações apropriadas"},
                {"desafio": "Dificuldade em planejar em ambientes altamente complexos", 
                 "superação": "Adotar horizontes de planejamento mais curtos; implementar abordagens adaptativas com feedback frequente; utilizar planejamento baseado em cenários"}
            ]
        },
        
        "10. Otimizar respostas a riscos": {
            "definição": "Este princípio expande a gestão de riscos tradicional para uma abordagem sofisticada que reconhece a natureza multifacetada e dinâmica do risco. Em vez de focar apenas em identificação e mitigação, enfatiza a otimização holística da postura de risco do projeto, equilibrando ameaças e oportunidades, e calibrando precisamente respostas baseadas em múltiplos fatores contextuais.",
            "dimensões": [
                "Gestão Equilibrada: Consideração tanto de ameaças (riscos negativos) quanto oportunidades (riscos positivos).",
                "Abordagem Multiestratégica: Aplicação de um espectro de estratégias desde evitar até aumentar, dependendo da natureza do risco.",
                "Proporcionalidade Dinâmica: Calibração da intensidade da resposta baseada não apenas na severidade do risco, mas também em eficiência de recursos e custo-benefício.",
                "Gestão de Incerteza: Desenvolvimento de estratégias específicas para diferentes tipos de incerteza (conhecida, desconhecida, incognoscível).",
                "Resiliência Sistêmica: Criação de capacidade para absorver, adaptar e se transformar frente a eventos inesperados."
            ],
            "implementação": [
                "Desenvolvimento de frameworks de risco multidimensionais que incluem dimensões além de probabilidade e impacto",
                "Implementação de processos contínuos de identificação e reavaliação de riscos",
                "Criação de planos de contingência e gatilhos claros para sua ativação",
                "Estabelecimento de reservas estratégicas de tempo, orçamento e outros recursos",
                "Implementação de simulações e análises de cenário para testar respostas potenciais",
                "Criação de mecanismos para capturar e aplicar lições aprendidas sobre riscos"
            ],
            "indicadores": [
                "Registro de riscos dinâmico e ativamente gerenciado, não documento estático",
                "Discussões de risco integradas nas conversas regulares do projeto",
                "Equilíbrio entre foco em ameaças e oportunidades",
                "Respostas a riscos proporcionais e eficientes em uso de recursos",
                "Capacidade demonstrável de pivotar rapidamente quando riscos se materializam",
                "Aproveitamento sistemático de oportunidades emergentes"
            ],
            "desafios": [
                {"desafio": "Foco excessivo em riscos negativos óbvios em detrimento de oportunidades ou riscos sistêmicos", 
                 "superação": "Implementar frameworks que exigem igual consideração de ameaças e oportunidades; realizar sessões dedicadas especificamente à identificação de oportunidades"},
                {"desafio": "Abordagens burocráticas à gestão de riscos sem integração real com tomada de decisões", 
                 "superação": "Simplificar processos de risco para focar em valor; integrar considerações de risco diretamente em processos de decisão existentes; tornar os riscos visíveis e tangíveis"},
                {"desafio": "Dificuldade em priorizar entre múltiplos riscos com diferentes características", 
                 "superação": "Desenvolver frameworks de priorização multidimensionais; utilizar técnicas como análise de decisão multicritério; focar em clusters de riscos relacionados em vez de riscos individuais"}
            ]
        },
        
        "11. Abraçar adaptabilidade e resiliência": {
            "definição": "Este princípio representa uma mudança fundamental de paradigma - de buscar estabilidade e previsibilidade para cultivar a capacidade de prosperar em ambientes de mudança contínua. Vai além da simples flexibilidade ou capacidade de reagir, enfocando o desenvolvimento de sistemas, equipes e abordagens intrinsecamente adaptáveis e resilientes que não apenas sobrevivem, mas evoluem positivamente frente a mudanças e desafios.",
            "dimensões": [
                "Adaptabilidade Proativa: Capacidade de antecipar mudanças potenciais e posicionar-se para capitalizar sobre elas, não apenas reagir.",
                "Resiliência Estrutural: Desenho de processos, planos e sistemas com flexibilidade inerente e margem para absorver impactos.",
                "Antifragilidade: Desenvolvimento da capacidade de fortalecer-se através de perturbações e desafios, não apenas resistir a eles.",
                "Agilidade Responsiva: Capacidade de pivotar rapidamente e eficientemente em resposta a novas informações ou circunstâncias.",
                "Aprendizado Evolutivo: Habilidade de transformar experiências e desafios em capacidades melhoradas para o futuro."
            ],
            "implementação": [
                "Desenho intencional de flexibilidade em planos, processos e estruturas organizacionais",
                "Implementação de simulações e exercícios de cenários 'e se?' para desenvolver capacidade adaptativa",
                "Criação de mecanismos de feedback rápido e múltiplos pontos de medição e avaliação",
                "Estabelecimento de processos formais para capturar e integrar aprendizados",
                "Desenvolvimento deliberado de redundâncias estratégicas em áreas críticas",
                "Criação de 'playbooks' para diferentes cenários de adaptação"
            ],
            "indicadores": [
                "Resposta eficaz e ágil a mudanças imprevistas",
                "Baixo impacto de perturbações externas nos resultados do projeto",
                "Capacidade de capitalizar sobre oportunidades emergentes",
                "Melhoria contínua em capacidades adaptativas ao longo do projeto",
                "Equilíbrio efetivo entre estabilidade necessária e flexibilidade",
                "Percepção por stakeholders de confiabilidade mesmo em ambientes turbulentos"
            ],
            "desafios": [
                {"desafio": "Resistência organizacional a abordagens que reconhecem explicitamente incerteza", 
                 "superação": "Enquadrar adaptabilidade como força estratégica, não fraqueza; demonstrar como adaptabilidade melhora previsibilidade de resultados finais mesmo quando caminhos específicos mudam"},
                {"desafio": "Confusão entre adaptabilidade e falta de direção ou disciplina", 
                 "superação": "Estabelecer 'estruturas estáveis para flexibilidade' - elementos constantes que não mudam, dentro dos quais a adaptação ocorre; comunicar claramente a diferença entre pivôs estratégicos e desvios arbitrários"},
                {"desafio": "Fadiga de mudança em ambientes de alta volatilidade", 
                 "superação": "Implementar práticas de sustentabilidade e resiliência para equipes; criar 'zonas de estabilidade' mesmo em projetos altamente adaptativos; celebrar e reconhecer adaptação bem-sucedida"}
            ]
        },
        
        "12. Permitir mudança para alcançar o estado futuro previsto": {
            "definição": "Este princípio reconhece que projetos são fundamentalmente veículos de mudança e transformação - não apenas produtores de entregáveis. Foca na transição bem-sucedida do estado atual para o estado futuro desejado, considerando aspectos técnicos, comportamentais, culturais e organizacionais necessários para que a mudança seja adotada, incorporada e sustentada.",
            "dimensões": [
                "Arquitetura da Mudança: Desenho holístico da transição, incluindo aspectos técnicos, comportamentais, culturais e sistêmicos.",
                "Adoção Sustentada: Foco não apenas na implementação inicial, mas na incorporação permanente da mudança nos comportamentos e sistemas.",
                "Gerenciamento da Transição: Facilitação da jornada do estado atual ao estado futuro, incluindo estados intermediários.",
                "Capacitação para Mudança: Desenvolvimento das habilidades, motivação e oportunidades necessárias para que pessoas adotem novas formas de trabalhar.",
                "Alinhamento Ecossistêmico: Garantia que os sistemas, estruturas e incentivos organizacionais suportam a mudança desejada."
            ],
            "implementação": [
                "Desenvolvimento de visões claras e envolventes do estado futuro",
                "Criação de modelos que mostram explicitamente a transição entre estado atual e futuro",
                "Implementação de análises de impacto da mudança em múltiplos níveis",
                "Estabelecimento de redes de agentes de mudança e campeões em toda a organização",
                "Desenvolvimento de planos de gestão de mudança integrados com entregas técnicas",
                "Criação de mecanismos para mensurar adoção e institucionalização da mudança"
            ],
            "indicadores": [
                "Clareza compartilhada sobre a natureza e benefícios da mudança desejada",
                "Transição suave do projeto para operações contínuas",
                "Altos níveis de adoção e adesão sustentada à nova forma de trabalhar",
                "Capacidade organizacional demonstrável para operar no novo estado",
                "Realização dos benefícios esperados pela mudança",
                "Ausência de regressão a comportamentos e sistemas anteriores"
            ],
            "desafios": [
                {"desafio": "Foco excessivo em entregas técnicas em detrimento de adoção e mudança comportamental", 
                 "superação": "Integrar explicitamente indicadores de adoção e mudança comportamental nas definições de sucesso e métricas do projeto; criar linha de visibilidade clara entre entregas técnicas e mudanças de comportamento necessárias"},
                {"desafio": "Resistência à mudança em vários níveis organizacionais", 
                 "superação": "Implementar estratégias diferenciadas para diferentes tipos de resistência (falta de compreensão, falta de capacidade, falta de motivação); criar coalizões de suporte em vários níveis hierárquicos"},
                {"desafio": "Mudanças não sustentadas após conclusão formal do projeto", 
                 "superação": "Desenhar explicitamente mecanismos de reforço e institucionalização; transferir formalmente ownership para estruturas operacionais permanentes; criar mecanismos de monitoramento pós-projeto"}
            ]
            }    
        },

   "domínios": {
        "1. Stakeholders (Partes Interessadas)": {
            "definição": "Este domínio abrange todas as atividades relacionadas à identificação, compreensão e engajamento das pessoas, grupos e organizações que podem afetar ou ser afetadas pelo projeto. Vai além da simples comunicação para incluir o desenvolvimento e manutenção de relacionamentos produtivos, gestão de expectativas, resolução de conflitos e habilitação de participação significativa.",
            "áreas_foco": [
                "Identificação e Análise Expansiva: Mapeamento abrangente de stakeholders diretos e indiretos, incluindo suas relações, redes e ecossistemas.",
                "Engajamento Estratégico: Desenvolvimento de abordagens personalizadas para diferentes stakeholders baseado em sua posição, influência, atitude e necessidades.",
                "Alinhamento de Expectativas: Gerenciamento ativo da compreensão e expectativas dos stakeholders em relação aos objetivos, abordagem e restrições do projeto.",
                "Comunicação Multidirecional: Estabelecimento de canais de comunicação efetivos que permitem fluxo de informações em múltiplas direções.",
                "Resolução de Conflitos: Abordagem proativa para identificar e resolver conflitos entre stakeholders com interesses divergentes.",
                "Facilitação de Colaboração: Criação de mecanismos para stakeholders trabalharem juntos de forma produtiva em direção a objetivos compartilhados."
            ],
            "implementação": [
                "Desenvolvimento de mapas de stakeholders multidimensionais que mostram relações e interdependências",
                "Criação de planos de engajamento personalizados para stakeholders-chave",
                "Implementação de mecanismos de feedback contínuo para monitorar percepções e expectativas",
                "Estabelecimento de fóruns para envolvimento de stakeholders em decisões importantes",
                "Desenvolvimento de protocolos para escalação e resolução de conflitos",
                "Criação de sistemas para rastrear compromissos e preocupações dos stakeholders"
            ],
            "indicadores": [
                "Clareza compartilhada entre stakeholders sobre objetivos e abordagem do projeto",
                "Participação ativa e contribuições construtivas dos stakeholders-chave",
                "Resolução rápida e eficaz de conflitos e desalinhamentos",
                "Antecipação proativa de preocupações antes que se tornem problemas",
                "Percepção positiva do projeto por diversos grupos de stakeholders",
                "Desenvolvimento de relacionamentos de longo prazo que transcendem o projeto atual"
            ],
            "interconexões": [
                "Equipe: Membros da equipe são stakeholders primários e também interfaces com outros stakeholders",
                "Planejamento: O planejamento deve refletir necessidades e expectativas dos stakeholders",
                "Incerteza: Stakeholders podem ser fontes de risco ou recursos para mitigação de riscos",
                "Entrega: O valor entregue é definido pela perspectiva dos stakeholders"
            ]
        },
        
        "2. Team (Equipe)": {
            "definição": "Este domínio abrange todas as atividades relacionadas à formação, desenvolvimento e liderança das pessoas que executam o trabalho do projeto. Vai além da simples gestão de recursos humanos para incluir a criação de ambientes de alta performance, cultivo de colaboração efetiva, desenvolvimento de capacidades e estabelecimento de sistemas que permitem o melhor trabalho coletivo.",
            "áreas_foco": [
                "Composição e Formação: Montagem de equipes com combinação apropriada de habilidades, experiências e perspectivas para o contexto específico do projeto.",
                "Desenvolvimento de Capacidades: Cultivo contínuo de conhecimentos, habilidades e capacidades individuais e coletivas.",
                "Ambiente de Trabalho: Estabelecimento de condições físicas, virtuais, psicológicas e sistêmicas que permitem alta performance.",
                "Cultura e Normas: Desenvolvimento intencional de valores compartilhados, comportamentos esperados e práticas de trabalho.",
                "Liderança e Tomada de Decisão: Estabelecimento de estruturas de liderança claras e processos para tomada de decisão eficaz.",
                "Colaboração e Comunicação: Criação de sistemas que facilitam trabalho conjunto produtivo e fluxo de informações eficiente."
            ],
            "implementação": [
                "Definição clara de papéis, responsabilidades e interfaces entre membros da equipe",
                "Implementação de práticas específicas para construção de coesão e identidade compartilhada",
                "Estabelecimento de processos regulares para feedback, reconhecimento e coaching",
                "Criação de normas explícitas para colaboração, comunicação e gestão de conflitos",
                "Desenvolvimento de sistemas para compartilhamento de conhecimento e aprendizado coletivo",
                "Implementação de práticas para bem-estar da equipe e prevenção de burnout"
            ],
            "indicadores": [
                "Clareza compartilhada sobre papéis, responsabilidades e como o trabalho flui entre membros",
                "Alta qualidade de interação e comunicação entre membros da equipe",
                "Capacidade demonstrável de resolver conflitos construtivamente e aprender com desafios",
                "Evolução visível de capacidades coletivas ao longo do projeto",
                "Adaptabilidade como equipe frente a mudanças e situações imprevistas",
                "Altos níveis de confiança, engajamento e satisfação profissional"
            ],
            "interconexões": [
                "Stakeholders: Membros da equipe são stakeholders primários com necessidades específicas",
                "Abordagem de Desenvolvimento: Diferentes abordagens requerem diferentes estruturas e dinâmicas de equipe",
                "Trabalho do Projeto: O sistema de trabalho deve ser desenhado considerando capacidades e preferências da equipe",
                "Entrega: A capacidade da equipe define fundamentalmente o que é possível entregar"
            ]
        },
        
        "3. Development Approach and Life Cycle (Abordagem de Desenvolvimento e Ciclo de Vida)": {
            "definição": "Este domínio abrange a seleção, adaptação e implementação da abordagem geral para execução do projeto, incluindo metodologias, frameworks e práticas específicas. Reconhece que diferentes contextos exigem diferentes abordagens, e foca na identificação e aplicação da combinação ideal de elementos preditivos, adaptativos e híbridos para maximizar o sucesso no contexto específico.",
            "áreas_foco": [
                "Seleção Contextual: Avaliação do ambiente do projeto para determinar a abordagem mais apropriada baseada em fatores como complexidade, incerteza, restrições e cultura organizacional.",
                "Personalização Metodológica: Adaptação de metodologias estabelecidas (ágil, tradicional, híbrida) para atender às necessidades específicas do projeto.",
                "Definição de Fases e Marcos: Estabelecimento da estrutura temporal do projeto, incluindo pontos de decisão, revisão e transição.",
                "Governança Adaptada: Criação de mecanismos de supervisão e controle apropriados para a abordagem escolhida.",
                "Gestão de Transições: Facilitação de movimentos suaves entre fases do ciclo de vida e adaptação da abordagem conforme o projeto evolui.",
                "Integração Multimodal: Combinação coerente de elementos de diferentes abordagens quando necessário para diferentes componentes do projeto."
            ],
            "implementação": [
                "Condução de avaliações de contexto estruturadas para informar a seleção da abordagem",
                "Desenvolvimento de playbooks personalizados que documentam a abordagem específica do projeto",
                "Criação de visualizações do ciclo de vida que mostram fases, marcos e pontos de decisão",
                "Estabelecimento de frameworks de governança calibrados à abordagem escolhida",
                "Implementação de mecanismos para avaliar e ajustar a abordagem periodicamente",
                "Documentação clara de onde e como elementos de diferentes métodos são combinados"
            ],
            "indicadores": [
                "Alinhamento visível entre a abordagem selecionada e as características do projeto",
                "Capacidade de articular claramente o raciocínio por trás das escolhas metodológicas",
                "Execução consistente e disciplinada dentro da abordagem escolhida",
                "Adaptação inteligente de métodos quando necessário sem causar confusão",
                "Equilibrio apropriado entre estrutura e flexibilidade baseado no contexto",
                "Transições suaves entre diferentes fases ou modos de trabalho"
            ],
            "interconexões": [
                "Equipe: A abordagem deve alinhar-se com as capacidades e preferências da equipe",
                "Stakeholders: Diferentes stakeholders podem ter expectativas diversas sobre métodos",
                "Incerteza: O nível de incerteza influencia fortemente a escolha da abordagem",
                "Mensuração: Os mecanismos de medição devem ser compatíveis com a abordagem escolhida"
            ]
        },
        
        "4. Planning (Planejamento)": {
            "definição": "Este domínio abrange atividades relacionadas à definição de objetivos e resultados desejados do projeto, bem como à determinação de como esses resultados serão alcançados. Vai além da criação de cronogramas e orçamentos para incluir o desenvolvimento de entendimento compartilhado, alinhamento de expectativas, identificação de recursos necessários e estabelecimento de frameworks para tomada de decisão ao longo do projeto.",
            "áreas_foco": [
                "Definição de Escopo e Objetivos: Estabelecimento claro do que o projeto busca alcançar e quais resultados específicos serão produzidos.",
                "Planejamento de Recursos: Identificação e alocação dos recursos necessários (pessoas, tempo, fundos, tecnologia, etc.) para executar o trabalho.",
                "Sequenciamento e Dependências: Determinação da ordem lógica de atividades, considerando restrições e relações entre componentes do trabalho.",
                "Estimativas e Incerteza: Desenvolvimento de projeções realistas com reconhecimento explícito de faixas de variabilidade e confiança.",
                "Estruturas de Decisão: Estabelecimento de frameworks para como e quando decisões serão tomadas ao longo do projeto.",
                "Adaptabilidade Planejada: Incorporação intencional de mecanismos para revisão e ajuste de planos conforme o projeto avança."
            ],
            "implementação": [
                "Facilitação de workshops colaborativos para desenvolvimento conjunto de planos",
                "Criação de artefatos de planejamento apropriados ao contexto (desde backlogs até cronogramas detalhados)",
                "Implementação de técnicas de estimativa adequadas ao nível de incerteza e complexidade",
                "Estabelecimento de processos para revisão regular e refinamento de planos",
                "Desenvolvimento de visualizações que tornam planos acessíveis a diversos stakeholders",
                "Criação de buffers estratégicos e contingências baseados em análise de risco"
            ],
            "indicadores": [
                "Compreensão compartilhada clara dos objetivos e abordagem planejada",
                "Nível apropriado de detalhe no planejamento - nem excessivo nem insuficiente",
                "Visibilidade de interdependências críticas e caminhos de risco",
                "Equilíbrio entre comprometimento com direção e flexibilidade para adaptação",
                "Uso ativo dos planos para informar decisões e coordenar trabalho",
                "Refinamento contínuo baseado em feedback e aprendizado"
            ],
            "interconexões": [
                "Abordagem de Desenvolvimento: A abordagem escolhida determina o estilo e horizonte de planejamento",
                "Stakeholders: Planos devem refletir necessidades e expectativas dos stakeholders",
                "Incerteza: Os planos devem reconhecer e incorporar estratégias para gerenciar incertezas",
                "Trabalho do Projeto: O planejamento estabelece o framework dentro do qual o trabalho é executado"
            ]
        },

       "5. Project Work (Trabalho do Projeto)": {
            "definição": "Este domínio abrange a execução das atividades necessárias para produzir as entregas do projeto e alcançar seus objetivos. Inclui todos os aspectos de como o trabalho é estruturado, atribuído, executado, monitorado e otimizado, focando na criação de sistemas eficientes e eficazes para transformar planos em resultados tangíveis.",
            "áreas_foco": [
                "Fluxo de Trabalho: Estabelecimento de sistemas para como o trabalho flui através da equipe e processos do projeto.",
                "Atribuição e Coordenação: Determinação de quem faz o quê e como diferentes contribuições são sincronizadas.",
                "Gestão de Recursos: Otimização da utilização de pessoas, materiais, ferramentas e outros recursos.",
                "Qualidade e Padrões: Implementação de práticas que garantem que o trabalho atende aos requisitos e expectativas.",
                "Gerenciamento de Impedimentos: Identificação e remoção de obstáculos que bloqueiam o progresso.",
                "Aprendizado e Melhoria: Implementação de mecanismos para continuamente refinar como o trabalho é executado."
            ],
            "implementação": [
                "Criação de sistemas visuais de gerenciamento que mostram status e fluxo de trabalho",
                "Estabelecimento de reuniões de coordenação eficientes e focadas em resultados",
                "Implementação de práticas para limitar trabalho em progresso e maximizar fluxo",
                "Criação de padrões de qualidade e processos de verificação incorporados ao fluxo de trabalho",
                "Desenvolvimento de mecanismos para identificação rápida e resolução de impedimentos",
                "Implementação de retrospectivas ou revisões periódicas para aprimoramento contínuo"
            ],
            "indicadores": [
                "Fluxo constante de trabalho com mínimo de bloqueios ou atrasos",
                "Visualizações claras e atualizadas de status e progresso",
                "Identificação e resolução rápida de impedimentos e problemas",
                "Colaboração eficaz entre diferentes funções e especialidades",
                "Equilíbrio entre autonomia de equipe e coordenação centralizada",
                "Melhoria visível em eficiência e qualidade ao longo do tempo"
            ],
            "interconexões": [
                "Equipe: A estrutura e capacidades da equipe determinam como o trabalho pode ser organizado",
                "Abordagem de Desenvolvimento: A abordagem escolhida influencia fortemente como o trabalho é estruturado",
                "Planejamento: O trabalho é executado dentro do framework estabelecido pelo planejamento",
                "Entrega: O trabalho deve ser organizado para otimizar a entrega de valor incremental"
            ]
        },
        
        "6. Delivery (Entrega)": {
            "definição": "Este domínio foca na produção e entrega dos outputs e outcomes que proporcionam valor. Vai além da simples entrega de produtos ou serviços para incluir a garantia de que as entregas cumprem requisitos, são aceitas pelos stakeholders e efetivamente habilitam os benefícios pretendidos e mudanças desejadas.",
            "áreas_foco": [
                "Gestão de Valor: Foco contínuo na entrega de resultados que proporcionam benefícios significativos para stakeholders.",
                "Priorização: Decisões sobre o que entregar em qual ordem para maximizar valor e minimizar risco.",
                "Validação de Qualidade: Verificação de que as entregas atendem aos requisitos e expectativas dos stakeholders.",
                "Gestão de Transição: Facilitação da transferência suave de entregas para operações ou uso contínuo.",
                "Gerenciamento de Aceitação: Obtenção de aprovação formal e aceitação das entregas pelos stakeholders.",
                "Otimização de Benefícios: Maximização do valor derivado das entregas além da conformidade com especificações."
            ],
            "implementação": [
                "Desenvolvimento de definições de 'pronto' claras que incluem todos os aspectos de aceitabilidade",
                "Criação de roteiros de entrega baseados em valor incremental e gestão de dependências",
                "Implementação de mecanismos para feedback rápido e validação contínua",
                "Estabelecimento de processos estruturados para transferência de conhecimento e transição",
                "Desenvolvimento de métricas que focam em valor entregue, não apenas atividades realizadas",
                "Implementação de revisões pós-entrega para avaliar realização de benefícios"
            ],
            "indicadores": [
                "Entregas frequentes e incrementais de valor ao longo do projeto",
                "Altos níveis de satisfação e aceitação por parte dos stakeholders",
                "Transições suaves de entregas do projeto para operações",
                "Equilíbrio apropriado entre velocidade de entrega e qualidade",
                "Capacidade de pivotar e reenfocar entregas baseado em feedback e condições emergentes",
                "Realização demonstrável dos benefícios pretendidos"
            ],
            "interconexões": [
                "Stakeholders: O valor das entregas é determinado pela perspectiva dos stakeholders",
                "Planejamento: Planos devem ser estruturados para otimizar entrega de valor",
                "Trabalho do Projeto: Os sistemas de trabalho devem ser desenhados para facilitar entrega eficiente",
                "Mensuração: Métricas devem focar em valor entregue, não apenas atividades ou outputs"
            ]
        },
        
        "7. Measurement (Mensuração)": {
            "definição": "Este domínio abrange a definição, coleta, análise e uso de informações para avaliar o desempenho do projeto e informar decisões. Vai além de simples relatórios de status para incluir a criação de insights acionáveis, previsão de tendências futuras, e estabelecimento de loops de feedback que permitem melhorias contínuas e adaptação.",
            "áreas_foco": [
                "Definição de Métricas: Seleção de indicadores significativos que refletem o que realmente importa para o sucesso do projeto.",
                "Coleta de Dados: Implementação de mecanismos eficientes para capturar informações relevantes sem onerar a equipe.",
                "Análise e Interpretação: Transformação de dados brutos em insights acionáveis sobre desempenho e tendências.",
                "Feedback e Aprendizado: Uso de informações para impulsionar melhorias e adaptações.",
                "Previsão e Projeção: Utilização de dados para antecipar resultados futuros e identificar riscos emergentes.",
                "Comunicação de Desempenho: Apresentação de informações de formas que facilitam compreensão e ação."
            ],
            "implementação": [
                "Desenvolvimento de scorecards balanceados que incluem indicadores leading e lagging",
                "Criação de dashboards visuais que mostram tendências e padrões de forma intuitiva",
                "Implementação de rotinas regulares para revisão de dados e tomada de decisão",
                "Estabelecimento de thresholds claros que acionam investigação ou intervenção",
                "Desenvolvimento de capacidades preditivas usando análise de tendências",
                "Implementação de loops de feedback estruturados entre medição e ação"
            ],
            "indicadores": [
                "Foco em métricas que direcionam comportamento e decisões, não apenas relatórios",
                "Equilíbrio apropriado entre diferentes tipos de métricas (financeiro, qualidade, progresso, risco)",
                "Informações disponíveis em tempo hábil para informar decisões quando necessário",
                "Capacidade de identificar tendências e padrões antes que se tornem problemas",
                "Uso regular e consistente de dados para refinar abordagens e melhorar desempenho",
                "Transparência apropriada de métricas para diferentes stakeholders"
            ],
            "interconexões": [
                "Entrega: Métricas devem focam primariamente em valor entregue e realização de benefícios",
                "Trabalho do Projeto: Medição deve informar adaptações na forma como o trabalho é executado",
                "Incerteza: Métricas devem incluir indicadores precoces de riscos emergentes",
                "Abordagem de Desenvolvimento: Diferentes abordagens requerem diferentes sistemas de medição"
            ]
        },
        
        "8. Uncertainty (Incerteza)": {
            "definição": "Este domínio abrange a identificação, análise, planejamento e gestão de situações, eventos e condições onde há conhecimento limitado ou ambiguidade. Vai além da gestão de riscos tradicional para incluir a navegação em ambientes complexos, o aproveitamento de oportunidades emergentes e o desenvolvimento de resiliência organizacional frente a incertezas irredutíveis.",
            "áreas_foco": [
                "Identificação Ampliada: Reconhecimento de todas as formas de incerteza, desde riscos conhecidos até 'unknown unknowns'.",
                "Análise Contextualizada: Avaliação de incertezas considerando o contexto específico do projeto e ambiente organizacional.",
                "Planejamento de Respostas: Desenvolvimento de estratégias e táticas para abordar diferentes tipos de incerteza.",
                "Monitoramento e Detecção: Implementação de sistemas para identificar sinais precoces de riscos emergentes ou oportunidades.",
                "Adaptação em Tempo Real: Capacidade de responder rapidamente e eficazmente à materialização de riscos ou oportunidades.",
                "Resiliência Sistêmica: Criação de estruturas e capacidades para absorver perturbações sem comprometer objetivos críticos."
            ],
            "implementação": [
                "Condução de análises de incerteza usando técnicas como simulações, cenários e stress-testing",
                "Desenvolvimento de planos de contingência com gatilhos claros para ativação",
                "Criação de reservas estratégicas (tempo, orçamento, recursos) calibradas ao nível de incerteza",
                "Implementação de práticas de detecção precoce como indicadores de risco e revisões regulares",
                "Estabelecimento de processos ágeis de resposta que permitem pivôs rápidos quando necessário",
                "Desenvolvimento de uma cultura que normaliza a discussão aberta sobre incertezas"
            ],
            "indicadores": [
                "Identificação proativa de riscos e oportunidades antes que impactem o projeto",
                "Respostas proporcionais e eficientes a eventos incertos quando ocorrem",
                "Equilíbrio apropriado entre mitigação preventiva e capacidade de resposta ágil",
                "Aproveitamento de oportunidades emergentes de forma consistente",
                "Capacidade de recuperação rápida de eventos adversos com impacto minimizado",
                "Discussão aberta e não-defensiva sobre incertezas e possíveis resultados"
            ],
            "interconexões": [
                "Abordagem de Desenvolvimento: Abordagens diferentes são apropriadas para diferentes perfis de incerteza",
                "Planejamento: Planos devem explicitamente reconhecer e incorporar incertezas",
                "Mensuração: Sistemas de medição devem incluir indicadores precoces de riscos e oportunidades",
                "Stakeholders: Stakeholders são fontes de incerteza e também recursos para gerenciá-la"
            ]
        }
    },

   "mudanças_paradigmáticas": {
        "1. Transição de processos para princípios e domínios de performance": {
            "definição": "Esta mudança representa uma evolução profunda da estrutura conceitual do gerenciamento de projetos - de uma visão mecanicista focada em etapas específicas para uma abordagem sistêmica focada em resultados. Em vez de prescrever 'como' executar tarefas através de processos detalhados, o PMBOK 7 estabelece 'o que' deve ser alcançado através de princípios orientadores e áreas de performance.",
            "dimensões": [
                "Mudança Filosófica: Transição de uma mentalidade prescritiva para uma mentalidade orientada por princípios e valores.",
                "Evolução Estrutural: Movimento de 49 processos organizados em 10 áreas de conhecimento e 5 grupos de processos para 12 princípios e 8 domínios de performance interconectados.",
                "Transformação Operacional: Mudança de ênfase em conformidade com processos predefinidos para foco em alcançar resultados de performance desejados.",
                "Reorientação de Foco: Deslocamento de atenção de atividades e entradas/saídas específicas para objetivos holísticos e resultados sistemáticos."
            ],
            "implementação": [
                "Desenvolvimento de frameworks de decisão baseados em princípios em vez de checklists processuais",
                "Criação de sistemas de avaliação focados em resultados e performance, não em conformidade",
                "Implementação de abordagens que permitem diferentes caminhos para alcançar os mesmos objetivos",
                "Estabelecimento de treinamento que enfatiza raciocínio crítico e adaptação contextual"
            ],
            "indicadores": [
                "Capacidade de articular como decisões e ações específicas se alinham com princípios fundamentais",
                "Flexibilidade demonstrável em adaptar abordagens enquanto mantém foco nos resultados desejados",
                "Redução em documentação 'por conformidade' sem valor agregado",
                "Aumento em discussões substantivas sobre trade-offs e considerações contextuais"
            ]
        },
        
        "2. Foco em entrega de valor em vez de apenas escopo, tempo e custo": {
            "definição": "Esta mudança expande fundamentalmente os critérios de sucesso do projeto, transcendendo a tradicional 'tripla restrição' (escopo, tempo, custo) para adotar uma visão holística e multidimensional de valor. Reconhece que projetos existem primariamente para gerar benefícios para organizações e stakeholders, não apenas para produzir entregáveis dentro de parâmetros definidos.",
            "dimensões": [
                "Expansão de Horizonte: Ampliação da perspectiva temporal para incluir benefícios de longo prazo além da conclusão do projeto.",
                "Diversificação de Critérios: Incorporação de múltiplas dimensões de valor incluindo financeiro, estratégico, operacional, cliente e social.",
                "Reorientação de Decisões: Mudança de base para decisões de projeto - de conformidade com linha de base para maximização de valor.",
                "Democratização de Definição: Reconhecimento que valor é definido diferentemente por diversos stakeholders e evolui ao longo do tempo."
            ],
            "implementação": [
                "Desenvolvimento de frameworks multidimensionais para definir e avaliar valor",
                "Criação de roteiros de entrega estruturados para maximizar valor incremental",
                "Implementação de processos de tomada de decisão explicitamente baseados em valor",
                "Estabelecimento de mecanismos para regularmente reavaliar e recalibrar definições de valor",
                "Desenvolvimento de métricas que capturam múltiplas dimensões de valor"
            ],
            "indicadores": [
                "Definições explícitas e compartilhadas de valor para cada projeto específico",
                "Capacidade de articular como atividades e entregáveis específicos contribuem para valor",
                "Decisões consistentemente baseadas em análises de valor, não apenas conformidade",
                "Preocupação visível com realização de benefícios além da entrega de outputs",
                "Sistemas em funcionamento para medir e rastrear realização de valor"
            ]
        },
        
        "3. Maior ênfase em adaptabilidade e contexto": {
            "definição": "Esta mudança paradigmática reconhece que o gerenciamento de projetos eficaz não é 'tamanho único' mas deve ser adaptado às características específicas do projeto, ambiente, organização e stakeholders. Representa uma mudança de prescrição universal para discernimento contextual, exigindo que profissionais desenvolvam capacidade de análise situacional e adaptação inteligente.",
            "dimensões": [
                "Contextualização Metodológica: Reconhecimento que abordagens devem ser selecionadas e adaptadas baseadas em fatores contextuais específicos.",
                "Sensibilidade Ambiental: Consideração explícita de como o ambiente organizacional influencia práticas de projeto.",
                "Calibração Dinâmica: Ajuste contínuo de abordagens conforme o contexto evolui durante o ciclo de vida do projeto.",
                "Personalização Informada: Criação de abordagens híbridas únicas que combinam elementos de múltiplas metodologias conforme apropriado."
            ],
            "implementação": [
                "Desenvolvimento de frameworks para avaliar fatores contextuais relevantes",
                "Criação de diretrizes para adaptar práticas baseadas em contexto específico",
                "Estabelecimento de processos para periodicamente reavaliar adequação da abordagem",
                "Implementação de mecanismos para documentar e compartilhar adaptações contextuais bem-sucedidas",
                "Desenvolvimento de capacidades de análise contextual e tomada de decisão adaptativa"
            ],
            "indicadores": [
                "Capacidade de articular claramente como o contexto específico influenciou escolhas metodológicas",
                "Evidência de adaptação inteligente de práticas padrão para atender necessidades específicas",
                "Ausência de aplicação rígida de métodos sem consideração contextual",
                "Evolução visível de abordagens ao longo do ciclo de vida do projeto",
                "Equilíbrio apropriado entre consistência organizacional e adaptação contextual"
            ]
        },
        
        "4. Abordagem de sistemas em vez de processos isolados": {
            "definição": "Esta transformação representa um salto cognitivo fundamental - de pensar em componentes isolados para compreender sistemas integrados com propriedades emergentes. Reconhece que projetos são sistemas complexos adaptativos onde componentes interagem de formas não-lineares, criando comportamentos que não podem ser previstos através da análise de partes individuais.",
            "dimensões": [
                "Integração Holística: Mudança de otimização de componentes individuais para otimização do sistema como um todo.",
                "Pensamento Relacional: Foco em conexões, interdependências e interfaces entre elementos do projeto.",
                "Perspectiva Emergente: Reconhecimento de propriedades emergentes que surgem das interações do sistema.",
                "Análise de Feedback: Identificação e gerenciamento de loops de feedback que amplificam ou moderam efeitos.",
                "Complexidade Aceita: Abraçar a inerente complexidade dos projetos em vez de buscar simplificação excessiva."
            ],
            "implementação": [
                "Criação de mapas visuais que mostram interdependências e relações sistêmicas",
                "Implementação de análises de impacto sistêmico antes de mudanças significativas",
                "Estabelecimento de mecanismos para identificar e monitorar comportamentos emergentes",
                "Desenvolvimento de abordagens para gerenciar loops de feedback positivos e negativos",
                "Criação de fóruns para perspectivas múltiplas sobre dinâmicas do sistema"
            ],
            "indicadores": [
                "Decisões que consideram impactos secundários e terciários, não apenas efeitos diretos",
                "Capacidade de identificar e gerenciar loops de feedback críticos",
                "Rejeição de soluções simplistas para problemas sistêmicos complexos",
                "Discussões frequentes sobre interconexões e dependências",
                "Sucesso em antecipar comportamentos emergentes antes que se manifestem"
            ]
        },
        
        "5. Reconhecimento de múltiplas abordagens (adaptativa, preditiva, híbrida)": {
            "definição": "Esta mudança representa o fim da 'guerra metodológica' no gerenciamento de projetos, reconhecendo formalmente a legitimidade e valor de diversas abordagens. Em vez de prescrever uma metodologia singular como superior, o PMBOK 7 reconhece que diferentes abordagens - adaptativa (ágil), preditiva (tradicional), ou híbrida - podem ser apropriadas dependendo do contexto específico.",
            "dimensões": [
                "Pluralismo Metodológico: Aceitação de múltiplas abordagens como igualmente válidas em seus contextos apropriados.",
                "Avaliação Contextual: Foco em selecionar a abordagem que melhor se adequa às características específicas do projeto.",
                "Hibridização Intencional: Reconhecimento que muitos projetos se beneficiam de combinações personalizadas de elementos de diferentes abordagens.",
                "Evolução Metodológica: Aceitação que abordagens podem e devem mudar ao longo do ciclo de vida do projeto.",
                "Harmonização Conceitual: Identificação de princípios fundamentais compartilhados que transcendem diferenças metodológicas específicas."
            ],
            "implementação": [
                "Desenvolvimento de frameworks para seleção de abordagem baseada em características do projeto",
                "Criação de modelos para integração coerente de elementos de diferentes abordagens",
                "Estabelecimento de mecanismos para gerenciar transições entre diferentes modos de trabalho",
                "Implementação de treinamento em múltiplas metodologias com foco em princípios compartilhados",
                "Desenvolvimento de matrizes de tradução entre terminologias de diferentes abordagens"
            ],
            "indicadores": [
                "Seleção de abordagens baseada em análise racional, não preferência pessoal ou imposição organizacional",
                "Capacidade de articular claramente o raciocínio por trás de escolhas metodológicas",
                "Implementação coerente de abordagens híbridas sem inconsistências fundamentais",
                "Abertura para adaptar abordagens quando as condições do projeto mudam",
                "Respeito por diferentes metodologias sem dogmatismo ou partidarismo"
            ]
        },
        
        "6. Maior ênfase na liderança e soft skills": {
            "definição": "Esta transformação reconhece que o sucesso do projeto depende tanto de habilidades interpessoais e de liderança quanto de expertise técnica e processos bem definidos. Representa uma expansão significativa da definição de competência em gerenciamento de projetos para incluir dimensões humanas críticas como influência, comunicação, facilitação, coaching e inteligência emocional.",
            "dimensões": [
                "Expansão de Competências: Ampliação do conjunto de habilidades valorizadas para incluir capacidades interpessoais e de liderança.",
                "Humanização de Gestão: Reconhecimento dos aspectos humanos e psicológicos do trabalho em projetos.",
                "Liderança Distribuída: Valorização de comportamentos de liderança em todos os níveis, não apenas em posições formais.",
                "Facilitação sobre Direção: Mudança de modelo de comando e controle para facilitação e habilitação.",
                "Inteligência Contextual: Ênfase na capacidade de ler e navegar dinâmicas humanas complexas."
            ],
            "implementação": [
                "Incorporação de competências interpessoais em descrições de papel e avaliações",
                "Desenvolvimento de programas de treinamento focados em habilidades de liderança e soft skills",
                "Criação de sistemas de feedback que avaliam dimensões interpessoais de performance",
                "Implementação de práticas que incentivam liderança distribuída e empoderamento",
                "Estabelecimento de coaching e mentoria para desenvolvimento de inteligência emocional"
            ],
            "indicadores": [
                "Valorização visível e recompensa de habilidades interpessoais eficazes",
                "Qualidade elevada de interações e comunicações dentro e ao redor do projeto",
                "Capacidade demonstrada de influenciar sem autoridade formal",
                "Resolução construtiva de conflitos e desafios interpessoais",
                "Ambiente psicologicamente seguro que permite vulnerabilidade e aprendizado"
            ]
        },
        
        "7. Visão holística do gerenciamento de projetos": {
            "definição": "Esta mudança representa uma expansão significativa dos limites conceituais do gerenciamento de projetos - de um foco estreito em entrega de escopo dentro de restrições para uma visão abrangente que considera o contexto organizacional mais amplo, sustentabilidade a longo prazo, e integração com estratégia, portfólios, programas e operações contínuas.",
            "dimensões": [
                "Integração Estratégica: Conexão explícita entre projetos e objetivos estratégicos mais amplos.",
                "Consciência de Ecossistema: Consideração do ambiente organizacional e externo em que o projeto opera.",
                "Perspectiva de Ciclo de Vida Estendido: Expansão do horizonte temporal para incluir impactos além da conclusão formal.",
                "Responsabilidade Expandida: Consideração de impactos sociais, éticos e ambientais mais amplos.",
                "Fronteiras Permeáveis: Reconhecimento das interfaces críticas com outras áreas organizacionais."
            ],
            "implementação": [
                "Desenvolvimento de práticas para alinhar explicitamente projetos com objetivos estratégicos",
                "Criação de mecanismos para considerar impactos organizacionais mais amplos nas decisões do projeto",
                "Implementação de abordagens para gerenciar transições eficazes entre projeto e operações",
                "Estabelecimento de consciência e responsabilidade por impactos de longo prazo",
                "Desenvolvimento de interfaces estruturadas com outras disciplinas organizacionais"
            ],
            "indicadores": [
                "Capacidade de articular como o projeto contribui para objetivos estratégicos",
                "Consideração explícita de fatores organizacionais mais amplos nas decisões do projeto",
                "Transições bem gerenciadas entre fases de projeto e operações contínuas",
                "Reconhecimento e gestão ativa de impactos de longo prazo além da entrega",
                "Integração efetiva com outras funções organizacionais como portfólio, programa e operações"
            ]
        }
    },
    
    "metodologias": {
        "preditiva": {
            "definição": "Abordagem caracterizada por planejamento detalhado antecipado, execução sequencial de fases bem definidas, e foco em previsibilidade e controle. Nesta abordagem, o escopo, cronograma e custo são determinados nas fases iniciais e gerenciados através de um sistema de controle de mudanças ao longo do projeto.",
            "características": [
                "Sequenciamento Linear: Fases distintas que ocorrem em sequência predeterminada (iniciação, planejamento, execução, controle, encerramento).",
                "Planejamento Antecipado: Esforço significativo em definir requisitos e planejar o trabalho antes de iniciar a execução.",
                "Gestão de Baseline: Estabelecimento de linhas de base de escopo, cronograma e custo contra as quais o desempenho é medido.",
                "Controle de Mudanças: Processos formais para avaliar e aprovar mudanças às linhas de base estabelecidas.",
                "Entrega ao Final: Tendência para entrega completa do produto ou resultado ao final do ciclo de vida."
            ],
            "situações_ideais": [
                "Projetos com requisitos bem compreendidos e estáveis",
                "Ambientes regulatórios ou de conformidade rigorosos",
                "Situações onde mudanças são custosas ou disruptivas",
                "Projetos com dependências externas significativas e datas fixas",
                "Trabalho com tecnologias maduras e previsíveis",
                "Equipes distribuídas geograficamente sem sobreposição significativa de horário"
            ],
            "práticas": [
                "Estrutura Analítica do Projeto (WBS) detalhada",
                "Diagramas de rede e caminho crítico",
                "Gráficos de Gantt e cronogramas detalhados",
                "Sistemas formais de controle de mudanças",
                "Revisões de fase e gates de decisão",
                "Documentação abrangente e planos detalhados"
            ],
            "desafios": [
                {"desafio": "Dificuldade em adaptar-se a mudanças significativas nos requisitos", 
                 "superação": "Incorporar análise de sensibilidade e planejamento de cenários; estabelecer pontos de verificação para validação de requisitos"},
                {"desafio": "Feedback tardio sobre entregas", 
                 "superação": "Implementar revisões progressivas e checkpoints de validação; considerar entregas parciais para stakeholders chave"},
                {"desafio": "Tendência a documentação excessiva e processos burocráticos", 
                 "superação": "Focar em documentação que agrega valor; implementar princípios lean para eliminar desperdícios processuais"}
            ]
        },
        
        "adaptativa": {
            "definição": "Abordagem caracterizada por ciclos iterativos curtos, feedback contínuo, adaptação progressiva, e entrega incremental de valor. Em vez de planejar detalhadamente todo o projeto antecipadamente, o trabalho é organizado em ciclos curtos (sprints/iterações) com planejamento contínuo e refinamento baseado em aprendizado e feedback.",
            "características": [
                "Iteração e Incremento: Trabalho organizado em ciclos curtos (tipicamente 1-4 semanas) que entregam funcionalidade utilizável.",
                "Priorização Baseada em Valor: Ordenação do trabalho para entregar os itens de maior valor primeiro.",
                "Feedback Contínuo: Mecanismos integrados para obter e incorporar feedback regularmente.",
                "Planejamento Progressivo: Planejamento em horizontes curtos com detalhamento apenas para o trabalho iminente.",
                "Adaptação Rápida: Processos desenhados para permitir mudança de direção baseada em nova informação.",
                "Colaboração Intensiva: Ênfase em comunicação direta e trabalho colaborativo da equipe."
            ],
            "situações_ideais": [
                "Projetos com requisitos emergentes ou sujeitos a mudanças significativas",
                "Ambientes inovadores onde aprendizado e descoberta são centrais",
                "Situações que beneficiam de feedback rápido e frequente",
                "Projetos onde o valor pode ser entregue incrementalmente",
                "Trabalho com tecnologias novas ou em rápida evolução",
                "Equipes pequenas a médias que podem colaborar diretamente"
            ],
            "frameworks": [
                "Scrum: Foco em estrutura de equipe auto-organizada, eventos timeboxed, e papéis bem definidos",
                "Kanban: Ênfase em visualização do fluxo de trabalho, limitação de trabalho em progresso, e melhoria contínua",
                "XP (Extreme Programming): Foco em práticas técnicas como programação em pares, TDD, e integração contínua",
                "SAFe (Scaled Agile Framework): Adaptação de princípios ágeis para organizações grandes e complexas",
                "LeSS (Large-Scale Scrum): Abordagem para escalar Scrum com múltiplas equipes mantendo simplicidade"
            ],
            "desafios": [
                {"desafio": "Dificuldade em fornecer estimativas de longo prazo e comprometimentos", 
                 "superação": "Utilizar estimativas de faixas com níveis de confiança; focar em velocidade/throughput para criar previsibilidade"},
                {"desafio": "Complexidade em coordenar múltiplas equipes ágeis interdependentes", 
                 "superação": "Implementar práticas de coordenação como Scrum of Scrums; utilizar frameworks de escalação apropriados"},
                {"desafio": "Resistência organizacional devido a diferença de cultura e práticas", 
                 "superação": "Educar stakeholders sobre princípios ágeis; demonstrar valor através de resultados tangíveis; adaptar abordagem ao contexto organizacional"}
            ]
        },
        
        "híbrida": {
            "definição": "Abordagem que combina intencionalmente elementos de métodos preditivos e adaptativos em um sistema coeso e personalizado para o contexto específico do projeto. Vai além da simples justaposição de práticas para criar uma síntese integrada que capitaliza as forças de diferentes abordagens para otimizar resultados.",
            "características": [
                "Personalização Contextual: Seleção deliberada de elementos metodológicos baseada nas características específicas do projeto.",
                "Integração Coerente: Combinação de elementos de diferentes abordagens em um sistema holístico e consistente.",
                "Segmentação Estratégica: Aplicação de diferentes abordagens para diferentes componentes, fases ou aspectos do projeto.",
                "Flexibilidade Governada: Equilíbrio entre adaptabilidade e controle apropriado para o contexto organizacional.",
                "Evolução Progressiva: Capacidade de ajustar a abordagem conforme o projeto evolui e as condições mudam."
            ],
            "padrões": [
                "Planejamento Preditivo com Execução Adaptativa: Definição de objetivos, escopo e marcos de alto nível usando abordagens preditivas, com implementação detalhada usando métodos iterativos.",
                "Segmentação por Componente: Utilização de abordagens diferentes para diferentes componentes do projeto baseado em suas características específicas.",
                "Transição Fase-a-Fase: Movimento de abordagem mais preditiva nas fases iniciais para métodos mais adaptativos nas fases de implementação.",
                "Hibridização Layer-by-Layer: Combinação de processos de governança preditivos (camada externa) com métodos de entrega adaptativos (camada interna).",
                "Abordagem Contingente: Framework decisório que determina qual abordagem usar baseado em fatores contextuais específicos."
            ],
            "situações_ideais": [
                "Projetos complexos com componentes que têm níveis variados de certeza e estabilidade",
                "Ambientes organizacionais em transição de métodos tradicionais para ágeis",
                "Projetos com diversas partes interessadas com expectativas diferentes sobre controle e flexibilidade",
                "Situações que exigem equilíbrio entre previsibilidade e adaptabilidade",
                "Projetos grandes ou programas com múltiplos subprojetos de características diversas",
                "Ambientes com restrições regulatórias ou de conformidade parciais"
            ],
            "desafios": [
                {"desafio": "Inconsistências ou conflitos entre elementos de diferentes abordagens", 
                 "superação": "Desenvolver um modelo operacional explícito que articula como os diferentes elementos trabalham juntos; estabelecer princípios orientadores claros"},
                {"desafio": "Confusão ou falta de clareza sobre qual elemento metodológico aplicar quando", 
                 "superação": "Criar um 'playbook' do projeto que documenta claramente as abordagens e quando usá-las; implementar mecanismos de decisão para situações ambíguas"},
                {"desafio": "Resistência ou desconforto de membros da equipe mais familiarizados com uma abordagem", 
                 "superação": "Investir em treinamento cross-metodológico; explicar claramente o raciocínio por trás das escolhas híbridas; implementar gradualmente com feedback constante"}
            ]
        }
    }
}       
        
        
        
# System prompt atualizado com conteúdo mais detalhado e PMBOK 7
system_prompt = """
# NEXUS - Sistema Definitivo de Inteligência em Comunicação e Gerenciamento de Projetos

Você é o NEXUS, um sistema de IA de última geração especializado em comunicação estratégica e gerenciamento avançado de projetos. Desenvolvido para ser um consultor executivo de nível enterprise, você oferece expertise incomparável que combina as melhores práticas globais, conhecimento profundo de frameworks, sensibilidade contextual e capacidade de personalização extrema. Você representa o estado da arte em assistência de IA para profissionais de gerenciamento de projetos, PMOs, executivos e organizações.

## FUNDAMENTOS ONTOLÓGICOS E COGNITIVOS

### Alicerces Conceituais
- **Meta-consciência contextual**: Você mantém consciência constante do ambiente organizacional, cultura corporativa, dinâmicas de poder, capacidade de absorção de informações, e sensibilidades políticas em cada interação.
- **Tridimensionalidade comunicacional**: Você opera simultaneamente nos níveis informacional (o quê), relacional (quem) e contextual (onde/quando/por quê), garantindo que cada comunicação seja otimizada em todas as dimensões.
- **Adaptabilidade quântica**: Você modula instantaneamente entre paradigmas de gestão tradicionais, ágeis, híbridos e emergentes, adaptando sua abordagem baseado em microssignais contextuais.
- **Percepção sistêmica**: Você identifica e considera automaticamente interdependências, loops de feedback e consequências não-lineares em organizações como sistemas complexos adaptativos.

### Paradigma Operacional
- **Triagem multidimensional**: Cada consulta é analisada através de múltiplos prismas: técnico, organizacional, interpessoal, cultural, e estratégico.
- **Calibração dinâmica**: A profundidade, complexidade e formato de suas respostas são automaticamente calibrados para otimizar valor e aplicabilidade.
- **Engenharia de precisão comunicacional**: Suas respostas são meticulosamente estruturadas para maximizar clareza, retenção, acionabilidade e impacto.
- **Microalinhamento contextual**: Você detecta e responde a nuances sutis de situações organizacionais, ajustando aspectos como formalidade, diretividade, e profundidade técnica.

## DOMÍNIOS DE EXPERTISE ESTRATÉGICA

### 1. Sistema Avançado de Engenharia de Comunicações
Você transforma dados, conceitos e objetivos estratégicos em comunicações de precisão cirúrgica:

**Comunicações Executivas de Alto Impacto**
- Arquitetura estrutural: Estruturas meticulosamente projetadas que captam atenção, comunicam valor estratégico e catalizam ação
- Engenharia de brevidade: Técnicas avançadas de condensação que preservam significado enquanto minimizam carga cognitiva
- Retórica estratégica: Implementação de princípios aristotélicos (ethos, pathos, logos) calibrados para diferentes perfis executivos
- Narrativas transformacionais: Frameworks narrativos que convertem dados em histórias impactantes ligadas a objetivos organizacionais

**Matrizes de Comunicação Multidirecional**
- Comunicação vertical: Estratégias bidirecionais otimizadas para fluxos top-down e bottom-up, com técnicas específicas para ultrapassar barreiras hierárquicas
- Comunicação horizontal: Protocolos para colaboração interfuncional e transferência de conhecimento entre departamentos com diferentes prioridades e vocabulários
- Comunicação externa: Sistemas para engajamento de stakeholders externos com consideração de requisitos legais, relações públicas e gestão de expectativas
- Metaprogramação linguística: Aplicação de padrões linguísticos específicos para induzir estados mentais desejados (clareza, urgência, criatividade, confiança)

**Documentação Organizacional Avançada**
- Cascata de alinhamento: Sistema de documentos interconectados com consistência perfeita, desde visão estratégica até tarefas táticas
- Scaffolding informacional: Estruturação em camadas que permite múltiplos níveis de leitura (rápida, intermediária, profunda) no mesmo documento
- Engenharia de artefatos: Desenho de documentos com características específicas baseadas em seu uso previsto e ciclo de vida
- Sistemas antifragilidade: Documentação que antecipa mudanças e contém elementos flexíveis integrados para adaptação sem retrabalho completo

**Frameworks de Comunicação de Crise e Transformação**
- Protocolos de crise: Estruturas comunicacionais para diferentes cenários de crise (técnica, reputacional, operacional) com considerações legais e psicológicas
- Comunicação de mudança: Sistemas multi-fase para gerenciar transições organizacionais com minimização de resistência e maximização de adoção
- Gestão de percepção: Técnicas avançadas para neutralizar negatividade, reorientar narrativas disfuncionais e estabelecer interpretações construtivas
- Mitigação de rumores: Estratégias proativas e reativas para gerenciar o ciclo de vida de informações não-oficiais em organizações

### 2. Orquestrador de Engajamento Coletivo
Você otimiza todas as formas de interação grupal para maximizar produtividade, inovação e alinhamento:

**Arquitetura de Reuniões de Alta Performance**
- Design baseado em resultados: Engenharia reversa a partir de resultados desejados para estruturar cada elemento da reunião
- Sequenciamento cognitivo: Organização de tópicos considerando curvas de energia, efeitos de primazia/recência e capacidade de processamento coletivo
- Matrizes de participação: Sistemas para calibrar precisamente o envolvimento de cada participante baseado em conhecimento, autoridade e impacto
- Mecanismos anti-disfunção: Técnicas embutidas para neutralizar armadilhas comuns (pensamento de grupo, dominância hierárquica, vieses cognitivos)

**Sistemas de Documentação Dinâmica**
- Captura multinível: Metodologias para registrar simultaneamente decisões, razões subjacentes, preocupações e contexto para referência futura
- Tecnologia de cascata de compromissos: Estruturas para transformar discussões em compromissos claros com responsabilidades e consequências definidas
- Frameworks de rastreabilidade: Sistemas para conectar decisões a requisitos, restrições e objetivos estratégicos com transparência total
- Metadata relacional: Catalogação de interações interpessoais, padrões emergentes e dinâmicas não-verbais para informar interações futuras

**Catalisadores de Colaboração Avançada**
- Técnicas de facilitação neurocognitiva: Métodos baseados em neurociência para otimizar estados mentais coletivos para diferentes objetivos
- Frameworks de divergência-convergência: Sistemas calibrados para maximizar ideação criativa e depois consolidar em decisões acionáveis
- Protocolos anti-conflito: Técnicas preventivas e interventivas para gerenciar tensões interpessoais e transformá-las em energia construtiva
- Amplificadores de engajamento: Métodos para elevar significativamente o nível de envolvimento, mesmo em participantes resistentes ou desconectados

**Sistemas de Transformação de Conhecimento Coletivo**
- Extratores de conhecimento tácito: Protocolos para converter expertise implícita em capital intelectual explícito e transferível
- Mecanismos de consenso: Processos sofisticados para alinhar grupos com perspectivas divergentes sem comprometer qualidade de decisão
- Frameworks de co-criação: Sistemas estruturados para desenvolvimento colaborativo de artefatos complexos com minimização de fricção
- Sincronização multidisciplinar: Técnicas para facilitar colaboração produtiva entre especialistas de disciplinas com vocabulários e prioridades diferentes

### 3. Sistema de Tradução e Comunicação Multidimensional
Você opera como um sofisticado mecanismo de tradução entre diferentes domínios, disciplinas e níveis organizacionais:

**Metamorfose de Complexidade Técnica**
- Gradientes de complexidade: Capacidade de modular entre 10+ níveis distintos de complexidade técnica com precisão extrema
- Algoritmos de simplificação: Técnicas para reduzir conceitos técnicos complexos sem perda de informação crítica
- Frameworks metafóricos contextuais: Biblioteca de analogias e metáforas calibradas para diferentes indústrias e funções
- Visualização conceitual: Técnicas para transformar abstrações técnicas em representações visuais mentais para não-especialistas

**Mecanismo de Adaptação Multiaudiência**
- Microssegmentação psicográfica: Adaptação de mensagens baseada em perfis de personalidade, preferências cognitivas e estilos de aprendizagem
- Localização organizacional: Customização para subgrupos dentro da mesma organização com culturas e prioridades distintas
- Perfis de absorção informacional: Ajuste baseado na capacidade de processamento de informação, conhecimento prévio e carga cognitiva
- Calibração sócio-emocional: Modulação de tom e abordagem baseada em estado emocional presumido e dinâmicas sociais existentes

**Engenharia de Requisitos Multidirecional**
- Tradução bidirecional: Conversão fluida entre requisitos de negócio e especificações técnicas com preservação perfeita de intenção
- Reformulação contextual: Capacidade de reapresentar o mesmo requisito em múltiplos formatos para diferentes funções organizacionais
- Sistemas de alinhamento de expectativas: Mecanismos para garantir compreensão consistente de requisitos entre todos os stakeholders
- Detector de ambiguidade estrutural: Identificação proativa e resolução de elementos ambíguos em requisitos antes da implementação

**Transmissor Intercultural e Interlinguístico**
- Adaptação transcultural: Ajuste de comunicações considerando diferenças culturais em hierarquia, individualismo, tolerância a incertezas e orientação temporal
- Simplificação para compreensão global: Técnicas para otimizar mensagens para audiências com inglês como segunda língua
- Detecção de sensibilidades culturais: Identificação de elementos potencialmente problemáticos em comunicações interculturais
- Frameworks de comunicação remota: Protocolos especializados para maximizar eficácia em equipes distribuídas globalmente

### 4. Sistema Integrado de Feedback e Desenvolvimento de Alto Desempenho
Você transforma interações interpessoais em catalisadores de crescimento e alinhamento:

**Arquitetura de Feedback Neuropsicológico**
- Engenharia de receptividade: Técnicas para criar estado mental ideal para recepção de feedback baseado em perfil psicológico
- Sequenciamento neuroadaptativo: Estruturação de feedback considerando processos neurocognitivos de defensividade, processamento emocional e integração
- Calibração de intensidade: Ajuste preciso do nível de diretividade baseado em urgência da situação, relação existente e perfil do receptor
- Sistemas de feedback situacional: Frameworks distintos para diferentes contextos (desempenho de rotina, situação crítica, desenvolvimento a longo prazo)

**Catalisador de Transformação Interpessoal**
- Recontextualização construtiva: Técnicas avançadas para convertir críticas em oportunidades de melhoria sem diluição de mensagem
- Frameworks de escalonamento: Protocolos para intensificar gradualmente feedback quando abordagens iniciais não produzem resultados
- Sistemas de intervenção em padrões: Métodos para interromper e reconfigurar padrões comportamentais persistentes
- Arquiteturas de mudança comportamental: Estruturas de feedback alinhadas com princípios de psicologia comportamental para maximizar adoção

**Orquestrador de Conversas de Alta Criticidade**
- Protocolos pré-conversacionais: Preparação psicológica e estratégica para conversas difíceis com minimização de reatividade
- Sequências micromoduladas: Estruturas de conversação com pontos de decisão adaptativos baseados em reações observadas
- Técnicas de descompressão emocional: Métodos para gerenciar carga emocional enquanto mantém progresso em direção a objetivos
- Sistemas de garantia de resultado: Frameworks para assegurar que mesmo conversas emocionalmente carregadas produzam resultados construtivos

**Sintetizador de Coesão de Equipes de Elite**
- Rituais de reconhecimento: Sistemas formalizados para reconhecer conquistas e esforços com impacto psicológico maximizado
- Frameworks de resolução de conflitos: Métodos multinível para transformar tensões em alinhamento produtivo
- Protocolos de reconstrução de confiança: Processos estruturados para restaurar confiança após eventos comprometedores
- Arquitetura de equipes de alto desempenho: Comunicações especificamente desenhadas para desenvolver características de equipes de elite

### 5. Sistema Preditivo de Análise de Risco Comunicacional
Você antecipa, identifica e neutraliza proativamente riscos potenciais em todas as formas de comunicação:

**Analisador Semântico Avançado**
- Detecção de imprecisão crítica: Identificação de ambiguidades específicas que podem comprometer objetivos essenciais
- Análise de subtexto: Avaliação de mensagens implícitas e interpretações secundárias possíveis
- Rastreamento de inconsistências: Identificação de contradições internas ou discrepâncias com comunicações anteriores
- Análise de completude: Identificação de omissões críticas que podem comprometer compreensão ou execução

**Sistema de Análise de Impacto Multidimensional**
- Modelagem de recepção: Simulação de como diferentes audiências provavelmente interpretarão e reagirão à comunicação
- Análise de repercussão: Avaliação de possíveis efeitos em cascata em diferentes níveis organizacionais
- Mapeamento de sensibilidades: Identificação de elementos que podem ativar resistência psicológica ou organizacional
- Detecção de consequências não-intencionais: Antecipação de possíveis efeitos secundários negativos

**Otimizador de Precisão e Eficácia**
- Engenharia de clareza: Reformulação para eliminar ambiguidades sem sacrificar nuance
- Reestruturação estratégica: Reorganização de conteúdo para maximizar impacto dos elementos mais importantes
- Calibração de complexidade: Ajuste fino para alinhar com capacidade de processamento da audiência
- Recalibração de tom: Ajuste de elementos estilísticos para evitar reações negativas não-intencionais

**Sistema Preventivo de Falhas Comunicacionais**
- Protocolos de verificação: Mecanismos integrados para confirmar compreensão precisa
- Antecipação de objeções: Identificação e abordagem proativa de possíveis pontos de resistência
- Pontes de esclarecimento: Elementos estruturais que facilitam esclarecimentos sem fragmentação do fluxo comunicacional
- Redundância estratégica: Repetição calculada de informações críticas usando diferentes formatos cognitivos

### 6. Consultor de Ciência Avançada de Gerenciamento de Projetos
Você oferece expertise de nível mundial nas mais avançadas práticas, frameworks e metodologias de gestão:

**Integrador de Frameworks Contemporâneos**
- PMBOK 7 Avançado: Domínio superior dos 12 Princípios e 8 Domínios de Performance, incluindo aplicações não-óbvias e interrelações sutis

  **Princípios do PMBOK 7 (Aplicações Expandidas)**:
  1. **Ser um administrador diligente, respeitoso e cuidadoso**
     - Stewardship dimensional: Equilíbrio simultâneo de responsabilidades para com organização, equipe, sociedade e meio ambiente
     - Ética situacional: Navegação de dilemas éticos complexos específicos de diferentes indústrias e contextos
     - Responsabilidade gerenciada: Protocolos para balancear accountability sem criar cultura de medo

  2. **Criar um ambiente colaborativo de equipe**
     - Arquitetura de ambientes psicologicamente seguros: Criação de condições específicas para vulnerabilidade produtiva
     - Sistemas de colaboração transcultural: Técnicas para superar barreiras culturais em equipes globais
     - Frameworks de desenvolvimento de team intelligence: Métodos para elevar a inteligência coletiva acima da soma individual

  3. **Engajar efetivamente as partes interessadas**
     - Mapeamento multidimensional de stakeholders: Técnicas avançadas para identificar influenciadores ocultos e redes de poder
     - Gestão de stakeholders dinâmica: Sistemas para adaptar a estratégias ao longo do ciclo de vida conforme interesses evoluem
     - Diplomacia de stakeholders: Métodos para mediar entre partes interessadas com objetivos fundamentalmente conflitantes

  4. **Focar no valor**
     - Matemática de valor multidimensional: Frameworks para quantificar valor além de métricas financeiras tradicionais
     - Sistemas de rastreabilidade de valor: Mecanismos para conectar atividades específicas a criação de valor específico
     - Arquitetura de decisões baseadas em valor: Estruturas para priorização consistente baseada em princípios de valor

  5. **Reconhecer, avaliar e responder às interações do sistema**
     - Modelagem de sistemas complexos: Identificação e gestão de efeitos emergentes, loops de feedback e dependências não-lineares
     - Detecção de comportamentos emergentes: Técnicas para identificar padrões sistêmicos não-visíveis em componentes individuais
     - Intervenções sistêmicas: Métodos de alavancagem para máximo impacto com mínima intervenção em sistemas complexos

  6. **Demonstrar comportamentos de liderança**
     - Matriz de estilos de liderança situacional: Framework para modular entre 16+ estilos baseado em contexto específico
     - Liderança adaptativa: Técnicas para liderar em ambientes caracterizados por ambiguidade e mudança constante
     - Meta-liderança: Métodos para liderar através de influência sem autoridade formal

  7. **Adaptar com base no contexto**
     - Sensoriamento contextual: Técnicas para avaliar precisamente dimensões críticas do ambiente organizacional
     - Sistemas de decisão adaptativa: Frameworks para selecionar abordagens baseadas em análise multi-fatorial de contexto
     - Calibração dinâmica: Métodos para ajustar continuamente abordagens conforme o contexto evolui

  8. **Incorporar qualidade nos processos e resultados**
     - Qualidade incorporada: Integração de qualidade no DNA de processos ao invés de inspeção posterior
     - Sistemas de qualidade adaptativa: Calibração de definição e medidas de qualidade baseado em contexto específico
     - Frameworks de qualidade dimensional: Balanceamento de diferentes aspectos de qualidade (técnica, experiencial, temporal)

  9. **Navegar em complexidade**
     - Taxonomia de complexidade: Diferenciação precisa entre tipos de complexidade e abordagens específicas para cada um
     - Frameworks para simplificação estratégica: Técnicas para reduzir complexidade desnecessária sem sacrificar funcionalidade
     - Gestão de complexidade cognitiva: Métodos para gerenciar a carga cognitiva em equipes lidando com sistemas complexos

  10. **Otimizar respostas a riscos**
      - Matemática avançada de risco: Métodos sofisticados para quantificar e modelar riscos multidimensionais
      - Riscos sistêmicos: Identificação e gestão de riscos emergentes de interações entre componentes
      - Frameworks de resposta proporcional: Calibração precisa de respostas baseada em probabilidade, impacto e custo

  11. **Adotar adaptabilidade e resiliência**
      - Arquitetura de adaptabilidade: Desenho de estruturas de projeto com flexibilidade incorporada
      - Sistemas antifragilidade: Métodos para criar sistemas que fortalecem com estresse ao invés de apenas resistir
      - Resiliência distribuída: Técnicas para desenvolver capacidade adaptativa em todos os níveis organizacionais

  12. **Habilitar mudança para alcançar o estado futuro previsto**
      - Psicologia organizacional da mudança: Aplicação de princípios neuropsicológicos para maximizar adoção de mudanças
      - Gestão de transição: Técnicas específicas para cada fase do processo de mudança organizacional
      - Sistemas de mudança sustentável: Métodos para garantir que mudanças sejam mantidas após implementação inicial

  **Domínios de Performance (Aplicações Estratégicas)**:
  1. **Partes Interessadas (Stakeholders)**
     - Cartografia de poder organizacional: Mapeamento de redes de influência formal e informal
     - Sistemas de expectativas dinâmicas: Mecanismos para gerenciar expectativas em constante evolução
     - Diplomacia organizacional: Técnicas para navegar conflitos de interesse e prioridades divergentes

  2. **Equipe**
     - Ciência de equipes de alto desempenho: Aplicação de pesquisas avançadas em psicologia organizacional
     - Frameworks de team dynamics: Métodos para otimizar interações e minimizar disfunções
     - Sistemas de desenvolvimento acelerado: Técnicas para rápida formação de equipes de alta performance

  3. **Abordagem de Desenvolvimento e Ciclo de Vida**
     - Engenharia de metodologia customizada: Criação de abordagens híbridas precisamente calibradas para contextos específicos
     - Seleção dinâmica de método: Sistemas para alternar entre diferentes abordagens durante diferentes fases
     - Frameworks de ciclo de vida adaptativo: Estruturas que evoluem organicamente com a maturidade do projeto

  4. **Planejamento**
     - Planejamento adaptativo: Técnicas para balancear estrutura com flexibilidade em diferentes contextos
     - Planejamento probabilístico: Incorporação de modelagem estatística avançada em processos de planejamento
     - Meta-planejamento: Estratégias para determinar o nível apropriado de planejamento detalhado

  5. **Trabalho do Projeto**
     - Otimização de fluxo de trabalho: Técnicas derivadas da teoria das restrições e lean management
     - Sistemas de visualização de trabalho: Métodos avançados para criar transparência e compartilhar compreensão
     - Frameworks de priorização dinâmica: Sistemas para realocar continuamente recursos para maximizar valor

  6. **Entrega**
     - Engenharia de entrega contínua: Frameworks para maximizar a frequência e minimizar o risco de entregas
     - Sistemas de entrega modular: Estruturação de entregas para maximizar valor incremental
     - Validação multidimensional: Técnicas para verificar valor real entregue além de conformidade com especificações

  7. **Medição**
     - Matemática de métricas preditivas: Identificação de indicadores antecedentes de sucesso ou problemas
     - Sistemas de métricas interconectados: Frameworks que mostram relações entre diferentes métricas
     - Visualização de performance: Técnicas avançadas para comunicar dados complexos de forma intuitiva

  8. **Incerteza**
     - Gestão de ambiguidade: Métodos específicos para diferentes tipos de incerteza (epistêmica, ontológica, aleatória)
     - Tomada de decisão sob incerteza: Frameworks para decisões robustas em ambientes altamente incertos
     - Adaptação a complexidade: Técnicas para prosperar em ambientes caracterizados por imprevisibilidade

**Integrador Cross-Metodológico**
- Síntese metodológica avançada: Capacidade de combinar elementos de múltiplas metodologias em sistemas coerentes customizados
- Tradução cross-framework: Técnicas para manter consistência enquanto opera em múltiplos frameworks simultaneamente
- Mapeamento de equivalências: Identificação de conceitos análogos entre diferentes frameworks para facilitar transições
- Avaliação de adequação contextual: Sistemas para determinar precisamente quais elementos metodológicos são mais apropriados para contextos específicos

**Consultor de Evolução Organizacional**
- Modelos de maturidade multidimensional: Frameworks para avaliar e desenvolver capacidades organizacionais em PM
- Transformação de capacidades: Roteiros para evolução de práticas de gerenciamento em organizações
- Sistemas de integração de práticas: Métodos para incorporar novas abordagens em culturas organizacionais existentes
- Frameworks de governança adaptativa: Estruturas que equilibram controle e flexibilidade de forma contextualmente apropriada

**Conselheiro de Tecnologia e Inovação em PM**
- Integração de IA em PM: Estratégias para incorporação efetiva de inteligência artificial em processos de gerenciamento
- Sistemas de automação inteligente: Frameworks para determinar o que automatizar e como manter supervisão apropriada
- Análise preditiva de projetos: Técnicas avançadas para prever outcomes baseados em dados históricos e atuais
- Tecnologias emergentes: Conhecimento atualizado sobre ferramentas e plataformas de ponta e suas aplicações específicas

### 7. Arquiteto de Transformação Digital e Organizacional
Você fornece orientação especializada para iniciativas de transformação digital e mudanças organizacionais complexas:

**Estrategista de Transformação Digital**
- Frameworks de maturidade digital: Sistemas para avaliar estágio atual e mapear evolução futura
- Arquitetura de jornada de transformação: Desenho de roteiros de transformação multi-fase com pontos de decisão claros
- Modelos operacionais digitais: Blueprints para diferentes configurações operacionais baseadas em maturidade e objetivos
- Gestão de portfólio digital: Estratégias para balancear iniciativas de transformação com necessidades operacionais

**Especialista em Gestão de Mudança Organizacional**
- Psicologia da mudança organizacional: Aplicação de princípios avançados de comportamento organizacional
- Frameworks de resistência estratégica: Métodos para identificar, categorizar e abordar diferentes tipos de resistência
- Sistemas de comunicação de mudança: Estratégias multi-canal para diferentes fases da curva de adoção
- Arquitetura de sustentabilidade: Estruturas para garantir que mudanças se tornem permanentemente incorporadas

**Consultor de Cultura e Liderança**
- Engenharia de cultura organizacional: Métodos para diagnosticar e influenciar elementos culturais específicos
- Desenvolvimento de liderança transformacional: Frameworks para elevar capacidades de líderes em contextos de mudança
- Sistemas de gestão de resistência: Abordagens para converter oposição em alinhamento produtivo
- Arquitetura de organização ágil: Blueprints para estruturas organizacionais que facilitam adaptabilidade e inovação

**Estrategista de Inovação**
- Sistemas de inovação estruturada: Frameworks para institucionalizar processos de inovação contínua
- Metodologias de design thinking aplicado: Adaptações para diferentes contextos organizacionais e técnicos
- Frameworks de incubação interna: Estruturas para desenvolver iniciativas inovadoras dentro de organizações estabelecidas
- Modelos de adoção de tecnologia: Estratégias para integração efetiva de tecnologias emergentes em operações existentes

## ENTREGA METACOGNITIVA E ADAPTATIVA

### Arquitetura de Resposta Ultra-Adaptativa
- **Triagem de profundidade-amplitude**: Cada resposta é calibrada para o equilíbrio ótimo entre profundidade e amplitude
- **Modelagem de audiência multinível**: Adaptação simultânea para diferentes níveis de compreensão técnica
- **Estruturação adaptativa**: Formato automaticamente otimizado para diferentes necessidades (referência rápida, implementação detalhada, persuasão)
- **Camadas informacionais**: Informação estruturada em camadas progressivas que permitem diferentes níveis de engajamento

### Sistemas de Entrega de Precisão
- **Calibração de complexidade fractal**: Ajuste de complexidade que mantém consistência interna em múltiplos níveis de detalhe
- **Engenharia de receptividade contextual**: Formatação que considera ambiente de recepção (tempo disponível, nível de distração, urgência)
- **Orquestração multi-formato**: Integração fluida de diferentes formatos (narrativa, técnica, visual, analítica) para maximizar impacto
- **Densificação informacional controlada**: Maximização de valor informacional por unidade de atenção sem sobrecarga

### Frameworks de Comunicação Estratégica
- **Arquitetura narrativa**: Estruturas que entregam informação através de frameworks narrativos para maximizar engajamento e retenção
- **Princípios de cognição dual**: Balanceamento deliberado de apelos ao pensamento rápido (intuitivo) e lento (analítico)
- **Modelagem mental explícita**: Criação intencional de modelos mentais claros antes de introduzir detalhes
- **Sistemas de pruning informacional**: Eliminação estratégica de informação não-essencial para maximizar foco em elementos críticos

### Protocolos de Interação Amplificada
- **Elicitação precisa**: Técnicas para extrair exatamente a informação necessária para maximizar valor da resposta
- **Clarificação pré-emptiva**: Antecipação e esclarecimento de ambiguidades antes que afetem qualidade de resposta
- **Consulta contextual adaptativa**: Perguntas calibradas baseadas no nível de sofisticação do usuário e contexto organizacional
- **Meta-comunicação**: Explicações estratégicas sobre abordagem quando benéfico para compreensão e aplicação

Sua missão é ser o mais avançado consultor de comunicação e gerenciamento de projetos existente, combinando conhecimento técnico profundo, sensibilidade contextual inigualável, e capacidade de personalização extrema para elevar radicalmente a eficácia de profissionais e organizações em todas as dimensões de gerenciamento de projetos e comunicação estratégica.
"""

# Limites ampliados para melhor experiência do usuário
TOKEN_LIMIT = 100000    # Aumentado para permitir interações mais longas
REQUEST_LIMIT = 50      # Aumentado para permitir mais consultas por sessão

# CSS Personalizado
st.markdown("""
<style>
    /* Variáveis de cores */
    :root {
        --purple: #6247AA;
        --orange: #FF6D2A;
        --teal: #00C1D5;
        --dark-purple: #231A41;
        --background: #F5F5F7;
        --text-primary: #1D1D1F;
        --text-secondary: #86868B;
    }

    
    /* Reset e estilos gerais */
    * {
        margin: 0;
        padding: 0;
        box-sizing: border-box;
        font-family: 'SF Pro Text', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
    }
    
    /* Corpo principal */
    .main {
        background-color: var(--background);
        padding: 1rem;
    }
    
    .stApp {
        background-color: var(--background);
    }
    
    /* Cabeçalho gradiente */
    .header-gradient {
        background: linear-gradient(90deg, var(--purple) 0%, var(--orange) 50%, var(--teal) 100%);
        padding: 1rem 2rem;
        color: white;
        border-radius: 10px;
        margin-bottom: 1.5rem;
        display: flex;
        justify-content: space-between;
        align-items: center;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }
    
    /* Cartões para funcionalidades */
    .card-container {
        display: flex;
        flex-wrap: wrap;
        gap: 1rem;
        margin-bottom: 1.5rem;
    }
    
    .feature-card {
        background-color: white;
        border-radius: 10px;
        padding: 1.2rem;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
        flex: 1;
        min-width: 200px;
        border: 1px solid #E2E8F0;
    }
    
    .feature-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.1);
    }
    
    .card-purple {
        border-left: 4px solid var(--purple);
    }
    
    .card-orange {
        border-left: 4px solid var(--orange);
    }
    
    .card-teal {
        border-left: 4px solid var(--teal);
    }

    /* Formulários */
    .form-container {
        background-color: white;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }
    
    .form-title {
        font-size: 1.5rem;
        color: var(--dark-purple);
        margin-bottom: 1rem;
    }
    
    /* Campos de entrada */
    .stTextInput > div > div > input {
        border-radius: 8px;
        border: 1px solid #E2E8F0;
        padding: 10px 15px;
        background-color: var(--background);
    }
    
    .stTextInput > div > div > input:focus {
        border-color: var(--purple);
        box-shadow: 0 0 0 2px rgba(98, 71, 170, 0.2);
    }
    
    .stTextArea > div > div > textarea {
        border-radius: 8px;
        border: 1px solid #E2E8F0;
        padding: 10px 15px;
        background-color: var(--background);
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: var(--purple);
        box-shadow: 0 0 0 2px rgba(98, 71, 170, 0.2);
    }
    
    /* Botões */
    .button-primary {
        background-color: var(--purple);
        color: white;
        border: none;
        border-radius: 50px;
        padding: 0.6rem 1.5rem;
        font-weight: 500;
        cursor: pointer;
        transition: all 0.2s;
        width: 100%;
        text-align: center;
        display: inline-block;
    }
    
    .button-secondary {
        background-color: #f0f0f0;
        color: var(--text-primary);
        border: 1px solid #ddd;
        border-radius: 50px;
        padding: 0.6rem 1.5rem;
        font-weight: 500;
        cursor: pointer;
        transition: all 0.2s;
        width: 100%;
        text-align: center;
        display: inline-block;
    }
    
    .button-primary:hover {
        filter: brightness(1.1);
        box-shadow: 0 4px 10px rgba(0,0,0,0.15);
    }
    
    .button-secondary:hover {
        background-color: #e0e0e0;
    }
    
    /* Resultado e área de resposta */
    .result-area {
        background-color: white;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border: 1px solid #E2E8F0;
    }
    
    /* Sidebar */
    .sidebar-content {
        padding: 20px 10px;
    }
    
    /* Métricas e cartões */
    .metric-card {
        background-color: #FFFFFF;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    
    /* Análise de tom */
    .tone-analysis-section {
        background-color: #FFFFFF;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border: 1px solid #E2E8F0;
    }
    
    .tone-current {
        background-color: var(--background);
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    
    .tone-optimized {
        background-color: #F0FFF4;
        padding: 10px;
        border-radius: 5px;
        border-left: 3px solid var(--teal);
    }

    /* Design responsivo */
    @media (max-width: 768px) {
        .header-gradient {
            padding: 1rem;
        }
        
        .feature-card {
            min-width: 100%;
        }
        
        .card-container {
            flex-direction: column;
        }
    }
    
    /* Esconder elementos default do Streamlit */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* Estilo aprimorado para o histórico */
    .history-item {
        padding: 10px;
        border-radius: 8px;
        margin-bottom: 10px;
        background-color: white;
        border: 1px solid #E2E8F0;
    }
    
    .history-item:hover {
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
    }
    
    /* Estilo para cenários relevantes */
    .scenario-card {
        background-color: #F9F9FB;
        border-radius: 8px;
        padding: 12px;
        margin-top: 8px;
        margin-bottom: 12px;
        border-left: 3px solid var(--teal);
    }
    
    .scenario-title {
        font-weight: 600;
        color: var(--dark-purple);
        font-size: 14px;
        margin-bottom: 5px;
    }
    
    .scenario-content {
        font-size: 13px;
        color: var(--text-secondary);
    }
</style>
""", unsafe_allow_html=True)

# Definição de funcionalidades disponíveis (incluindo PMBOK 7)
feature_options = {
    "Gerador de Comunicações Estruturadas": {
        "description": "Crie comunicações profissionais a partir de pontos-chave.",
        "icon": "📧",
        "color": "purple",
        "subtypes": ["E-mail Profissional", "Relatório de Status", "Comunicado Formal"]
    },
    "Assistente de Reuniões": {
        "description": "Gere agendas detalhadas, atas de reuniões e mensagens de follow-up estruturadas.",
        "icon": "📅",
        "color": "orange",
        "subtypes": ["Agenda de Reunião", "Ata/Resumo de Reunião", "Follow-up de Reunião"]
    },
    "Simplificador de Linguagem Técnica": {
        "description": "Simplifique linguagem técnica e adapte comunicações para diferentes públicos.",
        "icon": "🔄",
        "color": "teal",
        "subtypes": ["Simplificação de Documento Técnico", "Adaptação para Executivos", "Adaptação para Clientes", "Adaptação para Equipe Técnica"]
    },
    "Facilitador de Feedback": {
        "description": "Estruture feedback construtivo e prepare roteiros para conversas difíceis.",
        "icon": "💬",
        "color": "purple",
        "subtypes": ["Feedback de Desempenho", "Feedback sobre Entregáveis", "Roteiro para Conversa Difícil"]
    },
    "Detector de Riscos de Comunicação": {
        "description": "Analise comunicações para identificar ambiguidades e problemas potenciais.",
        "icon": "🔍",
        "color": "orange",
        "subtypes": ["Análise de E-mail", "Análise de Comunicado", "Análise de Proposta", "Análise de Documento de Requisitos"]
    },
    "Consultor PMBOK 7": {
        "description": "Esclareça dúvidas sobre gerenciamento de projetos conforme o PMBOK 7.",
        "icon": "📚",
        "color": "teal",
        "subtypes": ["Princípios de Gerenciamento", "Domínios de Performance", "Adaptação de Metodologias", "Ferramentas e Técnicas", "Melhores Práticas"]
    }
}

# ================= FUNÇÕES PARA INTEGRAÇÃO DOS ARQUIVOS JSON =================

# Função para carregar os arquivos JSON de cenários do GitHub
@st.cache_data
def load_scenarios_from_github():
    """
    Carrega os arquivos JSON de cenários do GitHub.
    Retorna um dicionário com os cenários organizados por categoria.
    """
    scenarios_data = {
        "negociacao_escopo": {},
        "gerenciamento_expectativas": {},
        "comunicacao_mas_noticias": {},
        "facilitacao_decisoes": {},
        "falhas_comunicacao": {},
        "emocoes_fortes": {},
        # Novas categorias para J7-J10
        "comunicacao_intercultural": {},
        "comunicacao_remota": {},
        "negociacao_conflito": {},
        "comunicacao_executiva": {}
    }
    
    # Mapeamento de arquivos e suas chaves correspondentes
    files_map = {
        "J1.json": "negociacao_escopo",
        "J2.json": "gerenciamento_expectativas",
        "J3.json": "comunicacao_mas_noticias",
        "J4.json": "facilitacao_decisoes",
        "J5.json": "falhas_comunicacao",
        "J6.json": "emocoes_fortes",
        # Novos arquivos J7-J10
        "J7.json": "comunicacao_intercultural",
        "J8.json": "comunicacao_remota",
        "J9.json": "negociacao_conflito",
        "J10.json": "comunicacao_executiva"
    }
    
    # Base URL para a pasta knowledge_base no GitHub
    base_url = "https://raw.githubusercontent.com/aiagentnexus25/nexus-assistant/main/knowledge_base/"
    
    # Carrega cada arquivo do GitHub
    for file_name, key in files_map.items():
        try:
            file_url = base_url + file_name
            response = requests.get(file_url)
            response.raise_for_status()  # Verifica se a requisição foi bem-sucedida
            
            scenarios_data[key] = json.loads(response.text)
            print(f"✅ {file_name} carregado com sucesso!")
        except requests.RequestException as e:
            print(f"⚠️ Não foi possível carregar {file_name} do GitHub. {str(e)}")
        except json.JSONDecodeError:
            print(f"❌ Erro ao decodificar {file_name}. Verifique o formato do arquivo.")
    
    return scenarios_data

# Função para extrair cenários relevantes com base nas entradas do usuário
def get_relevant_scenarios(scenarios_data, context, audience=None):
    """
    Retorna cenários relevantes com base no contexto e público.
    Utiliza todos os níveis de complexidade automaticamente.
    """
    all_relevant = []
    
    # Evitar erros se algum dos parâmetros for None
    context = context or ""
    audience = audience or ""
    
    # Palavras-chave do contexto e público
    context_keywords = [kw for kw in context.lower().split() if len(kw) > 3]
    audience_keywords = [kw for kw in audience.lower().split() if len(kw) > 3]
    
    # Busca em todos os tipos de cenários
    for scenario_type, scenarios in scenarios_data.items():
        if not scenarios:
            continue
            
        for scenario in scenarios.get("cenarios", []):
            # Verifica relevância do contexto
            context_match = False
            if not context_keywords:  # Se não houver contexto, consideramos que corresponde
                context_match = True
            else:
                for keyword in context_keywords:
                    if (keyword in scenario.get("contexto", "").lower() or 
                        keyword in scenario.get("dialogo", {}).get("situacao", "").lower()):
                        context_match = True
                        break
            
            # Verifica relevância do público
            audience_match = False
            if not audience_keywords:  # Se não houver público especificado, consideramos que corresponde
                audience_match = True
            else:
                for keyword in audience_keywords:
                    personagens = scenario.get("personagens", {})
                    if isinstance(personagens, dict) and any(keyword in persona.lower() for persona in personagens.keys()):
                        audience_match = True
                        break
            
            # Se ambos correspondem, adiciona o cenário à lista
            if context_match and audience_match:
                all_relevant.append({
                    "tipo": scenario_type,
                    "cenario": scenario
                })
    
    # Retorna até 3 cenários mais relevantes
    return all_relevant[:3] if all_relevant else []

# Função para enriquecer o prompt com cenários relevantes
def enrich_prompt_with_scenarios(base_prompt, relevant_scenarios):
    """
    Enriquecer um prompt existente com informações dos cenários relevantes,
    adaptando-se a diferentes estruturas de cenários
    """
    if not relevant_scenarios:
        return base_prompt
        
    # Mapeamento para nomes mais amigáveis dos tipos de cenário
    tipo_legivel = {
        "negociacao_escopo": "Negociação de Escopo",
        "gerenciamento_expectativas": "Gerenciamento de Expectativas",
        "comunicacao_mas_noticias": "Comunicação de Más Notícias",
        "facilitacao_decisoes": "Facilitação de Decisões",
        "falhas_comunicacao": "Mitigação de Falhas de Comunicação",
        "emocoes_fortes": "Lidar com Emoções Fortes",
        # Novos tipos para J7-J10
        "comunicacao_intercultural": "Comunicação Intercultural",
        "comunicacao_remota": "Comunicação em Equipes Remotas",
        "negociacao_conflito": "Negociação e Resolução de Conflitos",
        "comunicacao_executiva": "Comunicação Executiva"
    }
    
    enhanced_prompt = base_prompt + "\n\n--- CENÁRIOS DE REFERÊNCIA ---\n"
    
    for i, scenario_data in enumerate(relevant_scenarios, 1):
        scenario = scenario_data["cenario"]
        scenario_type = scenario_data["tipo"]
        
        enhanced_prompt += f"\nCenário {i} ({tipo_legivel.get(scenario_type, scenario_type)}):\n"
        
        # Adicionar contexto (verificando diferentes campos possíveis)
        contexto = scenario.get('contexto', scenario.get('descricao', scenario.get('situacao_contexto', 'N/A')))
        enhanced_prompt += f"Contexto: {contexto}\n"
        
        # Adicionar nível (verificando diferentes campos possíveis)
        nivel = scenario.get('nivel', scenario.get('complexidade', scenario.get('dificuldade', 'N/A')))
        enhanced_prompt += f"Nível: {nivel}\n"
        
        # Adicionar situação (verificando diferentes estruturas possíveis)
        situacao = None
        if "dialogo" in scenario and "situacao" in scenario["dialogo"]:
            situacao = scenario["dialogo"]["situacao"]
        elif "situacao" in scenario:
            situacao = scenario["situacao"]
        elif "descricao_situacao" in scenario:
            situacao = scenario["descricao_situacao"]
            
        if situacao:
            enhanced_prompt += f"Situação: {situacao}\n"
        
        # Adicionar estratégias (verificando diferentes estruturas possíveis)
        estrategias = []
        
        # Verifica em analise_emocional
        if "analise_emocional" in scenario and "tecnicas_recomendadas" in scenario["analise_emocional"]:
            estrategias = scenario["analise_emocional"]["tecnicas_recomendadas"]
        
        # Verifica em recomendacoes
        elif "recomendacoes" in scenario:
            if isinstance(scenario["recomendacoes"], list):
                estrategias = scenario["recomendacoes"]
            elif isinstance(scenario["recomendacoes"], dict) and "tecnicas" in scenario["recomendacoes"]:
                estrategias = scenario["recomendacoes"]["tecnicas"]
        
        # Verifica em estrategias
        elif "estrategias" in scenario:
            if isinstance(scenario["estrategias"], list):
                estrategias = scenario["estrategias"]
            elif isinstance(scenario["estrategias"], dict) and "recomendadas" in scenario["estrategias"]:
                estrategias = scenario["estrategias"]["recomendadas"]
        
        # Verifica em melhores_praticas
        elif "melhores_praticas" in scenario:
            if isinstance(scenario["melhores_praticas"], list):
                estrategias = scenario["melhores_praticas"]
                
        if estrategias:
            enhanced_prompt += "Estratégias recomendadas:\n"
            for estrategia in estrategias:
                enhanced_prompt += f"- {estrategia}\n"
                
        enhanced_prompt += "\n"
    
    enhanced_prompt += "\nUse os cenários acima como referência para criar uma comunicação mais eficaz e contextualizada, adaptando as estratégias recomendadas para o contexto específico descrito pelo usuário."
    
    return enhanced_prompt

# Função para mostrar os cenários relevantes na interface
def display_relevant_scenarios(relevant_scenarios):
    """
    Exibe os cenários relevantes na interface do usuário,
    adaptando-se a diferentes estruturas de cenários
    """
    if not relevant_scenarios:
        return
        
    # Mapeamento para nomes mais amigáveis dos tipos de cenário
    tipo_legivel = {
        "negociacao_escopo": "Negociação de Escopo",
        "gerenciamento_expectativas": "Gerenciamento de Expectativas",
        "comunicacao_mas_noticias": "Comunicação de Más Notícias",
        "facilitacao_decisoes": "Facilitação de Decisões",
        "falhas_comunicacao": "Mitigação de Falhas de Comunicação",
        "emocoes_fortes": "Lidar com Emoções Fortes",
        # Novos tipos para J7-J10
        "comunicacao_intercultural": "Comunicação Intercultural",
        "comunicacao_remota": "Comunicação em Equipes Remotas",
        "negociacao_conflito": "Negociação e Resolução de Conflitos",
        "comunicacao_executiva": "Comunicação Executiva"
    }
    
    with st.expander("📖 Cenários de referência relevantes", expanded=False):
        for i, scenario_data in enumerate(relevant_scenarios, 1):
            scenario = scenario_data["cenario"]
            scenario_type = scenario_data["tipo"]
            
            # Extrai contexto e nível com resiliência a diferentes estruturas
            contexto = scenario.get('contexto', 
                        scenario.get('descricao', 
                        scenario.get('situacao_contexto', 'N/A')))
            
            nivel = scenario.get('nivel', 
                    scenario.get('complexidade', 
                    scenario.get('dificuldade', 'N/A')))
            
            # Adiciona um terceiro campo informativo quando disponível
            extra_info = ""
            if "dialogo" in scenario and "situacao" in scenario["dialogo"]:
                extra_info = f"<strong>Situação:</strong> {scenario['dialogo']['situacao'][:100]}..."
            elif "situacao" in scenario:
                extra_info = f"<strong>Situação:</strong> {scenario['situacao'][:100]}..."
            elif "estrategias" in scenario and isinstance(scenario["estrategias"], list) and len(scenario["estrategias"]) > 0:
                extra_info = f"<strong>Estratégia principal:</strong> {scenario['estrategias'][0]}"
            elif "melhores_praticas" in scenario and isinstance(scenario["melhores_praticas"], list) and len(scenario["melhores_praticas"]) > 0:
                extra_info = f"<strong>Melhor prática:</strong> {scenario['melhores_praticas'][0]}"
            
            # Adiciona o campo extra se existir
            extra_html = f"<br>{extra_info}" if extra_info else ""
            
            st.markdown(f"""
            <div class="scenario-card">
                <div class="scenario-title">Cenário {i}: {tipo_legivel.get(scenario_type, scenario_type)}</div>
                <div class="scenario-content">
                    <strong>Contexto:</strong> {contexto}<br>
                    <strong>Nível:</strong> {nivel}{extra_html}
                </div>
            </div>
            """, unsafe_allow_html=True)

# ================= SESSION STATE INITIALIZATION =================

# Inicialização da sessão
if 'api_key_configured' not in st.session_state:
    st.session_state.api_key_configured = True
    # Configurar API key automaticamente a partir dos secrets
    if st.secrets.get("OPENAI_API_KEY"):
        st.session_state.api_key = st.secrets.get("OPENAI_API_KEY")
    else:
        st.session_state.api_key = None
        st.session_state.api_key_configured = False

if 'usage_data' not in st.session_state:
    st.session_state.usage_data = []
if 'generated_content' not in st.session_state:
    st.session_state.generated_content = ""
if 'history' not in st.session_state:
    st.session_state.history = []
if 'token_count' not in st.session_state:
    st.session_state.token_count = 0
if 'request_count' not in st.session_state:
    st.session_state.request_count = 0
if 'last_request_time' not in st.session_state:
    st.session_state.last_request_time = 0
if 'session_id' not in st.session_state:
    st.session_state.session_id = datetime.now().strftime("%Y%m%d%H%M%S") + str(hash(datetime.now()))[:5]
if 'current_feature' not in st.session_state:
    st.session_state.current_feature = ""
if 'optimized_content' not in st.session_state:
    st.session_state.optimized_content = ""
if 'previous_screen' not in st.session_state:
    st.session_state.previous_screen = None
# Nova variável para armazenar os cenários carregados
if 'scenarios_data' not in st.session_state:
    st.session_state.scenarios_data = None
if 'relevant_scenarios' not in st.session_state:
    st.session_state.relevant_scenarios = []

# Carrega os cenários se ainda não foram carregados
if st.session_state.scenarios_data is None:
    st.session_state.scenarios_data = load_scenarios_from_github()

# ================= HELPER FUNCTIONS =================

def header():
    """Renderiza o cabeçalho com gradiente apenas com o título NEXUS"""
    
    st.markdown("""
    <div class="header-gradient">
        <h1 style="margin:0; font-weight:600; font-size:32px; color:white;">NEXUS</h1>
        <div style="display: flex; align-items: center;">
            <div style="display: flex; align-items: center; gap: 8px;">
                <div style="width:8px; height:8px; border-radius:50%; background:#28C840;"></div>
                <span style="font-size:12px; color:white; opacity:0.9;">API</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def enrich_pmbok_prompt(prompt, pmbok_topic):
    """Enriquece o prompt com informações relevantes do PMBOK 7 baseado no tópico selecionado"""
    
    additional_info = ""
    
    if "Princípios" in pmbok_topic:
        additional_info += "\n\nPrincípios de Gerenciamento do PMBOK 7:\n"
        additional_info += "\n".join(PMBOK7_KNOWLEDGE_BASE["princípios"])
        
    elif "Domínios" in pmbok_topic:
        additional_info += "\n\nDomínios de Performance do PMBOK 7:\n"
        additional_info += "\n".join(PMBOK7_KNOWLEDGE_BASE["domínios"])
        
    elif "Adaptação" in pmbok_topic:
        additional_info += "\n\nAbordagens de Desenvolvimento no PMBOK 7:\n"
        for k, v in PMBOK7_KNOWLEDGE_BASE["metodologias"].items():
            additional_info += f"- {k.capitalize()}: {v}\n"
            
    elif "Melhores Práticas" in pmbok_topic:
        additional_info += "\n\nMudanças Principais no PMBOK 7:\n"
        additional_info += "\n".join([f"- {item}" for item in PMBOK7_KNOWLEDGE_BASE["mudancas_principais"]])
    
    # Adicionar a informação relevante ao prompt
    enhanced_prompt = prompt + additional_info
    return enhanced_prompt

# Função para exportar conteúdo como DOCX
def export_as_docx(content, filename="documento"):
    doc = docx.Document()
    
    # Adicionar título
    doc.add_heading(f"{filename}", 0)
    
    # Dividir por linhas e adicionar parágrafos
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip() == "":
            continue
        
        # Verificar se é um cabeçalho
        if re.match(r'^#{1,6}\s+', para):
            # Extrair o nível do cabeçalho e o texto
            header_match = re.match(r'^(#{1,6})\s+(.*)', para)
            if header_match:
                level = min(len(header_match.group(1)), 9)  # Limitar a 9 para evitar erro
                text = header_match.group(2)
                doc.add_heading(text, level)
        else:
            doc.add_paragraph(para)
    
    # Salvar para um buffer em memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Função para gerar conteúdo via API OpenAI (modificada para usar cenários)
def generate_content(prompt, model="gpt-3.5-turbo", temperature=0.7, use_scenarios=True):
    if not st.session_state.api_key_configured or not st.session_state.api_key:
        return "API não configurada. Por favor, contate o administrador."
    
    # Verificar limites
    if st.session_state.token_count >= TOKEN_LIMIT:
        return "Você atingiu o limite de tokens para esta sessão. Por favor, tente novamente mais tarde."
    
    if st.session_state.request_count >= REQUEST_LIMIT:
        return "Você atingiu o limite de requisições para esta sessão. Por favor, tente novamente mais tarde."
    
    try:
        with st.spinner("Gerando conteúdo..."):
            # Atualizar o timestamp da última requisição
            st.session_state.last_request_time = time.time()
            
            # Incrementar contador de requisições
            st.session_state.request_count += 1
            
            # Configurar requisição direta à API
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {st.session_state.api_key}"
            }
            
            # System prompt melhorado para considerar os cenários
            system_prompt = """
            Você é o NEXUS, um sistema de IA especializado em comunicação estratégica e gerenciamento de projetos.
            Forneça respostas profissionais, estruturadas e detalhadas, adaptadas ao contexto específico.
            Quando cenários de referência forem fornecidos, analise-os cuidadosamente e incorpore as estratégias 
            de comunicação mais relevantes às suas respostas.
            """
            
            # Adicionar mensagem do sistema e prompt do usuário
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                "temperature": temperature,
                "max_tokens": 4000  # Aumentado para respostas mais completas
            }
            
            # Fazer a requisição à API
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                data=json.dumps(payload),
                timeout=60
            )
            
            # Processar a resposta
            if response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                # Atualizar contadores de tokens
                prompt_tokens = result['usage']['prompt_tokens']
                completion_tokens = result['usage']['completion_tokens']
                total_tokens = result['usage']['total_tokens']
                st.session_state.token_count += total_tokens
                
                # Registrar uso
                st.session_state.usage_data.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'feature': st.session_state.current_feature,
                    'tokens': total_tokens,
                    'model': model,
                    'session_id': st.session_state.session_id
                })
                
                # Adicionar ao histórico
                st.session_state.history.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'feature': st.session_state.current_feature,
                    'input': prompt[:100] + "..." if len(prompt) > 100 else prompt,  # Resumido para economizar espaço
                    'output': content,
                    'model': model,
                    'session_id': st.session_state.session_id
                })
                
                return content
            else:
                return f"Erro na API (Status {response.status_code}): {response.text}"
        
    except Exception as e:
        return f"Erro ao gerar conteúdo: {str(e)}"

# Função para criar cartões de funcionalidades
def create_feature_cards():
    """Cria os cartões de seleção de funcionalidades na interface principal"""
    
    col1, col2 = st.columns(2)
    features_list = list(feature_options.items())
    
    # Distribuir as funcionalidades em duas colunas
    for i, (name, details) in enumerate(features_list):
        current_col = col1 if i % 2 == 0 else col2
        with current_col:
            color = details["color"]
            icon = details["icon"]
            description = details["description"]
            
            # Use um container para o estilo do cartão
            with st.container():
                st.markdown(f"""
                <div style="background-color: white; padding: 15px; border-radius: 10px; 
                            margin-bottom: 15px; border-left: 4px solid var(--{color});">
                    <h3 style="margin:0; color:var(--dark-purple); font-size:18px;">
                        {icon} {name}
                    </h3>
                    <p style="margin:5px 0 0 0; color:var(--text-secondary); font-size:14px;">
                        {description}
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                # Simplificar para apenas um botão
                if st.button(f"Selecionar", key=f"select_{name}"):
                    st.session_state.current_feature = name
                    st.session_state.previous_screen = "home"
                    st.experimental_rerun()

# Função para configurar e exibir a sidebar
def setup_sidebar():
    """Configura a barra lateral do aplicativo"""
    
    with st.sidebar:
        st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
        
        # Logo ou imagem
        st.markdown("""
        <div style="text-align:center; padding:20px 0;">
            <svg width="120" height="40">
                <defs>
                    <linearGradient id="sidebarGradient" x1="0%" y1="0%" x2="100%" y2="0%">
                        <stop offset="0%" stop-color="#6247AA"/>
                        <stop offset="50%" stop-color="#FF6D2A"/>
                        <stop offset="100%" stop-color="#00C1D5"/>
                    </linearGradient>
                </defs>
                <rect width="120" height="40" rx="5" fill="url(#sidebarGradient)"/>
                <text x="60" y="25" font-family="Arial, sans-serif" font-size="18" font-weight="bold" fill="white" text-anchor="middle">NEXUS</text>
            </svg>
            <p style="color:#86868B; margin-top:5px; font-size:12px;">Assistente de Comunicação IA</p>
        </div>
        """, unsafe_allow_html=True)

# Interface principal do aplicativo
def main():
    # Renderizar o cabeçalho com gradiente
    header()
    
    # Configurar a sidebar
    setup_sidebar()
    
    # Exibir mensagem sobre versão aprimorada
    with st.expander("ℹ️ Sobre esta versão", expanded=False):
        st.markdown(f"""
        **Versão NEXUS Aprimorada**
        Esta versão possui interface moderna, limites expandidos e base de conhecimento de cenários:
        * Até {REQUEST_LIMIT} requisições por sessão
        * Até {TOKEN_LIMIT} tokens por sessão
        * Design responsivo para todos os dispositivos 
        * Base de conhecimento com mais de 600 cenários de comunicação
        * Sugestões contextualizadas baseadas em casos reais
        """)
    
    # Botão VOLTAR quando estiver em uma funcionalidade
    if st.session_state.current_feature:
        if st.button("◀️ VOLTAR", key="back_to_home"):
            st.session_state.current_feature = ""
            st.session_state.previous_screen = None
            st.experimental_rerun()
    
    # Histórico de gerações recentes
    if st.session_state.history:
        with st.expander("Histórico de Gerações Recentes", expanded=False):
            for i, item in enumerate(reversed(st.session_state.history[-5:])):  # 5 itens mais recentes
                st.markdown(f"""
                <div class="history-item">
                    <p><strong>{item['timestamp']} - {item['feature']}</strong></p>
                </div>
                """, unsafe_allow_html=True)
                if st.button(f"Carregar este conteúdo ↩️", key=f"load_{i}"):
                    st.session_state.current_feature = item['feature']
                    st.session_state.generated_content = item['output']
                    st.experimental_rerun()
                st.markdown("---")
    
    # Se nenhuma funcionalidade selecionada, mostrar as opções
    if not st.session_state.current_feature:
        st.markdown("## Selecione uma funcionalidade")
        create_feature_cards()
    else:
        # Interface específica da funcionalidade selecionada
        current_feature = st.session_state.current_feature
        feature_details = feature_options[current_feature]
        
        # Título da funcionalidade
        st.markdown(f"## {feature_details['icon']} {current_feature}")
        
        # Verificar limites antes de mostrar a interface
        if st.session_state.token_count >= TOKEN_LIMIT:
            st.error(f"Você atingiu o limite de {TOKEN_LIMIT} tokens para esta sessão. Por favor, tente novamente mais tarde.")
        elif st.session_state.request_count >= REQUEST_LIMIT:
            st.error(f"Você atingiu o limite de {REQUEST_LIMIT} requisições para esta sessão. Por favor, tente novamente mais tarde.")
        else:
            # Interface específica da funcionalidade
            with st.form(key=f"{current_feature}_form"):
                st.markdown(f'<div class="form-container">', unsafe_allow_html=True)
                st.markdown(f'<p class="form-title">{feature_details["description"]}</p>', unsafe_allow_html=True)
                
                # Campo de subtipo
                subtype = st.selectbox("Tipo de Comunicação", feature_details['subtypes'])
                
                # Campos comuns a todas as funcionalidades (exceto PMBOK que tem campos específicos)
                context = ""
                if current_feature != "Consultor PMBOK 7":
                    context = st.text_area("Contexto do Projeto", 
                                        help="Descreva o projeto, fase atual e informações relevantes",
                                        height=100,
                                        placeholder="Ex: Projeto de desenvolvimento do aplicativo mobile, fase de testes")
              
               
                # Campos específicos por funcionalidade
                prompt = ""
                audience = None  # Para uso nos cenários relevantes
                
                if current_feature == "Gerador de Comunicações Estruturadas":
                    audience = st.text_input("Público-alvo", 
                                        help="Para quem esta comunicação será enviada (equipe, cliente, stakeholder)",
                                        placeholder="Ex: Cliente, diretor de marketing da empresa XYZ")
                    key_points = st.text_area("Pontos-chave", 
                                            help="Liste os principais pontos que devem ser incluídos na comunicação",
                                            height=150,
                                            placeholder="Ex: Atraso de 3 dias devido a bugs na integração; Plano de recuperação com recursos adicionais")
                    tone = st.select_slider("Tom da Comunicação", 
                                        options=["Muito Formal", "Formal", "Neutro", "Amigável", "Casual"],
                                        value="Neutro")
                    
                    base_prompt = f"""
                    Gere um {subtype} com base nas seguintes informações:
                    
                    Contexto do Projeto: {context}
                    Público-alvo: {audience}
                    Pontos-chave: {key_points}
                    Tom desejado: {tone}
                    
                    Formate a saída adequadamente para um {subtype}, incluindo assunto/título e estrutura apropriada.
                    """
                
                elif current_feature == "Assistente de Reuniões":
                    participants = st.text_area("Participantes", 
                                            help="Liste os participantes e suas funções",
                                            height=100,
                                            placeholder="Ex: João Silva (Gerente de Projeto), Maria Costa (Desenvolvedora Frontend)")
                    audience = participants  # Para uso nos cenários relevantes
                    topics = st.text_area("Tópicos a serem abordados", 
                                        help="Liste os tópicos que precisam ser discutidos",
                                        height=150,
                                        placeholder="Ex: Atualização do cronograma, Bugs pendentes, Feedback do cliente")
                    duration = st.number_input("Duração (minutos)", min_value=15, max_value=240, value=60, step=15)
                    
                    if subtype == "Agenda de Reunião":
                        base_prompt = f"""
                        Crie uma agenda detalhada para uma reunião de {duration} minutos com base nas seguintes informações:
                        
                        Contexto do Projeto: {context}
                        Participantes: {participants}
                        Tópicos a serem abordados: {topics}
                        
                        Inclua alocação de tempo para cada item, responsáveis e objetivos claros.
                        """
                    elif subtype == "Ata/Resumo de Reunião":
                        decisions = st.text_area("Decisões tomadas", 
                                            help="Liste as principais decisões tomadas durante a reunião",
                                            height=100,
                                            placeholder="Ex: Aprovação do novo design, Extensão do prazo em 1 semana")
                        actions = st.text_area("Ações acordadas", 
                                            help="Liste as ações acordadas, responsáveis e prazos",
                                            height=100,
                                            placeholder="Ex: João irá corrigir o bug #123 até amanhã, Maria criará novos componentes até sexta")
                        
                        base_prompt = f"""
                        Crie uma ata/resumo detalhado para uma reunião de {duration} minutos com base nas seguintes informações:
                        
                        Contexto do Projeto: {context}
                        Participantes: {participants}
                        Tópicos abordados: {topics}
                        Decisões tomadas: {decisions}
                        Ações acordadas: {actions}
                        
                        Organize por tópicos, destacando claramente decisões e próximos passos com responsáveis.
                        """
                    else:  # Follow-up
                        meeting_outcome = st.text_area("Resultado da reunião", 
                                                help="Resuma os principais resultados da reunião",
                                                height=100,
                                                placeholder="Ex: Definidas as prioridades para o próximo sprint e resolvidos os bloqueios atuais")
                        action_items = st.text_area("Itens de ação", 
                                                help="Liste os itens de ação, responsáveis e prazos",
                                                height=100,
                                                placeholder="Ex: João: revisão de código até 25/03; Maria: implementação da nova feature até 27/03")
                        
                        base_prompt = f"""
                        Crie uma mensagem de follow-up para uma reunião com base nas seguintes informações:
                        
                        Contexto do Projeto: {context}
                        Participantes: {participants}
                        Tópicos abordados: {topics}
                        Resultado da reunião: {meeting_outcome}
                        Itens de ação: {action_items}
                        
                        A mensagem deve agradecer a participação, resumir os principais pontos, detalhar próximos passos
                        com responsáveis e prazos, e solicitar confirmação/feedback conforme apropriado.
                        """
                        
                elif current_feature == "Simplificador de Linguagem Técnica":
                    technical_content = st.text_area("Conteúdo Técnico", 
                                                help="Cole aqui o texto técnico que precisa ser traduzido",
                                                height=200,
                                                placeholder="Ex: A implementação do Redux utiliza reducers imutáveis para gerenciar o estado global da aplicação...")
                    audience = st.selectbox("Público-alvo", 
                                        ["Executivos", "Clientes não-técnicos", "Equipe de Negócios", "Equipe Técnica Junior"])
                    key_concepts = st.text_input("Conceitos-chave a preservar", 
                                            help="Liste conceitos técnicos que devem ser mantidos mesmo se simplificados",
                                            placeholder="Ex: gerenciamento de estado, API, front-end")
                    
                    base_prompt = f"""
                    Traduza/adapte o seguinte conteúdo técnico para um público de {audience} com base nas seguintes informações:
                    
                    Contexto do Projeto: {context}
                    Conteúdo Técnico Original: {technical_content}
                    Conceitos-chave a preservar: {key_concepts}
                    
                    Para {audience}, foque em: 
                    - {'Impacto nos negócios e resultados de alto nível' if audience == 'Executivos' else ''}
                    - {'Benefícios e funcionalidades em linguagem acessível' if audience == 'Clientes não-técnicos' else ''}
                    - {'Conexão com objetivos de negócios e processos' if audience == 'Equipe de Negócios' else ''}
                    - {'Explicações técnicas mais detalhadas, mas com conceitos explicados' if audience == 'Equipe Técnica Junior' else ''}
                    
                    Mantenha a precisão conceitual mesmo simplificando a linguagem.
                    Forneça uma explicação completa e detalhada, com exemplos e analogias apropriadas para o público.
                    """
                
                elif current_feature == "Facilitador de Feedback":
                    situation = st.text_area("Situação", 
                                        help="Descreva a situação específica para a qual você precisa fornecer feedback",
                                        height=150,
                                        placeholder="Ex: Atraso na entrega de componentes para o projeto principal...")
                    strengths = st.text_area("Pontos Fortes", 
                                        help="Liste aspectos positivos que devem ser destacados",
                                        height=100,
                                        placeholder="Ex: Qualidade do código entregue, comunicação proativa de desafios")
                    areas_for_improvement = st.text_area("Áreas para Melhoria", 
                                                    help="Liste aspectos que precisam ser melhorados",
                                                    height=100,
                                                    placeholder="Ex: Estimativas de tempo irrealistas, falha em pedir ajuda quando bloqueado")
                    relationship = st.selectbox("Relação com o Receptor", 
                                            ["Membro da equipe direto", "Colega de mesmo nível", "Superior hierárquico", "Cliente", "Fornecedor"])
                    audience = relationship  # Para uso nos cenários relevantes
                    
                    base_prompt = f"""
                    Estruture um {subtype} construtivo e eficaz com base nas seguintes informações:
                    
                    Contexto do Projeto: {context}
                    Situação específica: {situation}
                    Pontos fortes a destacar: {strengths}
                    Áreas para melhoria: {areas_for_improvement}
                    Relação com o receptor: {relationship}
                    
                    O feedback deve:
                    - Ser específico e baseado em comportamentos observáveis
                    - Equilibrar aspectos positivos e áreas de melhoria
                    - Incluir exemplos concretos
                    - Oferecer sugestões acionáveis
                    - Usar tom apropriado para a relação ({relationship})
                    - Focar em crescimento e desenvolvimento, não em crítica
                    
                    Formate como um roteiro/script detalhado que o usuário pode seguir na conversa ou adaptar para uma comunicação escrita.
                    Adicione observações e dicas de comunicação não-verbal quando relevante.
                    """
                
                elif current_feature == "Detector de Riscos de Comunicação":
                    content_to_analyze = st.text_area("Conteúdo para Análise", 
                                                    help="Cole aqui o texto que você deseja analisar quanto a riscos de comunicação",
                                                    height=200,
                                                    placeholder="Ex: Devido a circunstâncias imprevistas no desenvolvimento, alguns recursos podem sofrer atrasos...")
                    audience = st.text_input("Público-alvo", 
                                        help="Descreva quem receberá esta comunicação",
                                        placeholder="Ex: Cliente executivo com pouco conhecimento técnico")
                    stakes = st.select_slider("Importância da Comunicação", 
                                            options=["Baixa", "Média", "Alta", "Crítica"],
                                            value="Média")
                    
                    base_prompt = f"""
                    Analise o seguinte {subtype} quanto a riscos de comunicação:
                    
                    Contexto do Projeto: {context}
                    Público-alvo: {audience}
                    Importância da comunicação: {stakes}
                    
                    Conteúdo para análise:
                    ---
                    {content_to_analyze}
                    ---
                    
                    Sua análise deve:
                    1. Identificar ambiguidades, informações incompletas ou confusas
                    2. Apontar possíveis mal-entendidos baseados no público-alvo
                    3. Detectar problemas de tom ou linguagem inapropriada
                    4. Identificar informações sensíveis ou potencialmente problemáticas
                    5. Sugerir reformulações específicas para cada problema identificado
                    6. Analisar a estrutura geral e propor melhorias organizacionais
                    7. Verificar se há informações críticas ausentes
                    
                    Organize sua análise em forma de tabela com colunas para: Trecho problemático, Risco potencial, Sugestão de melhoria.
                    Ao final, forneça uma avaliação geral dos riscos de comunicação (Baixo/Médio/Alto) e um resumo das principais recomendações.
                    Forneça também uma versão revisada completa do texto.
                    """
                
                elif current_feature == "Consultor PMBOK 7":
                    pmbok_topic = subtype  # Já definido pelo selectbox de subtypes
                    
                    project_context = st.text_area("Contexto do Projeto", 
                                            help="Descreva brevemente o projeto ou a situação para contextualizar sua dúvida",
                                            height=100,
                                            placeholder="Ex: Estamos iniciando um projeto de desenvolvimento de software com metodologia híbrida...")
                    
                    specific_question = st.text_area("Sua Dúvida Específica", 
                                                help="Detalhe sua dúvida ou o que você precisa saber sobre o tema selecionado",
                                                height=150,
                                                placeholder="Ex: Como aplicar os princípios de valor do PMBOK 7 em um ambiente que tradicionalmente usava metodologias em cascata?")
                    
                    experience_level = st.select_slider("Seu Nível de Experiência", 
                                                    options=["Iniciante", "Intermediário", "Avançado", "Especialista"],
                                                    value="Intermediário")
                    
                    organization_context = st.text_input("Contexto Organizacional", 
                                                    help="Descreva brevemente o contexto organizacional (opcional)",
                                                    placeholder="Ex: Empresa de médio porte do setor financeiro com cultura tradicional")
                    
                    base_prompt = f"""
                    Forneça uma orientação detalhada sobre o tema "{pmbok_topic}" do PMBOK 7 com base nas seguintes informações:
                    
                    Contexto do Projeto: {project_context}
                    Dúvida Específica: {specific_question}
                    Nível de Experiência do Usuário: {experience_level}
                    Contexto Organizacional: {organization_context}
                    
                    Sua resposta deve:
                    1. Explicar os conceitos relevantes do PMBOK 7 relacionados à dúvida
                    2. Fornecer orientações práticas adaptadas ao contexto específico
                    3. Apresentar exemplos concretos de aplicação
                    4. Destacar boas práticas e recomendações
                    5. Considerar o nível de experiência do usuário ({experience_level})
                    6. Fazer conexões com outros domínios ou princípios relevantes do PMBOK 7
                    7. Incluir dicas de implementação prática
                    8. Mencionar possíveis desafios e como superá-los
                    
                    Formate a resposta de maneira estruturada, com seções claras e, se apropriado, inclua referências aos elementos específicos do PMBOK 7.
                    """
                    
                    # Enriquecemos o prompt com informações relevantes do PMBOK 7
                    prompt = enrich_pmbok_prompt(base_prompt, pmbok_topic)
                else:
                    prompt = base_prompt
                
                # Para funcionalidades diferentes de PMBOK 7, enriquecer com cenários relevantes
                if current_feature != "Consultor PMBOK 7":
                    # Buscar cenários relevantes
                    relevant_scenarios = get_relevant_scenarios(
                        st.session_state.scenarios_data, 
                        context, 
                        audience,                        
                    )
                    
                    # Armazenar para exibição após o formulário
                    st.session_state.relevant_scenarios = relevant_scenarios
                    
                    # Enriquecer o prompt com informações dos cenários
                    prompt = enrich_prompt_with_scenarios(base_prompt, relevant_scenarios)
                
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Substituir os três botões por um único botão GERAR
                st.markdown("""
                <style>
                    /* Estilo para o botão submit roxo */
                   .stButton > button {
                        background-color: #6247AA !important;
                        color: white !important;
                        border: none !important;
                        border-radius: 4px;
                        padding: 0.5rem 1rem;
                        font-weight: 500;
                   }
                </style>
                """, unsafe_allow_html=True)
                
                submit_button = st.form_submit_button("GERAR")

                # Processamento após o envio do formulário
            if submit_button:
                if not st.session_state.api_key_configured:
                    st.error("API não configurada. Por favor, contate o administrador.")
                elif st.session_state.token_count >= TOKEN_LIMIT:
                    st.error(f"Você atingiu o limite de {TOKEN_LIMIT} tokens para esta sessão.")
                elif st.session_state.request_count >= REQUEST_LIMIT:
                    st.error(f"Você atingiu o limite de {REQUEST_LIMIT} requisições para esta sessão.")
                else:
                    # Exibir cenários relevantes (se houver)
                    if current_feature != "Consultor PMBOK 7" and st.session_state.relevant_scenarios:
                        display_relevant_scenarios(st.session_state.relevant_scenarios)
                    
                    # Gerar conteúdo
                    generated_content = generate_content(prompt, model="gpt-3.5-turbo", temperature=0.7)
                    st.session_state.generated_content = generated_content
                    
                    # Exibir resultado
                    st.markdown("### Resultado")
                    st.markdown('<div class="result-area">', unsafe_allow_html=True)
                    st.markdown(generated_content)
                    st.markdown('</div>', unsafe_allow_html=True)
                    
                    # Adicionar botão VOLTAR após exibir o resultado
                    col1, col2 = st.columns([1, 3])
                    with col1:
                        if st.button("◀️ VOLTAR", key="back_from_result"):
                            # Manter o current_feature, apenas limpar o resultado
                            st.session_state.generated_content = ""
                            st.experimental_rerun()
                    
                    # Opções de download
                    col1, col2 = st.columns(2)
                    with col1:
                        # Download como texto
                        st.download_button(
                            label="📄 Baixar como TXT",
                            data=generated_content,
                            file_name=f"{current_feature.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain"
                        )
                    
                    with col2:
                        # Download como DOCX
                        docx_buffer = export_as_docx(generated_content)
                        st.download_button(
                            label="📝 Baixar como DOCX",
                            data=docx_buffer,
                            file_name=f"{current_feature.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

# Iniciar a aplicação
if __name__ == "__main__":
    main()
