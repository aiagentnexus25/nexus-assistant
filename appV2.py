import streamlit as st
import requests
import json
import pandas as pd
from datetime import datetime
import plotly.express as px
import os
import re
import docx
from io import BytesIO
import time

# ================= CONFIGURATION =================

# Configuração da página
st.set_page_config(
    page_title="NEXUS - Assistente de Comunicação de Projetos",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Personalizado
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        text-align: center;
        margin-bottom: 1rem;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #3B82F6;
        text-align: center;
        margin-bottom: 2rem;
    }
    .feature-card {
        background-color: #F3F4F6;
        border-radius: 10px;
        padding: 20px;
        margin-bottom: 20px;
        border-left: 5px solid #3B82F6;
    }
    .result-area {
        background-color: #F0F9FF;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border: 1px solid #BAE6FD;
    }
    .sidebar-content {
        padding: 20px 10px;
    }
    .metric-card {
        background-color: #FAFAFA;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
    }
    .feedback-good {
        color: #059669;
        font-weight: bold;
    }
    .feedback-bad {
        color: #DC2626;
        font-weight: bold;
    }
    .stAlert {
        margin-top: 1rem;
        margin-bottom: 1rem;
    }
    .usage-info {
        background-color: #EFF6FF;
        border-radius: 8px;
        padding: 10px;
        margin: 10px 0;
        border: 1px solid #DBEAFE;
    }
    .tone-analysis-section {
        background-color: #F0FFF4;
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        border: 1px solid #C6F6D5;
    }
    .tone-current {
        background-color: #E6FFFA;
        padding: 10px;
        border-radius: 5px;
        margin-bottom: 10px;
    }
    .tone-optimized {
        background-color: #F0FFF4;
        padding: 10px;
        border-radius: 5px;
        border-left: 3px solid #38A169;
    }
</style>
""", unsafe_allow_html=True)

# ================= CONSTANTS =================

# Limites de uso restritivos para testes
TOKEN_LIMIT = 3000      # Limite total de tokens por sessão
REQUEST_LIMIT = 5       # Limite de solicitações por sessão
RATE_LIMIT_SECONDS = 60 # Tempo mínimo (em segundos) entre requisições

# System prompt base para o assistente
system_prompt = """
Você é o NEXUS, um assistente especializado em comunicação de projetos. Seu objetivo é ajudar gerentes de projetos, líderes de equipe e outros profissionais a comunicar-se de forma clara, eficaz e profissional em diversos contextos de projetos.

Você possui cinco habilidades principais:

1. Gerador de Comunicações Estruturadas: Criar e-mails profissionais, relatórios de status e comunicados formais.
2. Assistente de Reuniões: Gerar agendas detalhadas, atas e resumos de reuniões, e estruturar follow-ups.
3. Tradutor de Jargão Técnico: Simplificar linguagem técnica, adaptar comunicações para diferentes públicos e traduzir requisitos técnicos.
4. Facilitador de Feedback: Estruturar feedback construtivo, transformar críticas em sugestões e criar roteiros para conversas difíceis.
5. Detector de Riscos de Comunicação: Analisar comunicações, sugerir alternativas mais claras e avaliar adequação ao público.

Ao responder, você deve:
- Ser conciso mas completo
- Usar linguagem profissional e tom adequado para o contexto
- Estruturar o conteúdo de forma lógica e clara
- Focar em comunicação eficaz e construtiva
- Evitar jargões desnecessários, a menos que sejam apropriados para o público-alvo

Observação importante: Esta é uma versão de demonstração com limites de uso. Seja eficiente e direto em suas respostas.
"""

# Definição de funcionalidades disponíveis
feature_options = {
    "Gerador de Comunicações Estruturadas": {
        "description": "Crie e-mails profissionais, relatórios de status e comunicados formais a partir de pontos-chave.",
        "icon": "📧",
        "subtypes": ["E-mail Profissional", "Relatório de Status", "Comunicado Formal"]
    },
    "Assistente de Reuniões": {
        "description": "Gere agendas detalhadas, atas de reuniões e mensagens de follow-up estruturadas.",
        "icon": "📅",
        "subtypes": ["Agenda de Reunião", "Ata/Resumo de Reunião", "Follow-up de Reunião"]
    },
    "Tradutor de Jargão Técnico": {
        "description": "Simplifique linguagem técnica e adapte comunicações para diferentes públicos.",
        "icon": "🔄",
        "subtypes": ["Simplificação de Documento Técnico", "Adaptação para Executivos", "Adaptação para Clientes", "Adaptação para Equipe Técnica"]
    },
    "Facilitador de Feedback": {
        "description": "Estruture feedback construtivo e prepare roteiros para conversas difíceis.",
        "icon": "💬",
        "subtypes": ["Feedback de Desempenho", "Feedback sobre Entregáveis", "Roteiro para Conversa Difícil"]
    },
    "Detector de Riscos de Comunicação": {
        "description": "Analise comunicações para identificar ambiguidades e problemas potenciais.",
        "icon": "🔍",
        "subtypes": ["Análise de E-mail", "Análise de Comunicado", "Análise de Proposta", "Análise de Documento de Requisitos"]
    }
}

# ================= SESSION STATE INITIALIZATION =================

# Inicialização da sessão
if 'api_key_configured' not in st.session_state:
    st.session_state.api_key_configured = True
    # Configurar API key automaticamente a partir dos secrets
    if st.secrets.get("OPENAI_API_KEY"):
        st.session_state.api_key = st.secrets.get("OPENAI_API_KEY")
    else:
        st.session_state.api_key = None
        st.session_state.api_key_configured = False

if 'usage_data' not in st.session_state:
    st.session_state.usage_data = []
if 'generated_content' not in st.session_state:
    st.session_state.generated_content = ""
if 'history' not in st.session_state:
    st.session_state.history = []
if 'token_count' not in st.session_state:
    st.session_state.token_count = 0
if 'request_count' not in st.session_state:
    st.session_state.request_count = 0
if 'last_request_time' not in st.session_state:
    st.session_state.last_request_time = 0
if 'session_id' not in st.session_state:
    st.session_state.session_id = datetime.now().strftime("%Y%m%d%H%M%S") + str(hash(datetime.now()))[:5]
if 'current_feature' not in st.session_state:
    st.session_state.current_feature = ""
if 'optimized_content' not in st.session_state:
    st.session_state.optimized_content = ""

# ================= HELPER FUNCTIONS =================

# Função para gerar conteúdo via API OpenAI
def generate_content(prompt, model="gpt-3.5-turbo", temperature=0.7):
    if not st.session_state.api_key_configured or not st.session_state.api_key:
        return "API não configurada. Por favor, contate o administrador."
    
    # Verificar limites
    if st.session_state.token_count >= TOKEN_LIMIT:
        return "Você atingiu o limite de tokens para esta sessão. Por favor, tente novamente mais tarde."
    
    if st.session_state.request_count >= REQUEST_LIMIT:
        return "Você atingiu o limite de requisições para esta sessão. Por favor, tente novamente mais tarde."
    
    # Verificar limites de taxa (rate limits)
    current_time = time.time()
    if current_time - st.session_state.last_request_time < RATE_LIMIT_SECONDS and st.session_state.request_count > 0:
        wait_time = round(RATE_LIMIT_SECONDS - (current_time - st.session_state.last_request_time))
        return f"Por favor, aguarde {wait_time} segundos antes de fazer outra requisição."
    
    try:
        with st.spinner("Gerando conteúdo..."):
            # Atualizar o timestamp da última requisição
            st.session_state.last_request_time = current_time
            
            # Incrementar contador de requisições
            st.session_state.request_count += 1
            
            # Configurar requisição direta à API
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {st.session_state.api_key}"
            }
            
            # Adicionar mensagem do sistema e prompt do usuário
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt}
                ],
                "temperature": temperature,
                "max_tokens": 1500  # Limitado para economizar tokens
            }
            
            # Fazer a requisição à API
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                data=json.dumps(payload),
                timeout=60
            )
            
            # Processar a resposta
            if response.status_code == 200:
                result = response.json()
                content = result['choices'][0]['message']['content']
                
                # Atualizar contadores de tokens
                prompt_tokens = result['usage']['prompt_tokens']
                completion_tokens = result['usage']['completion_tokens']
                total_tokens = result['usage']['total_tokens']
                st.session_state.token_count += total_tokens
                
                # Registrar uso
                st.session_state.usage_data.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'feature': st.session_state.current_feature,
                    'tokens': total_tokens,
                    'model': model,
                    'session_id': st.session_state.session_id
                })
                
                # Adicionar ao histórico
                st.session_state.history.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'feature': st.session_state.current_feature,
                    'input': prompt[:100] + "..." if len(prompt) > 100 else prompt,  # Resumido para economizar espaço
                    'output': content,
                    'model': model,
                    'session_id': st.session_state.session_id
                })
                
                return content
            else:
                return f"Erro na API (Status {response.status_code}): {response.text}"
        
    except Exception as e:
        return f"Erro ao gerar conteúdo: {str(e)}"

# NOVA FUNÇÃO: Análise de Tom
def analyze_tone(text, target_audience, desired_impact):
    """
    Analisa o tom do texto e sugere melhorias baseadas no público-alvo
    e no impacto emocional desejado.
    
    Args:
        text (str): O texto a ser analisado
        target_audience (str): O público-alvo (cliente, equipe, executivos)
        desired_impact (str): O impacto desejado (tranquilizar, motivar, urgência)
        
    Returns:
        dict: Análise do tom e texto otimizado
    """
    if not st.session_state.api_key_configured or not st.session_state.api_key:
        return {
            'current_tone': "API não configurada",
            'emotional_impact': "Não foi possível analisar",
            'optimized_text': text
        }
    
    # Verificar limites
    if st.session_state.token_count >= TOKEN_LIMIT:
        return {
            'current_tone': "Limite de tokens atingido",
            'emotional_impact': "Não foi possível analisar",
            'optimized_text': text
        }
    
    if st.session_state.request_count >= REQUEST_LIMIT:
        return {
            'current_tone': "Limite de requisições atingido",
            'emotional_impact': "Não foi possível analisar",
            'optimized_text': text
        }
    
    # Verificar rate limit
    current_time = time.time()
    if current_time - st.session_state.last_request_time < RATE_LIMIT_SECONDS and st.session_state.request_count > 0:
        wait_time = round(RATE_LIMIT_SECONDS - (current_time - st.session_state.last_request_time))
        return {
            'current_tone': f"Aguarde {wait_time} segundos",
            'emotional_impact': "Não foi possível analisar",
            'optimized_text': text
        }
    
    try:
        with st.spinner("Analisando tom da comunicação..."):
            # Atualizar o timestamp da última requisição
            st.session_state.last_request_time = current_time
            
            # Incrementar contador de requisições
            st.session_state.request_count += 1
            
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {st.session_state.api_key}"
            }
            
            prompt = f"""
            Analise o tom do seguinte texto destinado a '{target_audience}':
            
            "{text}"
            
            Forneça:
            1. Uma análise do tom atual (formal, informal, urgente, técnico, etc.)
            2. O provável impacto emocional no leitor
            3. Uma versão otimizada do texto para ter o impacto de '{desired_impact}'
            
            Retorne a resposta como um dicionário Python com as chaves: 'current_tone', 'emotional_impact', 'optimized_text'
            """
            
            payload = {
                "model": "gpt-3.5-turbo",
                "messages": [
                    {"role": "system", "content": "Você é um especialista em análise de comunicação e impacto emocional em contexto profissional."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.7,
                "max_tokens": 1500
            }
            
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                data=json.dumps(payload),
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                response_text = result['choices'][0]['message']['content']
                
                # Atualizar contadores de tokens
                prompt_tokens = result['usage']['prompt_tokens']
                completion_tokens = result['usage']['completion_tokens']
                total_tokens = result['usage']['total_tokens']
                st.session_state.token_count += total_tokens
                
                # Registrar uso
                st.session_state.usage_data.append({
                    'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'feature': "Análise de Tom",
                    'tokens': total_tokens,
                    'model': "gpt-3.5-turbo",
                    'session_id': st.session_state.session_id
                })
                
                # Tenta extrair o dicionário Python da resposta
                dict_pattern = r'\{.*\}'
                dict_match = re.search(dict_pattern, response_text, re.DOTALL)
                
                if dict_match:
                    dict_str = dict_match.group(0)
                    try:
                        # Substitui aspas simples por aspas duplas para compatibilidade com JSON
                        dict_str = dict_str.replace("'", "\"")
                        result = json.loads(dict_str)
                    except json.JSONDecodeError:
                        # Fallback se o parsing falhar
                        result = {
                            'current_tone': 'Não foi possível analisar precisamente',
                            'emotional_impact': 'Não foi possível analisar precisamente',
                            'optimized_text': text  # Mantém o texto original
                        }
                else:
                    # Se não puder extrair um dicionário, cria um resultado estruturado manualmente
                    lines = response_text.split('\n')
                    result = {
                        'current_tone': next((line.replace('1.', '').strip() for line in lines if '1.' in line), 'Não identificado'),
                        'emotional_impact': next((line.replace('2.', '').strip() for line in lines if '2.' in line), 'Não identificado'),
                        'optimized_text': text  # Default para o texto original
                    }
                    
                    # Tenta encontrar o texto otimizado
                    optimized_section = response_text.split('3.')
                    if len(optimized_section) > 1:
                        result['optimized_text'] = optimized_section[1].strip()
                
                return result
            else:
                return {
                    'current_tone': f"Erro na API (Status {response.status_code})",
                    'emotional_impact': "Não foi possível analisar",
                    'optimized_text': text
                }
                
    except Exception as e:
        return {
            'current_tone': f"Erro na análise: {str(e)}",
            'emotional_impact': "Não foi possível analisar",
            'optimized_text': text
        }

# NOVA FUNÇÃO: Interface para análise de tom
def create_tone_analysis_section(content):
    """
    Cria a seção de análise de tom na interface
    
    Args:
        content (str): O texto gerado a ser analisado
        
    Returns:
        None: Atualiza a UI diretamente
    """
    st.markdown("## 🎭 Análise de Tom e Otimização")
    
    with st.expander("Analisar e otimizar o tom da comunicação", expanded=True):
        col1, col2 = st.columns(2)
        
        with col1:
            target_audience = st.selectbox(
                "Público-alvo:",
                ["Cliente", "Equipe técnica", "Gerência", "Executivos", "Stakeholders", "Público geral"],
                key="tone_audience"
            )
        
        with col2:
            desired_impact = st.selectbox(
                "Impacto desejado:",
                ["Tranquilizar", "Motivar", "Gerar urgência", "Informar objetivamente", 
                 "Persuadir", "Demonstrar empatia", "Mostrar autoridade"],
                key="tone_impact"
            )
        
        if st.button("Analisar e otimizar tom", key="analyze_tone_btn"):
            tone_analysis = analyze_tone(content, target_audience, desired_impact)
            
            st.markdown('<div class="tone-analysis-section">', unsafe_allow_html=True)
            
            st.markdown("### Análise do tom atual")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f'<div class="tone-current"><strong>Tom atual:</strong> {tone_analysis["current_tone"]}</div>', unsafe_allow_html=True)
            
            with col2:
                st.markdown(f'<div class="tone-current"><strong>Impacto emocional provável:</strong> {tone_analysis["emotional_impact"]}</div>', unsafe_allow_html=True)
            
            st.markdown("### Versão otimizada para o impacto desejado")
            st.markdown(f'<div class="tone-optimized">{tone_analysis["optimized_text"]}</div>', unsafe_allow_html=True)
            
            # Armazenar o texto otimizado na sessão
            st.session_state.optimized_content = tone_analysis["optimized_text"]
            
            col1, col2 = st.columns(2)
            with col1:
                # Download como texto
                st.download_button(
                    label="📄 Baixar versão otimizada (TXT)",
                    data=tone_analysis["optimized_text"],
                    file_name=f"comunicacao_otimizada_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain"
                )
            
            with col2:
                # Download como DOCX
                docx_buffer = export_as_docx(tone_analysis["optimized_text"], "Comunicação Otimizada")
                st.download_button(
                    label="📝 Baixar versão otimizada (DOCX)",
                    data=docx_buffer,
                    file_name=f"comunicacao_otimizada_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.markdown('</div>', unsafe_allow_html=True)

# Função para exportar conteúdo como DOCX
def export_as_docx(content, filename="documento"):
    doc = docx.Document()
    
    # Adicionar título
    doc.add_heading(f"{filename}", 0)
    
    # Dividir por linhas e adicionar parágrafos
    paragraphs = content.split('\n')
    for para in paragraphs:
        if para.strip() == "":
            continue
        
        # Verificar se é um cabeçalho
        if re.match(r'^#{1,6}\s+', para):
            # Extrair o nível do cabeçalho e o texto
            header_match = re.match(r'^(#{1,6})\s+(.*)', para)
            if header_match:
                level = min(len(header_match.group(1)), 9)  # Limitar a 9 para evitar erro
                text = header_match.group(2)
                doc.add_heading(text, level)
        else:
            doc.add_paragraph(para)
    
    # Salvar para um buffer em memória
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# ================= SIDEBAR =================

# Sidebar para configuração
with st.sidebar:
    st.markdown('<div class="sidebar-content">', unsafe_allow_html=True)
    st.image("https://via.placeholder.com/150x150.png?text=NEXUS", width=150)
    
    # Mostrar status da API
    st.markdown("### Status")
    if st.session_state.api_key_configured:
        st.success("✅ API configurada automaticamente")
    else:
        st.error("❌ API não configurada. Contate o administrador.")
    
    # Exibir estatísticas de uso
    st.markdown("### Seu Uso")
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Requisições", f"{st.session_state.request_count}/{REQUEST_LIMIT}")
        st.markdown('</div>', unsafe_allow_html=True)
    
    with col2:
        st.markdown('<div class="metric-card">', unsafe_allow_html=True)
        st.metric("Tokens", f"{st.session_state.token_count}/{TOKEN_LIMIT}")
        st.markdown('</div>', unsafe_allow_html=True)
    
    # Barras de progresso
    st.progress(st.session_state.request_count / REQUEST_LIMIT)
    st.caption("Uso de requisições")
    
    st.progress(st.session_state.token_count / TOKEN_LIMIT)
    st.caption("Uso de tokens")
    
    # Informações de sessão (apenas para rastreamento)
    st.markdown('<div class="usage-info">', unsafe_allow_html=True)
    st.caption(f"ID da sessão: {st.session_state.session_id}")
    st.caption(f"Início: {datetime.fromtimestamp(st.session_state.last_request_time).strftime('%H:%M:%S') if st.session_state.last_request_time > 0 else 'N/A'}")
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Feedback
    st.markdown("### Feedback")
    feedback = st.radio("Como está sua experiência?", ["😀 Excelente", "🙂 Boa", "😐 Neutra", "🙁 Ruim", "😞 Péssima"])
    feedback_text = st.text_area("Comentários adicionais")
    
    if st.button("Enviar Feedback"):
        st.success("Feedback enviado. Obrigado por nos ajudar a melhorar!")
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Informação sobre limites de uso
    st.markdown("---")
    st.caption("Esta é uma versão de demonstração com limites de uso para controlar custos. Para uso sem limites, implemente o NEXUS em seu próprio ambiente.")

# ================= MAIN INTERFACE =================

# Interface principal
st.markdown('<h1 class="main-header">NEXUS</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Comunicação de projetos clara, eficaz e profissional</p>', unsafe_allow_html=True)

# Mensagem sobre versão de teste
st.info(f"""
**Versão de Demonstração**
Esta é uma versão de demonstração do NEXUS com limites de uso:
- Máximo de {REQUEST_LIMIT} requisições por sessão
- Máximo de {TOKEN_LIMIT} tokens por sessão
- Tempo mínimo de {RATE_LIMIT_SECONDS} segundos entre requisições
""")

# Histórico de gerações recentes
if st.session_state.history:
    with st.expander("Histórico de Gerações Recentes", expanded=False):
        for i, item in enumerate(reversed(st.session_state.history[-3:])):  # Limitado a 3 itens mais recentes
            st.markdown(f"**{item['timestamp']} - {item['feature']}**")
            if st.button(f"Carregar este conteúdo ↩️", key=f"load_{i}"):
                st.session_state.current_feature = item['feature']
                st.session_state.generated_content = item['output']
                st.experimental_rerun()
            st.markdown("---")

# Seleção de funcionalidade
# Organizar opções em colunas
col1, col2 = st.columns(2)

count = 0
for feature, details in feature_options.items():
    if count % 2 == 0:
        with col1:
            with st.expander(f"{details['icon']} {feature}", expanded=False):
                st.markdown(f"**{details['description']}**")
                if st.button(f"Usar {feature}", key=f"select_{feature}"):
                    st.session_state.current_feature = feature
    else:
        with col2:
            with st.expander(f"{details['icon']} {feature}", expanded=False):
                st.markdown(f"**{details['description']}**")
                if st.button(f"Usar {feature}", key=f"select_{feature}"):
                    st.session_state.current_feature = feature
    count += 1

# ================= FEATURE INTERFACE =================

# Se uma funcionalidade foi selecionada na sessão atual ou anteriormente
if st.session_state.current_feature:
    current_feature = st.session_state.current_feature
    feature_details = feature_options[current_feature]
    
    st.markdown(f"## {feature_details['icon']} {current_feature}")
    
    # Verificar limites antes de mostrar a interface
    if st.session_state.token_count >= TOKEN_LIMIT:
        st.error(f"Você atingiu o limite de {TOKEN_LIMIT} tokens para esta sessão. Por favor, tente novamente mais tarde.")
    elif st.session_state.request_count >= REQUEST_LIMIT:
        st.error(f"Você atingiu o limite de {REQUEST_LIMIT} requisições para esta sessão. Por favor, tente novamente mais tarde.")
    elif time.time() - st.session_state.last_request_time < RATE_LIMIT_SECONDS and st.session_state.request_count > 0:
        wait_time = round(RATE_LIMIT_SECONDS - (time.time() - st.session_state.last_request_time))
        st.warning(f"Por favor, aguarde {wait_time} segundos antes de fazer outra requisição.")
    else:
        # Interface específica da funcionalidade
        with st.form(key=f"{current_feature}_form"):
            st.markdown(f"**{feature_details['description']}**")
            
            # Campo de subtipo
            subtype = st.selectbox("Tipo de Comunicação", feature_details['subtypes'])
            
            # Campos comuns a todas as funcionalidades
            context = st.text_area("Contexto do Projeto", 
                                help="Descreva o projeto, fase atual e informações relevantes",
                                height=100,
                                placeholder="Ex: Projeto de desenvolvimento do aplicativo mobile, fase de testes")
            
            # Campos específicos por funcionalidade
            prompt = ""
            
            if current_feature == "Gerador de Comunicações Estruturadas":
                audience = st.text_input("Público-alvo", 
                                    help="Para quem esta comunicação será enviada (equipe, cliente, stakeholder)",
                                    placeholder="Ex: Cliente, diretor de marketing da empresa XYZ")
                key_points = st.text_area("Pontos-chave", 
                                        help="Liste os principais pontos que devem ser incluídos na comunicação",
                                        height=150,
                                        placeholder="Ex: Atraso de 3 dias devido a bugs na integração; Plano de recuperação com recursos adicionais")
                tone = st.select_slider("Tom da Comunicação", 
                                    options=["Muito Formal", "Formal", "Neutro", "Amigável", "Casual"],
                                    value="Neutro")
                
                prompt = f"""
                Gere um {subtype} com base nas seguintes informações:
                
                Contexto do Projeto: {context}
                Público-alvo: {audience}
                Pontos-chave: {key_points}
                Tom desejado: {tone}
                
                Formate a saída adequadamente para um {subtype}, incluindo assunto/título e estrutura apropriada.
                """
                
            elif current_feature == "Assistente de Reuniões":
                participants = st.text_area("Participantes", 
                                        help="Liste os participantes e suas funções",
                                        height=100,
                                        placeholder="Ex: João Silva (Gerente de Projeto), Maria Costa (Desenvolvedora Frontend)")
                topics = st.text_area("Tópicos a serem abordados", 
                                    help="Liste os tópicos que precisam ser discutidos",
                                    height=150,
                                    placeholder="Ex: Atualização do cronograma, Bugs pendentes, Feedback do cliente")
                duration = st.number_input("Duração (minutos)", min_value=15, max_value=240, value=60, step=15)
                
                if subtype == "Agenda de Reunião":
                    prompt = f"""
                    Crie uma agenda detalhada para uma reunião de {duration} minutos com base nas seguintes informações:
                    
                    Contexto do Projeto: {context}
                    Participantes: {participants}
                    Tópicos a serem abordados: {topics}
                    
                    Inclua alocação de tempo para cada item, responsáveis e objetivos claros.
                    """
                elif subtype == "Ata/Resumo de Reunião":
                    decisions = st.text_area("Decisões tomadas", 
                                        help="Liste as principais decisões tomadas durante a reunião",
                                        height=100,
                                        placeholder="Ex: Aprovação do novo design, Extensão do prazo em 1 semana")
                    actions = st.text_area("Ações acordadas", 
                                        help="Liste as ações acordadas, responsáveis e prazos",
                                        height=100,
                                        placeholder="Ex: João irá corrigir o bug #123 até amanhã, Maria criará novos componentes até sexta")
                    
                    prompt = f"""
                    Crie uma ata/resumo detalhado para uma reunião de {duration} minutos com base nas seguintes informações:
                    
                    Contexto do Projeto: {context}
                    Participantes: {participants}
                    Tópicos abordados: {topics}
                    Decisões tomadas: {decisions}
                    Ações acordadas: {actions}
                    
                    Organize por tópicos, destacando claramente decisões e próximos passos com responsáveis.
                    """
                else:  # Follow-up
                    meeting_outcome = st.text_area("Resultado da reunião", 
                                                help="Resuma os principais resultados da reunião",
                                                height=100,
                                                placeholder="Ex: Definidas as prioridades para o próximo sprint e resolvidos os bloqueios atuais")
                    action_items = st.text_area("Itens de ação", 
                                            help="Liste os itens de ação, responsáveis e prazos",
                                            height=100,
                                            placeholder="Ex: João: revisão de código até 25/03; Maria: implementação da nova feature até 27/03")
                    
                    prompt = f"""
                    Crie uma mensagem de follow-up para uma reunião com base nas seguintes informações:
                    
                    Contexto do Projeto: {context}
                    Participantes: {participants}
                    Tópicos abordados: {topics}
                    Resultado da reunião: {meeting_outcome}
                    Itens de ação: {action_items}
                    
                    A mensagem deve agradecer a participação, resumir os principais pontos, detalhar próximos passos
                    com responsáveis e prazos, e solicitar confirmação/feedback conforme apropriado.
                    """
                    
            elif current_feature == "Tradutor de Jargão Técnico":
                technical_content = st.text_area("Conteúdo Técnico", 
                                            help="Cole aqui o texto técnico que precisa ser traduzido",
                                            height=200,
                                            placeholder="Ex: A implementação do Redux utiliza reducers imutáveis para gerenciar o estado global da aplicação...")
                audience = st.selectbox("Público-alvo", 
                                    ["Executivos", "Clientes não-técnicos", "Equipe de Negócios", "Equipe Técnica Junior"])
                key_concepts = st.text_input("Conceitos-chave a preservar", 
                                        help="Liste conceitos técnicos que devem ser mantidos mesmo se simplificados",
                                        placeholder="Ex: gerenciamento de estado, API, front-end")
                
                prompt = f"""
                Traduza/adapte o seguinte conteúdo técnico para um público de {audience} com base nas seguintes informações:
                
                Contexto do Projeto: {context}
                Conteúdo Técnico Original: {technical_content}
                Conceitos-chave a preservar: {key_concepts}
                
                Para {audience}, foque em: 
                - {'Impacto nos negócios e resultados de alto nível' if audience == 'Executivos' else ''}
                - {'Benefícios e funcionalidades em linguagem acessível' if audience == 'Clientes não-técnicos' else ''}
                - {'Conexão com objetivos de negócios e processos' if audience == 'Equipe de Negócios' else ''}
                - {'Explicações técnicas mais detalhadas, mas com conceitos explicados' if audience == 'Equipe Técnica Junior' else ''}
                
                Mantenha a precisão conceitual mesmo simplificando a linguagem.
                """
                
            elif current_feature == "Facilitador de Feedback":
                situation = st.text_area("Situação", 
                                    help="Descreva a situação específica para a qual você precisa fornecer feedback",
                                    height=150,
                                    placeholder="Ex: Atraso na entrega de componentes para o projeto principal...")
                strengths = st.text_area("Pontos Fortes", 
                                    help="Liste aspectos positivos que devem ser destacados",
                                    height=100,
                                    placeholder="Ex: Qualidade do código entregue, comunicação proativa de desafios")
                areas_for_improvement = st.text_area("Áreas para Melhoria", 
                                                help="Liste aspectos que precisam ser melhorados",
                                                height=100,
                                                placeholder="Ex: Estimativas de tempo irrealistas, falha em pedir ajuda quando bloqueado")
                relationship = st.selectbox("Relação com o Receptor", 
                                        ["Membro da equipe direto", "Colega de mesmo nível", "Superior hierárquico", "Cliente", "Fornecedor"])
                
                prompt = f"""
                Estruture um {subtype} construtivo e eficaz com base nas seguintes informações:
                
                Contexto do Projeto: {context}
                Situação específica: {situation}
                Pontos fortes a destacar: {strengths}
                Áreas para melhoria: {areas_for_improvement}
                Relação com o receptor: {relationship}
                
                O feedback deve:
                - Ser específico e baseado em comportamentos observáveis
                - Equilibrar aspectos positivos e áreas de melhoria
                - Incluir exemplos concretos
                - Oferecer sugestões acionáveis
                - Usar tom apropriado para a relação ({relationship})
                - Focar em crescimento e desenvolvimento, não em crítica
                
                Formate como um roteiro/script que o usuário pode seguir na conversa ou adaptar para uma comunicação escrita.
                """
                
            elif current_feature == "Detector de Riscos de Comunicação":
                content_to_analyze = st.text_area("Conteúdo para Análise", 
                                                help="Cole aqui o texto que você deseja analisar quanto a riscos de comunicação",
                                                height=200,
                                                placeholder="Ex: Devido a circunstâncias imprevistas no desenvolvimento, alguns recursos podem sofrer atrasos...")
                audience = st.text_input("Público-alvo", 
                                    help="Descreva quem receberá esta comunicação",
                                    placeholder="Ex: Cliente executivo com pouco conhecimento técnico")
                stakes = st.select_slider("Importância da Comunicação", 
                                        options=["Baixa", "Média", "Alta", "Crítica"],
                                        value="Média")
                
                prompt = f"""
                Analise o seguinte {subtype} quanto a riscos de comunicação:
                
                Contexto do Projeto: {context}
                Público-alvo: {audience}
                Importância da comunicação: {stakes}
                
                Conteúdo para análise:
                ---
                {content_to_analyze}
                ---
                
                Sua análise deve:
                1. Identificar ambiguidades, informações incompletas ou confusas
                2. Apontar possíveis mal-entendidos baseados no público-alvo
                3. Detectar problemas de tom ou linguagem inapropriada
                4. Identificar informações sensíveis ou potencialmente problemáticas
                5. Sugerir reformulações específicas para cada problema identificado
                
                Organize sua análise em forma de tabela com colunas para: Trecho problemático, Risco potencial, Sugestão de melhoria.
                Ao final, forneça uma avaliação geral dos riscos de comunicação (Baixo/Médio/Alto) e um resumo das principais recomendações.
                """
            
            submit_button = st.form_submit_button(f"Gerar {current_feature}")
        
        if submit_button:
            if not st.session_state.api_key_configured:
                st.error("API não configurada. Por favor, contate o administrador.")
            elif st.session_state.token_count >= TOKEN_LIMIT:
                st.error(f"Você atingiu o limite de {TOKEN_LIMIT} tokens para esta sessão. Por favor, tente novamente mais tarde.")
            elif st.session_state.request_count >= REQUEST_LIMIT:
                st.error(f"Você atingiu o limite de {REQUEST_LIMIT} requisições para esta sessão. Por favor, tente novamente mais tarde.")
            elif time.time() - st.session_state.last_request_time < RATE_LIMIT_SECONDS and st.session_state.request_count > 0:
                wait_time = round(RATE_LIMIT_SECONDS - (time.time() - st.session_state.last_request_time))
                st.warning(f"Por favor, aguarde {wait_time} segundos antes de fazer outra requisição.")
            else:
                # Gerar conteúdo
                generated_content = generate_content(prompt, model="gpt-3.5-turbo", temperature=0.7)
                st.session_state.generated_content = generated_content
                
                # Exibir resultado
                st.markdown("### Resultado")
                st.markdown('<div class="result-area">', unsafe_allow_html=True)
                st.markdown(generated_content)
                st.markdown('</div>', unsafe_allow_html=True)
                
                # Opções de download
                col1, col2 = st.columns(2)
                with col1:
                    # Download como texto
                    st.download_button(
                        label="📄 Baixar como TXT",
                        data=generated_content,
                        file_name=f"{current_feature.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain"
                    )
                
                with col2:
                    # Download como DOCX
                    docx_buffer = export_as_docx(generated_content)
                    st.download_button(
                        label="📝 Baixar como DOCX",
                        data=docx_buffer,
                        file_name=f"{current_feature.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # NOVA SEÇÃO: Análise de Tom
                create_tone_analysis_section(generated_content)
                
                # Feedback sobre o resultado
                st.markdown("### Este resultado foi útil?")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("👍 Sim, foi útil"):
                        st.markdown('<p class="feedback-good">Obrigado pelo feedback positivo!</p>', unsafe_allow_html=True)
                
                with col2:
                    if st.button("👎 Não, preciso de melhoria"):
                        st.markdown('<p class="feedback-bad">Lamentamos que não tenha atendido suas expectativas. Por favor, forneça detalhes no campo de feedback na barra lateral para podermos melhorar.</p>', unsafe_allow_html=True)

# ================= FOOTER =================

# Nota para o artigo
st.write("")
st.write("")
st.markdown("""
---
### Sobre esta demonstração do NEXUS

Este aplicativo é uma versão de demonstração do NEXUS, um assistente de IA para comunicação de projetos. Foi criado como parte do artigo "**Comunicação eficiente em projetos: como a IA pode ajudar gerentes e equipes**".

Para implementar o NEXUS em seu próprio ambiente:
1. Acesse o código-fonte no GitHub: [github.com/seu-usuario/nexus-assistant](https://github.com/seu-usuario/nexus-assistant)
2. Siga as instruções de instalação no README
3. Configure com sua própria chave API da OpenAI
4. Personalize para as necessidades específicas da sua equipe

**Experimente as cinco funcionalidades principais do NEXUS e veja como a IA pode transformar a comunicação nos seus projetos!**
""")

# Rodapé com créditos
st.markdown("""
<div style="text-align: center; color: gray; font-size: 0.8rem;">
    NEXUS | Assistente de Comunicação de Projetos | © 2025
</div>
""", unsafe_allow_html=True)
